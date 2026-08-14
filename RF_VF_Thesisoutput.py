# ============================================
# RANDOM FOREST MODEL SCRIPT
# --------------------------------------------
# Main outputs:
# 1) RF performance (train/test RMSE, MAE, R2)
# 2) Percentage-error metrics (MAPE, MdAPE, sMAPE)
# 3) Group-level error diagnostics
# 4) Feature importance + permutation importance
# 5) Optional Germany residual map by ZIP
#
# Scenario logic:
# - PLZ observation filters (N0, N3, N4, N5, N10)
# - No-imputation vs imputation setup
#
# Report outputs:
# - Word (.docx)
# - Excel (.xlsx)
# ============================================
# SCRIPT NAVIGATION (READ FIRST)
# 1) CONFIGURATION - USER INPUTS
#    Edit target, grouping, debug mode, and map controls.
# 2) FEATURE/ENGINEERING DEFINITIONS
#    Central place for explanatory variables and derived variables.
# 3) HELPER FUNCTIONS
#    Reusable routines for reporting, diagnostics, and plotting.
# 4) RF TRAINING + TUNING
#    Pipeline, randomized search, cross-validation, and evaluation.
# 5) RESIDUAL MAP + EXPORTS
#    ZIP residual aggregation, Germany map rendering, Excel/Word export.
# 6) FINAL COMPARISON OUTPUT
#    Consolidated metrics for scenario-level decision making.

import numpy as np
import pandas as pd
import warnings
from io import BytesIO
from datetime import datetime
from time import perf_counter

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.ticker import FuncFormatter, MaxNLocator
import seaborn as sns

from sklearn.model_selection import GroupShuffleSplit, GroupKFold, RandomizedSearchCV
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder
from sklearn.impute import SimpleImputer
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
from sklearn.inspection import permutation_importance, PartialDependenceDisplay
from sklearn.tree import export_text

from docx import Document
from docx.shared import Inches

from src.geo_residual_map import (
    build_residual_zip_tables,
    plot_germany_residual_map,
)


# ============================================================
# CONFIGURATION - USER INPUTS
# ============================================================

# ------------------------------
# USER GUIDE
# ------------------------------
# In most experiments, you only need to edit:
# 1) MODEL_TARGET              -> total rent or rent_per_sqm
# 2) GROUP_COL                 -> grouped split unit (ZIP recommended)
# 3) DEBUG_MODE settings       -> one scenario vs full grid
# 4) Residual map settings     -> color metric, palette, clipping
#
# Good default practice:
# - Keep GROUP_COL = "obj_zipCode" to preserve spatial validation logic.


FILE_PATH = r"C:\Users\borgesd8828\OneDrive - ARCADIS\Dokumente\GitHub\Real-Estate-Valuation-ImmoScout24\apr20_rental_no_duplicates_for_python.csv"
CSV_SEP = ","


# ============================================================
# MODEL TARGET CONFIGURATION
# ============================================================

TARGET = "obj_totalRent"
TARGET_RENT_SQM = "rent_per_sqm"

# Choose the target used by RF and reporting.
# Option A: total monthly rent (EUR/month).
# Option B: rent per sqm (EUR/sqm).
MODEL_TARGET = TARGET           # total monthly rent
#MODEL_TARGET = TARGET_RENT_SQM    # rent per sqm

TARGET_LABELS = {
    TARGET: "Total Rent",
    TARGET_RENT_SQM: "Rent per sqm"
}

TARGET_UNITS = {
    TARGET: "EUR/month",
    TARGET_RENT_SQM: "EUR/sqm"
}

MODEL_TARGET_LABEL = TARGET_LABELS.get(MODEL_TARGET, MODEL_TARGET)
MODEL_TARGET_UNIT = TARGET_UNITS.get(MODEL_TARGET, "")


# ============================================================
# GROUPING CONFIGURATION
# ============================================================

GROUP_COL = "obj_zipCode"
#GROUP_COL = "obj_regio1"


# ============================================================
# MODEL CONFIGURATION
# ============================================================

RANDOM_STATE = 42
TEST_SIZE = 0.25

REFERENCE_YEAR = 2020

N_ITER_SEARCH = 15
#N_ITER_SEARCH = 100
CV_FOLDS = 5
#CV_FOLDS = 10
# Increase N_ITER_SEARCH/CV_FOLDS for stronger tuning at the cost of runtime.

PERMUTATION_SAMPLE_SIZE = 5000
PERMUTATION_N_REPEATS = 5

MIN_GROUP_ERROR_N = 3

GENERATE_PDP = True

# Global outlier clipping for model inputs.
# This is independent from map color clipping and affects model data directly.
ENABLE_GLOBAL_OUTLIER_CLIPPING = True
# Outlier handling mode:
# - "clip_values": keep all rows and cap extreme numeric values at percentile limits.
# - "remove_rows": remove rows containing outliers in monitored numeric variables.
OUTLIER_TREATMENT_MODE = "remove_rows"
# Percent to cut from each tail (example: 1.0 means P1 and P99).
OUTLIER_CLIP_TAIL_PCT = 0.25
# If True, clipping is also applied to the model target.
OUTLIER_CLIP_APPLY_TO_TARGET = True
# Optional feature names to exclude from clipping.
OUTLIER_CLIP_EXCLUDE_FEATURES = []

# Worked-example configuration for report readability.
EXAMPLE_N_CASES = 5
EXAMPLE_CITY_PATTERNS = [
    ("Berlin", r"berlin"),
    ("Hamburg", r"hamburg"),
    ("Munich", r"muenchen|munich"),
    ("Cologne", r"koeln|cologne"),
    ("Frankfurt", r"frankfurt")
]

# Optional geospatial residual map.
ENABLE_GERMANY_RESIDUAL_MAP = True
ZIP_COL_FOR_MAP = "obj_zipCode"
# Minimum listings per ZIP to include in map summary.
# 1 = broader coverage; 2/3 = more stable ZIP-level averages.
RESIDUAL_MAP_MIN_LISTINGS_PER_ZIP = 0.25
RESIDUAL_MAP_MAX_POINTS = 1200
# Color-scale controls for map bubbles.
# Options for RESIDUAL_MAP_COLOR_COLUMN:
# - mean_residual       (signed error in target unit)
# - mean_abs_residual   (absolute error in target unit)
# - mean_pct_error      (signed percentage error)
# - mape_pct            (absolute percentage error)
# - smape_pct           (symmetric absolute percentage error)
RESIDUAL_MAP_COLOR_COLUMN = "mean_residual"
# Friendly palette name used in this script (user-friendly label).
# red_blue_diverging_reversed -> Matplotlib "RdBu_r" (negative = blue, positive = red)
# red_blue_diverging          -> Matplotlib "RdBu"
# yellow_orange_red           -> Matplotlib "YlOrRd"
# blue_green                  -> Matplotlib "YlGnBu"
RESIDUAL_MAP_PALETTE = "red_blue_diverging_reversed"

MAP_PALETTE_TO_CMAP = {
    "red_blue_diverging_reversed": "RdBu_r",
    "red_blue_diverging": "RdBu",
    "yellow_orange_red": "YlOrRd",
    "blue_green": "YlGnBu",
}

RESIDUAL_MAP_CMAP = MAP_PALETTE_TO_CMAP.get(
    RESIDUAL_MAP_PALETTE,
    "RdBu_r"
)
# Symmetric scale around zero is recommended for signed errors.
RESIDUAL_MAP_SYMMETRIC_SCALE = True
# Percentile clipping reduces outlier domination in color scale.
RESIDUAL_MAP_CLIP_PCT_LOW = 2.0
RESIDUAL_MAP_CLIP_PCT_HIGH = 98.0

if RESIDUAL_MAP_COLOR_COLUMN == "mean_residual" and MODEL_TARGET_UNIT:
    RESIDUAL_MAP_COLOR_LABEL = f"Mean residual ({MODEL_TARGET_UNIT})"
elif RESIDUAL_MAP_COLOR_COLUMN in {"mean_pct_error", "mean_abs_pct_error", "mape_pct", "smape_pct"}:
    RESIDUAL_MAP_COLOR_LABEL = "Percentage error (%)"
else:
    RESIDUAL_MAP_COLOR_LABEL = None


# ============================================================
# EXECUTION MODE
# ============================================================
# DEBUG_MODE = True  -> runs one scenario only (faster iteration).
# DEBUG_MODE = False -> runs full scenario grid (final comparison tables).

DEBUG_MODE = False

# Options: "N0", "N3", "N4", "N5", "N6", "N10"
DEBUG_SCENARIO_CODE = "N6"

# False = NO IMPUTATION
# True  = WITH IMPUTATION
DEBUG_USE_IMPUTATION = True


# ============================================================
# PLZ OBSERVATION SCENARIOS
# ============================================================

PLZ_OBS_SCENARIOS = [
    {"scenario_code": "N0", "min_plz_obs": 0},
    {"scenario_code": "N1", "min_plz_obs": 1},
    {"scenario_code": "N2", "min_plz_obs": 2},
    {"scenario_code": "N3", "min_plz_obs": 3},
    {"scenario_code": "N4", "min_plz_obs": 4},
    {"scenario_code": "N5", "min_plz_obs": 5},
    {"scenario_code": "N6", "min_plz_obs": 6},
    {"scenario_code": "N7", "min_plz_obs": 7},
    {"scenario_code": "N8", "min_plz_obs": 8},
    {"scenario_code": "N9", "min_plz_obs": 9},
    {"scenario_code": "N10", "min_plz_obs": 10}
]

IMPUTATION_SCENARIOS = [False, True]


# ============================================================
# FEATURES
# ============================================================

BASE_FEATURES_FULL = [
    # "obj_yearConstructed",
    "obj_heatingType",
    # "obj_serviceCharge",
    "obj_ExclusiveExpose",
    # "obj_newlyConst",
    "obj_regio1",
    "obj_balcony",
    "obj_cellar",
    "obj_hasKitchen",
    "obj_picturecount",
    "obj_lift",
    "obj_petsAllowed",
    # "obj_interiorQual",
    "obj_condition",
    "obj_livingSpace",
    # "obj_depositLink",
    # "ga_cd_via",
    "obj_typeOfFlat",
    "obj_garden",
    "obj_barrierFree",
    # "obj_zipCode"  # replaced by obj_zip3 (3-digit regional prefix)
]

if MODEL_TARGET == TARGET_RENT_SQM:
    BASE_FEATURES = [
        feature for feature in BASE_FEATURES_FULL
        if feature != "obj_livingSpace"
    ]
else:
    BASE_FEATURES = BASE_FEATURES_FULL.copy()




RAW_ENGINEERING_FEATURES = [
    "obj_energyType",
   # "obj_thermalChar",
    "obj_numberOfFloors",
    "obj_noParkSpaces",
    "obj_lastRefurbish",
    "obj_yearConstructed"
]

ENGINEERED_FEATURES = [
    "obj_energyType_cat",

  #  "obj_hasThermalCharInfo",
  #  "obj_thermalChar_num",

    "obj_hasNumberOfFloorsInfo",
    "obj_numberOfFloors_num",

    "obj_hasParkingInfo",
    "obj_noParkSpaces_num",

    "obj_hasLastRefurbishInfo",
    "obj_yearsSinceLastRefurbish",

   # "obj_hasYearConstructedInfo",
    "obj_buildingAge",

   # "obj_refurbishedAfterConstruction",
    "obj_yearsBetweenConstructionAndRefurbish",
  #  "obj_hasConstructionAndRefurbishInfo",
    "obj_zip3"
]

FEATURES = BASE_FEATURES + ENGINEERED_FEATURES

BINARY_FEATURES = [
    "obj_ExclusiveExpose",
    # "obj_newlyConst",
    "obj_balcony",
    "obj_cellar",
    "obj_hasKitchen",
    "obj_lift",
    "obj_garden",
    "obj_barrierFree",
    # "obj_depositLink",

   # "obj_hasThermalCharInfo",
    "obj_hasNumberOfFloorsInfo",
    "obj_hasParkingInfo",
    "obj_hasLastRefurbishInfo",
  #  "obj_hasYearConstructedInfo",
   # "obj_refurbishedAfterConstruction",
   # "obj_hasConstructionAndRefurbishInfo"
]


# ============================================================
# OUTPUT FILES
# ============================================================

timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
safe_target_name = MODEL_TARGET.replace("/", "_").replace(" ", "_")

OUTPUT_EXCEL = f"RandomForest_{safe_target_name}_PLZObs_{timestamp}.xlsx"
DOCX_OUTPUT = f"RandomForest_{safe_target_name}_PLZObs_{timestamp}.docx"

warnings.filterwarnings("ignore")

# Execution-time marker for full script runtime.
SCRIPT_START_TIME = perf_counter()


# ============================================================
# GLOBAL PLOT STYLE
# ============================================================

plt.rcParams.update({
    "figure.facecolor": "white",
    "axes.facecolor": "white",
    "axes.edgecolor": "black",
    "axes.titlesize": 12,
    "axes.labelsize": 10,
    "xtick.labelsize": 9,
    "ytick.labelsize": 9,
    "legend.fontsize": 9,
    "font.size": 10
})


# ============================================================
# DOCX HELPERS
# ============================================================

doc = Document()
doc.add_heading(
    f"Random Forest Report - {MODEL_TARGET_LABEL}",
    0
)


def add_section(title):
    doc.add_heading(str(title), level=1)


def add_subsection(title):
    doc.add_heading(str(title), level=2)


def add_text(text):
    doc.add_paragraph(str(text))


def add_list(items):
    if not items:
        doc.add_paragraph("(no items)")
    else:
        for item in items:
            doc.add_paragraph(str(item), style="List Bullet")


def add_current_figure_to_doc(width=6):
    """
    Save the current Matplotlib figure to an in-memory buffer
    and insert it directly into the Word document.
    """
    buf = BytesIO()
    plt.savefig(buf, format="png", dpi=300, bbox_inches="tight")
    buf.seek(0)
    doc.add_picture(buf, width=Inches(width))
    buf.close()


def add_figure_to_doc(fig, width=6):
    """
    Save a Matplotlib figure to an in-memory buffer
    and insert it directly into the Word document.
    """
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=300, bbox_inches="tight")
    buf.seek(0)
    doc.add_picture(buf, width=Inches(width))
    buf.close()


def add_model_target_info():
    """
    Add model target information to the Word report.
    """
    add_subsection("Model Target")

    add_text(f"Target variable: {MODEL_TARGET}")
    add_text(f"Description: {MODEL_TARGET_LABEL}")

    if MODEL_TARGET_UNIT:
        add_text(f"Unit: {MODEL_TARGET_UNIT}")

    if MODEL_TARGET == TARGET_RENT_SQM:
        add_text(
            "The target variable represents rent per square meter and was calculated as "
            "obj_totalRent divided by obj_livingSpace."
        )
        add_text(
            "Living space was excluded from the explanatory variables because it is embedded "
            "in the dependent variable."
        )

    elif MODEL_TARGET == TARGET:
        add_text(
            "The target variable represents total monthly rent. Living space is included as an "
            "explanatory variable."
        )


# ============================================================
# GENERAL HELPERS
# ============================================================

def filter_by_min_plz_observations(X, y, groups, min_plz_obs):
    """
    Filter observations according to the minimum number of listings per PLZ.

    N0  -> no minimum PLZ-count filter
    N3  -> keep only PLZs with at least 3 observations
    N5  -> keep only PLZs with at least 5 observations
    N10 -> keep only PLZs with at least 10 observations
    """

    if min_plz_obs <= 0:
        removed_group_counts = pd.DataFrame(columns=[GROUP_COL, "n_obs"])
        return X.copy(), y.copy(), groups.copy(), removed_group_counts

    group_counts = groups.value_counts()

    valid_groups = group_counts[
        group_counts >= min_plz_obs
    ].index

    removed_group_counts = (
        group_counts[group_counts < min_plz_obs]
        .reset_index()
    )

    removed_group_counts.columns = [
        GROUP_COL,
        "n_obs"
    ]

    valid_mask = groups.isin(valid_groups)

    X_filtered = X.loc[valid_mask].copy()
    y_filtered = y.loc[valid_mask].copy()
    groups_filtered = groups.loc[valid_mask].copy()

    return X_filtered, y_filtered, groups_filtered, removed_group_counts


def _pretty_tick(value, _pos):
    """Readable numeric ticks using EU separators (e.g., 1.200,5)."""
    def _format_eu_number(val, decimals):
        text = f"{val:,.{decimals}f}"
        return text.replace(",", "#").replace(".", ",").replace("#", ".")

    abs_val = abs(value)
    if abs_val >= 1000:
        return _format_eu_number(value, 0)
    if abs_val >= 100:
        return _format_eu_number(value, 1)
    if abs_val >= 1:
        return _format_eu_number(value, 2)
    return _format_eu_number(value, 3)


def format_axis_pretty(ax):
    """Apply consistent tick density and formatting to report plots."""
    ax.xaxis.set_major_locator(MaxNLocator(nbins=6))
    ax.yaxis.set_major_locator(MaxNLocator(nbins=6))
    ax.xaxis.set_major_formatter(FuncFormatter(_pretty_tick))
    ax.yaxis.set_major_formatter(FuncFormatter(_pretty_tick))


def get_symmetric_residual_ylim(values, percentile=99, pad_ratio=0.12):
    """Build symmetric y-limits around zero for residual-style plots."""
    arr = np.asarray(values, dtype=float)
    arr = arr[np.isfinite(arr)]

    if arr.size == 0:
        return (-1.0, 1.0)

    robust_abs = float(np.nanpercentile(np.abs(arr), percentile))
    robust_abs = max(robust_abs, 1e-9)
    lim = robust_abs * (1.0 + pad_ratio)
    return (-lim, lim)


def get_robust_limits(values, low_pct=1, high_pct=99, pad_ratio=0.08):
    """Compute robust axis limits from percentiles with small padding."""
    arr = np.asarray(values, dtype=float)
    arr = arr[np.isfinite(arr)]

    if arr.size == 0:
        return (-1.0, 1.0)

    low = float(np.nanpercentile(arr, low_pct))
    high = float(np.nanpercentile(arr, high_pct))
    span = max(high - low, 1e-9)
    pad = span * pad_ratio
    return (low - pad, high + pad)


def add_debug_variable_inventory(df_model: pd.DataFrame, selected_features):
    """Add a debug-only inventory of columns available to the RF scenario report."""
    if not DEBUG_MODE:
        return

    available_columns = sorted(df_model.columns.tolist())
    selected_set = set(selected_features)
    non_model_columns = sorted([col for col in available_columns if col not in selected_set])

    add_subsection("DEBUG VARIABLE INVENTORY")
    add_text(f"Total columns available in scenario dataframe: {len(available_columns)}")
    add_text("All columns currently available to the scenario:")
    add_list(available_columns)
    add_text(f"Columns not used as direct model features: {len(non_model_columns)}")
    add_list(non_model_columns)


OK_COLOR = "#2ca02c"
NOT_OK_COLOR = "#d62728"
CHECK_COLOR = "#ff7f0e"


def pass_fail_label_from_threshold(value, ok_threshold, check_threshold=None, lower_is_better=True):
    """Heuristic status label for RF diagnostics where formal tests are not available."""
    if pd.isna(value):
        return "CHECK", CHECK_COLOR

    if lower_is_better:
        if value <= ok_threshold:
            return "OK", OK_COLOR
        if check_threshold is not None and value <= check_threshold:
            return "CHECK", CHECK_COLOR
        return "NOT OK", NOT_OK_COLOR

    if value >= ok_threshold:
        return "OK", OK_COLOR
    if check_threshold is not None and value >= check_threshold:
        return "CHECK", CHECK_COLOR
    return "NOT OK", NOT_OK_COLOR


def add_status_box(ax, top_text, bottom_text, label, label_color):
    """Add a compact status annotation box inside an axis."""
    annotation_text = f"{top_text}\n{bottom_text}"

    ax.text(
        0.03,
        0.97,
        annotation_text,
        transform=ax.transAxes,
        va="top",
        ha="left",
        fontsize=9,
        bbox=dict(
            boxstyle="round,pad=0.3",
            facecolor="white",
            edgecolor="lightgray",
            alpha=0.95
        )
    )

    ax.text(
        0.97,
        0.06,
        label,
        transform=ax.transAxes,
        va="bottom",
        ha="right",
        fontsize=16,
        fontweight="bold",
        color=label_color
    )


def create_feature_missing_report(X):
    """
    Create a descriptive missing-value report.
    This report is exported only for documentation.
    It is NOT used for feature selection.
    """

    missing_share = X.isna().mean().sort_values(ascending=False)

    missing_share_df = (
        missing_share
        .reset_index()
        .rename(columns={"index": "feature", 0: "missing_share"})
    )

    return missing_share_df


def normalize_binary_series(s: pd.Series):
    """
    Normalize binary variables to 0.0 / 1.0.
    Accepted values include Y/N, TRUE/FALSE, 1/0, 1.0/0.0.
    """

    s = s.astype("object")
    s = s.where(pd.notna(s), np.nan)

    def _norm(v):
        if pd.isna(v):
            return np.nan

        if isinstance(v, (int, float, np.integer, np.floating)):
            if np.isclose(v, 1):
                return 1.0
            if np.isclose(v, 0):
                return 0.0
            return np.nan

        v = str(v).strip().upper()

        if v in {"Y", "YES", "TRUE", "T", "1", "1.0"}:
            return 1.0

        if v in {"N", "NO", "FALSE", "F", "0", "0.0"}:
            return 0.0

        return np.nan

    return s.map(_norm).astype("float")


def apply_global_outlier_clipping(
    X: pd.DataFrame,
    y: pd.Series,
    tail_pct: float,
    apply_to_target: bool,
    mode: str = "clip_values",
    exclude_features=None
):
    """
        Apply global outlier treatment on numeric variables by percentile tails.

        mode="clip_values":
            - keep all rows and clip values to [P_low, P_high].

        mode="remove_rows":
            - remove rows where at least one monitored variable is outside [P_low, P_high].
    """

    if exclude_features is None:
        exclude_features = []

    X_out = X.copy()
    y_out = y.copy()

    try:
        tail_pct = float(tail_pct)
    except Exception:
        tail_pct = 0.0

    if tail_pct <= 0 or tail_pct >= 50:
        empty_thresholds = pd.DataFrame(columns=["variable", "low", "high", "n_rows_flagged"])
        empty_excluded = pd.DataFrame(columns=["row_index", "trigger_variables", "trigger_details", "n_trigger_variables"])
        return X_out, y_out, empty_thresholds, empty_excluded

    low_q = tail_pct
    high_q = 100.0 - tail_pct

    clipping_rows = []
    outlier_flags = pd.DataFrame(index=X_out.index)

    numeric_cols = X_out.select_dtypes(include=["number"]).columns.tolist()
    for col in numeric_cols:
        if col in exclude_features:
            continue

        non_na = X_out[col].dropna()
        if non_na.empty:
            continue

        unique_vals = set(non_na.unique())
        if unique_vals.issubset({0, 1}):
            continue

        low_val = np.nanpercentile(non_na.to_numpy(dtype=float), low_q)
        high_val = np.nanpercentile(non_na.to_numpy(dtype=float), high_q)

        if np.isfinite(low_val) and np.isfinite(high_val) and low_val < high_val:
            col_values = X_out[col].to_numpy(dtype=float)
            col_flags = (col_values < low_val) | (col_values > high_val)
            outlier_flags[col] = col_flags

            if mode == "clip_values":
                X_out[col] = X_out[col].clip(lower=low_val, upper=high_val)

            clipping_rows.append(
                {
                    "variable": col,
                    "low": low_val,
                    "high": high_val,
                    "n_rows_flagged": int(np.nansum(col_flags))
                }
            )

    if apply_to_target:
        y_non_na = y_out.dropna()
        if not y_non_na.empty:
            y_low = np.nanpercentile(y_non_na.to_numpy(dtype=float), low_q)
            y_high = np.nanpercentile(y_non_na.to_numpy(dtype=float), high_q)
            if np.isfinite(y_low) and np.isfinite(y_high) and y_low < y_high:
                y_values = y_out.to_numpy(dtype=float)
                y_flags = (y_values < y_low) | (y_values > y_high)
                outlier_flags[MODEL_TARGET] = y_flags

                if mode == "clip_values":
                    y_out = y_out.clip(lower=y_low, upper=y_high)

                clipping_rows.append(
                    {
                        "variable": MODEL_TARGET,
                        "low": y_low,
                        "high": y_high,
                        "n_rows_flagged": int(np.nansum(y_flags))
                    }
                )

    excluded_rows_df = pd.DataFrame(columns=["row_index", "trigger_variables", "trigger_details", "n_trigger_variables"])

    if mode == "remove_rows" and not outlier_flags.empty:
        row_outlier_mask = outlier_flags.any(axis=1)
        excluded_index = outlier_flags.index[row_outlier_mask]
        threshold_map = {
            row["variable"]: (row["low"], row["high"])
            for _, row in pd.DataFrame(clipping_rows).iterrows()
        }

        excluded_rows = []
        for idx in excluded_index:
            triggers = outlier_flags.columns[outlier_flags.loc[idx].to_numpy(dtype=bool)].tolist()
            details = []

            for var in triggers:
                low_val, high_val = threshold_map.get(var, (np.nan, np.nan))

                if var == MODEL_TARGET:
                    observed = y_out.loc[idx] if idx in y_out.index else np.nan
                else:
                    observed = X_out.loc[idx, var] if (idx in X_out.index and var in X_out.columns) else np.nan

                side = "inside"
                if pd.notna(observed) and np.isfinite(low_val) and observed < low_val:
                    side = "below_low"
                elif pd.notna(observed) and np.isfinite(high_val) and observed > high_val:
                    side = "above_high"

                if pd.notna(observed):
                    observed_txt = f"{float(observed):.4f}"
                else:
                    observed_txt = "NA"

                details.append(
                    f"{var}: observed={observed_txt}, ref=[{low_val:.4f}, {high_val:.4f}], side={side}"
                )

            excluded_rows.append(
                {
                    "row_index": int(idx) if isinstance(idx, (int, np.integer)) else str(idx),
                    "trigger_variables": ", ".join(triggers),
                    "trigger_details": " | ".join(details),
                    "n_trigger_variables": len(triggers)
                }
            )

        excluded_rows_df = pd.DataFrame(excluded_rows)

        keep_mask = ~row_outlier_mask
        X_out = X_out.loc[keep_mask].copy()
        y_out = y_out.loc[keep_mask].copy()

    clipping_df = pd.DataFrame(clipping_rows)
    return X_out, y_out, clipping_df, excluded_rows_df


def make_ohe():
    """
    Create OneHotEncoder compatible with different sklearn versions.
    """

    try:
        return OneHotEncoder(
            drop="first",
            handle_unknown="ignore",
            sparse_output=False
        )
    except TypeError:
        return OneHotEncoder(
            drop="first",
            handle_unknown="ignore",
            sparse=False
        )


def add_engineered_features(
    df_input: pd.DataFrame,
    reference_year=REFERENCE_YEAR
):
    """
    Create engineered features from variables with many missing values.

    Important:
    - The dataset represents ImmoScout24 listings from April 2020.
    - Therefore, age-related variables are calculated using REFERENCE_YEAR = 2020.
    """

    df = df_input.copy()
    current_year = reference_year

    # Energy type
    if "obj_energyType" in df.columns:
        df["obj_energyType_cat"] = (
            df["obj_energyType"]
            .astype("object")
            .where(pd.notna(df["obj_energyType"]), "Unknown")
            .astype(str)
        )
    else:
        df["obj_energyType_cat"] = "Unknown"

    # Thermal characteristic
    if "obj_thermalChar" in df.columns:
        thermal_num = pd.to_numeric(df["obj_thermalChar"], errors="coerce")
        df["obj_thermalChar_num"] = thermal_num.fillna(0)
    else:
        df["obj_thermalChar_num"] = 0.0

    # Number of floors
    if "obj_numberOfFloors" in df.columns:
        number_of_floors_num = pd.to_numeric(
            df["obj_numberOfFloors"],
            errors="coerce"
        )

        valid_floors = number_of_floors_num.notna()

        df["obj_hasNumberOfFloorsInfo"] = valid_floors.astype(float)
        df["obj_numberOfFloors_num"] = number_of_floors_num.fillna(0)
    else:
        df["obj_hasNumberOfFloorsInfo"] = 0.0
        df["obj_numberOfFloors_num"] = 0.0

    # Parking spaces
    if "obj_noParkSpaces" in df.columns:
        parking_num = pd.to_numeric(df["obj_noParkSpaces"], errors="coerce")
        valid_parking = parking_num.notna()

        df["obj_hasParkingInfo"] = valid_parking.astype(float)
        df["obj_noParkSpaces_num"] = parking_num.fillna(0)
    else:
        df["obj_hasParkingInfo"] = 0.0
        df["obj_noParkSpaces_num"] = 0.0

    # Last refurbishment
    if "obj_lastRefurbish" in df.columns:
        last_refurbish_year = pd.to_numeric(
            df["obj_lastRefurbish"],
            errors="coerce"
        )

        valid_refurbish = (
            last_refurbish_year.notna()
            & (last_refurbish_year > 1800)
            & (last_refurbish_year <= current_year)
        )

        df["obj_hasLastRefurbishInfo"] = valid_refurbish.astype(float)

        df["obj_yearsSinceLastRefurbish"] = np.where(
            valid_refurbish,
            current_year - last_refurbish_year,
            0
        )
    else:
        last_refurbish_year = pd.Series(np.nan, index=df.index)
        valid_refurbish = pd.Series(False, index=df.index)

        df["obj_hasLastRefurbishInfo"] = 0.0
        df["obj_yearsSinceLastRefurbish"] = 0.0

    # Year constructed
    if "obj_yearConstructed" in df.columns:
        constructed_year = pd.to_numeric(
            df["obj_yearConstructed"],
            errors="coerce"
        )

        valid_constructed = (
            constructed_year.notna()
            & (constructed_year > 1800)
            & (constructed_year <= current_year)
        )

        df["obj_buildingAge"] = np.where(
            valid_constructed,
            current_year - constructed_year,
            0
        )
    else:
        constructed_year = pd.Series(np.nan, index=df.index)
        valid_constructed = pd.Series(False, index=df.index)

        df["obj_buildingAge"] = 0.0

    # Construction-refurbishment difference
    valid_both = valid_refurbish & valid_constructed

    df["obj_yearsBetweenConstructionAndRefurbish"] = np.where(
        valid_both & (last_refurbish_year >= constructed_year),
        last_refurbish_year - constructed_year,
        0
    )

    # ----------------------------------------
    # 8) ZIP 3-digit regional prefix
    # ----------------------------------------
    if "obj_zipCode" in df.columns:
        df["obj_zip3"] = (
            df["obj_zipCode"].astype(str)
            .str.replace(r"\.0$", "", regex=True)
            .str.zfill(5)
            .str[:3]
        )
    else:
        df["obj_zip3"] = "unknown"

    return df


def make_preprocessor(numeric_features, categorical_features, use_imputation=False):
    """
    Build the preprocessing object for both scenarios:

    - no imputation:
        numeric features are passed directly
        categorical features are one-hot encoded

    - with imputation:
        numeric features are imputed using the median
        categorical features are imputed using the most frequent category
        and then one-hot encoded
    """

    if use_imputation:
        numeric_transformer = Pipeline([
            ("imputer", SimpleImputer(strategy="median"))
        ])

        categorical_transformer = Pipeline([
            ("imputer", SimpleImputer(strategy="most_frequent")),
            ("onehot", make_ohe())
        ])
    else:
        numeric_transformer = "passthrough"

        categorical_transformer = Pipeline([
            ("onehot", make_ohe())
        ])

    preprocessor = ColumnTransformer(
        transformers=[
            ("num", numeric_transformer, numeric_features),
            ("cat", categorical_transformer, categorical_features)
        ],
        remainder="drop",
        sparse_threshold=0
    )

    return preprocessor


def regression_metrics(y_true, y_pred, prefix):
    """
    Compute RMSE, MAE and R2.
    """

    rmse = float(np.sqrt(mean_squared_error(y_true, y_pred)))
    mae = float(mean_absolute_error(y_true, y_pred))
    r2 = float(r2_score(y_true, y_pred))

    return {
        f"{prefix}_RMSE": rmse,
        f"{prefix}_MAE": mae,
        f"{prefix}_R2": r2
    }


def add_pearson_correlation_plot(X_input, y_input, target, scenario_name, use_imputation=False):
    """
    Create Pearson correlation visualizations using numeric variables only.
    Returns a DataFrame with feature-to-target Pearson correlations.
    """
    add_section(f"PEARSON CORRELATION - {scenario_name}")

    df_corr = X_input.copy()
    df_corr[target] = y_input

    numeric_corr_df = df_corr.select_dtypes(include=["number"]).copy()

    if target not in numeric_corr_df.columns:
        add_text("The target is not available as a numeric variable for Pearson correlation.")
        return pd.DataFrame(columns=["feature", "pearson_corr"])

    numeric_corr_df = numeric_corr_df.dropna(axis=1, how="all")

    if numeric_corr_df.shape[1] < 2:
        add_text("Too few numeric variables are available for Pearson correlation analysis.")
        return pd.DataFrame(columns=["feature", "pearson_corr"])

    if use_imputation:
        for col in numeric_corr_df.columns:
            if numeric_corr_df[col].isna().any():
                numeric_corr_df[col] = numeric_corr_df[col].fillna(numeric_corr_df[col].median())
        add_text("Pearson correlation was computed after median imputation for numeric variables.")
    else:
        add_text("Pearson correlation was computed using available data. Pandas uses pairwise valid observations.")

    corr_matrix = numeric_corr_df.corr(method="pearson")

    fig, ax = plt.subplots(figsize=(14, 10))
    sns.heatmap(
        corr_matrix,
        annot=True,
        fmt=".2f",
        cmap="coolwarm",
        center=0,
        square=False,
        linewidths=0.3,
        annot_kws={"size": 10},
        cbar_kws={"shrink": 0.8},
        ax=ax
    )
    ax.set_title(f"Pearson Correlation Matrix - {scenario_name}", fontsize=16)
    plt.xticks(rotation=45, ha="right", fontsize=8)
    plt.yticks(rotation=0, fontsize=8)
    plt.tight_layout()

    add_figure_to_doc(fig, width=7.2)
    plt.close(fig)

    add_subsection(f"Pearson Correlation with {target}")

    target_corr = (
        corr_matrix[target]
        .drop(target)
        .sort_values(key=lambda x: x.abs(), ascending=False)
    )

    add_text("Top Pearson correlations with the target:")
    for feature, corr_value in target_corr.head(25).items():
        add_text(f"{feature}: {corr_value:.4f}")

    target_corr_plot = target_corr.head(25).sort_values()

    fig, ax = plt.subplots(figsize=(10, 8))
    ax.barh(target_corr_plot.index, target_corr_plot.values, color="forestgreen")
    ax.axvline(x=0, color="black", linewidth=0.8)
    ax.set_xlabel("Pearson correlation with target")
    ax.set_ylabel("Feature")
    ax.set_title(f"Top Pearson Correlations with {target} - {scenario_name}")
    plt.tight_layout()

    add_figure_to_doc(fig, width=6.8)
    plt.close(fig)

    add_text(
        "Note: Pearson correlation measures linear relationships among numeric variables. "
        "Pure categorical variables were excluded before one-hot encoding."
    )

    return pd.DataFrame({
        "feature": target_corr.index,
        "pearson_corr": target_corr.values
    })


def add_rf_prediction_examples_and_tree(
    scenario_name,
    X_test,
    y_test,
    y_pred_test,
    rf_model,
    feature_names_out,
    n_examples=5
):
    """
    Add worked RF examples and one readable decision-tree excerpt to the Word report.
    """
    if len(y_test) == 0:
        return

    add_section(f"WORKED RF EXAMPLES - {scenario_name}")
    add_text(
        "The examples below compare actual and predicted rent values for selected observations."
    )

    x_index = X_test.index
    selected_idx = []

    if "obj_regio1" in X_test.columns:
        regio = X_test["obj_regio1"].astype(str)
        for _, pattern in EXAMPLE_CITY_PATTERNS:
            city_idx = x_index[regio.str.contains(pattern, case=False, na=False, regex=True)]
            if len(city_idx) > 0 and city_idx[0] not in selected_idx:
                selected_idx.append(city_idx[0])
            if len(selected_idx) >= n_examples:
                break

    if len(selected_idx) < n_examples:
        remaining = [idx for idx in x_index if idx not in selected_idx]
        for idx in remaining:
            selected_idx.append(idx)
            if len(selected_idx) >= n_examples:
                break

    selected_idx = selected_idx[:n_examples]

    y_test_series = y_test.copy()
    y_pred_series = pd.Series(np.asarray(y_pred_test, dtype=float), index=y_test.index)

    for k, idx in enumerate(selected_idx, start=1):
        add_subsection(f"Example {k}")
        add_text(f"Row index: {idx}")
        if "obj_regio1" in X_test.columns:
            add_text(f"Region: {X_test.loc[idx, 'obj_regio1']}")
        if "obj_zipCode" in X_test.columns:
            add_text(f"ZIP: {X_test.loc[idx, 'obj_zipCode']}")

        actual_val = float(y_test_series.loc[idx]) if idx in y_test_series.index else np.nan
        pred_val = float(y_pred_series.loc[idx]) if idx in y_pred_series.index else np.nan
        err_val = actual_val - pred_val if np.isfinite(actual_val) and np.isfinite(pred_val) else np.nan

        add_text(f"Actual rent: {actual_val:.4f}")
        add_text(f"Predicted rent: {pred_val:.4f}")
        add_text(f"Residual (actual - predicted): {err_val:.4f}")

    add_subsection("Example Decision Tree (text excerpt)")
    try:
        tree_text = export_text(
            rf_model.estimators_[0],
            feature_names=list(feature_names_out),
            max_depth=3,
            decimals=2
        )
        add_text("This is one tree from the forest (depth limited for readability).")
        doc.add_paragraph(tree_text)
    except Exception as exc:
        add_text(f"Decision tree text export failed: {exc}")


def clean_sheet_name(name):
    """
    Excel sheet names must be <= 31 characters.
    """

    return (
        str(name)
        .replace(" ", "_")
        .replace("-", "_")
        .replace("/", "_")
        .replace("\\", "_")
    )[:31]


# ============================================================
# SCENARIO RUNNER
# ============================================================

def run_rf_scenario(df_original, scenario_name, use_imputation, min_plz_obs):
    """
    Run one Random Forest scenario.

    Scenario structure:
    - N0, N3, N5, N10 based on minimum number of observations per PLZ
    - NO IMPUTATION or WITH IMPUTATION
    - Grouped train/test split by postal code
    - Grouped RandomizedSearchCV
    - Train/test metrics
    - OOB R2
    - Internal feature importance
    - Permutation importance
    - PDP
    - Group-level error analysis
    """

    add_section(f"SCENARIO: {scenario_name}")
    add_model_target_info()

    # =========================
    # 1) PREPARE DATA
    # =========================

    df_model = df_original.copy()

    df_model[TARGET] = pd.to_numeric(df_model[TARGET], errors="coerce")

    if TARGET_RENT_SQM in df_model.columns:
        df_model[TARGET_RENT_SQM] = pd.to_numeric(
            df_model[TARGET_RENT_SQM],
            errors="coerce"
        )

    # Create engineered features before selecting model columns
    df_model = add_engineered_features(df_model)

    rows_initial = len(df_model)

    # The model target is never imputed
    df_model = df_model.dropna(subset=[MODEL_TARGET])

    # Group column is required for PLZ filtering and grouped validation
    df_model = df_model.dropna(subset=[GROUP_COL])

    rows_after_model_target_group_drop = len(df_model)

    X = df_model[FEATURES].copy()
    y = df_model[MODEL_TARGET].copy()
    groups = df_model[GROUP_COL].copy()

    selected_features = FEATURES.copy()
    feature_missing_share_df = create_feature_missing_report(X)

    add_subsection("PLZ Observation Scenario")
    add_text(f"Scenario name: {scenario_name}")
    add_text(f"Model target: {MODEL_TARGET}")
    add_text(f"Minimum number of observations per PLZ: {min_plz_obs}")
    add_text(f"Number of model features before preprocessing: {len(selected_features)}")

    add_text("Model features:")
    add_list(selected_features)
    add_debug_variable_inventory(df_model, selected_features)

    # =========================
    # 2) NORMALIZE BINARY FEATURES
    # =========================

    for col in BINARY_FEATURES:
        if col in X.columns:
            X[col] = normalize_binary_series(X[col])

    binary_all_nan = [
        c for c in BINARY_FEATURES
        if c in X.columns and X[c].notna().sum() == 0
    ]

    if binary_all_nan:
        add_subsection("Binary Feature Warning")
        add_text(
            "The following binary variables were not recognized and were treated as categorical:"
        )
        add_list(binary_all_nan)

        for c in binary_all_nan:
            X[c] = df_model[c].astype("object")

    # =========================
    # 2.1) GLOBAL OUTLIER CLIPPING (USER-CONFIGURED)
    # =========================

    if ENABLE_GLOBAL_OUTLIER_CLIPPING:
        rows_before_outlier_treatment = len(X)

        X, y, clipping_df, excluded_rows_df = apply_global_outlier_clipping(
            X=X,
            y=y,
            tail_pct=OUTLIER_CLIP_TAIL_PCT,
            apply_to_target=OUTLIER_CLIP_APPLY_TO_TARGET,
            mode=OUTLIER_TREATMENT_MODE,
            exclude_features=OUTLIER_CLIP_EXCLUDE_FEATURES
        )

        groups = groups.loc[X.index].copy()

        rows_after_outlier_treatment = len(X)
        rows_removed_outliers = rows_before_outlier_treatment - rows_after_outlier_treatment
        variables_flagged = int((clipping_df["n_rows_flagged"] > 0).sum()) if not clipping_df.empty else 0

        add_subsection("Global Outlier Clipping")
        add_text(f"Enabled: {ENABLE_GLOBAL_OUTLIER_CLIPPING}")
        add_text(f"Mode: {OUTLIER_TREATMENT_MODE}")
        add_text(f"Tail percentage per side: {OUTLIER_CLIP_TAIL_PCT}%")
        add_text(f"Target clipping enabled: {OUTLIER_CLIP_APPLY_TO_TARGET}")
        add_text(f"Excluded features: {OUTLIER_CLIP_EXCLUDE_FEATURES}")
        add_text(f"Variables monitored: {len(clipping_df)}")
        add_text(f"Variables that triggered outliers: {variables_flagged}")
        add_text(f"Rows removed by outlier rule: {rows_removed_outliers}")

        add_text(
            "Detailed outlier listings are hidden to keep the report concise. "
            "Use this summary to compare scenarios."
        )
    else:
        add_subsection("Global Outlier Clipping")
        add_text("Enabled: False")

    # =========================
    # 3) MISSING VALUE HANDLING
    # =========================

    if not use_imputation:
        # No imputation:
        # observations with missing values in any selected model feature are removed.
        valid_mask = X.notna().all(axis=1) & y.notna() & groups.notna()

        X = X.loc[valid_mask].copy()
        y = y.loc[valid_mask].copy()
        groups = groups.loc[valid_mask].copy()

    else:
        # With imputation:
        # only target and group must be non-missing.
        # feature missing values are handled inside the preprocessing pipeline.
        valid_mask = y.notna() & groups.notna()

        X = X.loc[valid_mask].copy()
        y = y.loc[valid_mask].copy()
        groups = groups.loc[valid_mask].copy()

    rows_after_missing_handling = len(X)

    # =========================
    # 4) FILTER BY MINIMUM PLZ OBSERVATIONS
    # =========================

    rows_before_plz_filter = len(X)
    groups_before_plz_filter = groups.nunique()

    X, y, groups, removed_plz_groups_df = filter_by_min_plz_observations(
        X=X,
        y=y,
        groups=groups,
        min_plz_obs=min_plz_obs
    )

    rows_after_plz_filter = len(X)
    groups_after_plz_filter = groups.nunique()

    add_subsection("PLZ Group Filtering")
    add_text(f"Rows before PLZ filtering: {rows_before_plz_filter}")
    add_text(f"PLZ groups before filtering: {groups_before_plz_filter}")
    add_text(f"Rows after PLZ filtering: {rows_after_plz_filter}")
    add_text(f"PLZ groups after filtering: {groups_after_plz_filter}")
    add_text(f"Removed PLZ groups: {len(removed_plz_groups_df)}")

    if not removed_plz_groups_df.empty:
        add_text("Examples of removed PLZ groups:")
        for _, row in removed_plz_groups_df.head(20).iterrows():
            add_text(f"{row[GROUP_COL]}: n={row['n_obs']}")

    rows_final = len(X)

    add_subsection("Rows Used")
    add_text(f"Initial rows: {rows_initial}")
    add_text(
        f"Rows after dropping missing model target/group: "
        f"{rows_after_model_target_group_drop}"
    )
    add_text(f"Rows after missing value handling: {rows_after_missing_handling}")
    add_text(f"Rows used after PLZ filtering: {rows_final}")
    add_text(f"Use imputation: {use_imputation}")

    if ZIP_COL_FOR_MAP not in df_model.columns:
        add_text(
            f"Optional residual map warning: '{ZIP_COL_FOR_MAP}' is not available. "
            "The map export will be skipped for this run."
        )

    if rows_final == 0:
        add_text("No rows are available for this scenario. The scenario was skipped.")
        return None

    # =========================
    # 5) VARIABLE TYPES
    # =========================

    numeric_features = X.select_dtypes(include=["number"]).columns.tolist()
    categorical_features = [c for c in X.columns if c not in numeric_features]

    binary_detected = []
    numeric_non_binary = []

    for col in numeric_features:
        unique_vals = set(X[col].dropna().unique())

        if unique_vals.issubset({0, 1}):
            binary_detected.append(col)
        else:
            numeric_non_binary.append(col)

    add_subsection("Variable Type Summary")

    add_text(f"Number of binary variables: {len(binary_detected)}")
    add_list(binary_detected)

    add_text(f"Number of numeric non-binary variables: {len(numeric_non_binary)}")
    add_list(numeric_non_binary)

    add_text(f"Number of categorical variables: {len(categorical_features)}")
    add_list(categorical_features)

    # =========================
    # 5.1) PEARSON CORRELATION
    # =========================
    pearson_target_corr_df = add_pearson_correlation_plot(
        X_input=X,
        y_input=y,
        target=MODEL_TARGET,
        scenario_name=scenario_name,
        use_imputation=use_imputation
    )

    # =========================
    # 6) GROUP CHECK
    # =========================

    n_groups = groups.nunique()

    if n_groups < CV_FOLDS:
        add_text(
            f"Scenario skipped because GROUP_COL='{GROUP_COL}' has only "
            f"{n_groups} unique groups, but CV_FOLDS={CV_FOLDS}."
        )
        return None

    # =========================
    # 7) GROUPED TRAIN/TEST SPLIT
    # =========================

    gss = GroupShuffleSplit(
        n_splits=1,
        test_size=TEST_SIZE,
        random_state=RANDOM_STATE
    )

    train_idx, test_idx = next(gss.split(X, y, groups=groups))

    X_train = X.iloc[train_idx].copy()
    X_test = X.iloc[test_idx].copy()

    y_train = y.iloc[train_idx].copy()
    y_test = y.iloc[test_idx].copy()

    groups_train = groups.iloc[train_idx].copy()
    groups_test = groups.iloc[test_idx].copy()

    train_groups_unique = set(groups_train.unique())
    test_groups_unique = set(groups_test.unique())
    overlap_groups = train_groups_unique.intersection(test_groups_unique)

    add_subsection("Grouped Train/Test Split")
    add_text(f"Grouping variable: {GROUP_COL}")
    add_text(f"Test size: {TEST_SIZE}")
    add_text(f"Training observations: {len(X_train)}")
    add_text(f"Test observations: {len(X_test)}")
    add_text(f"Training groups: {len(train_groups_unique)}")
    add_text(f"Test groups: {len(test_groups_unique)}")
    add_text(f"Overlapping groups: {len(overlap_groups)}")

    if len(overlap_groups) == 0:
        add_text("OK: no group appears simultaneously in training and testing.")
    else:
        add_text("WARNING: overlapping groups exist between training and testing.")
        add_list(list(overlap_groups)[:30])

    # =========================
    # 8) PREPROCESSOR
    # =========================

    preprocessor = make_preprocessor(
        numeric_features=numeric_features,
        categorical_features=categorical_features,
        use_imputation=use_imputation
    )

    # =========================
    # 9) RANDOM FOREST PIPELINE
    # =========================

    rf_base = RandomForestRegressor(
        random_state=RANDOM_STATE,
        n_jobs=-1,
        bootstrap=True,
        oob_score=True
    )

    rf_pipeline = Pipeline([
        ("preprocessor", preprocessor),
        ("model", rf_base)
    ])

    # =========================
    # 10) GROUPED CROSS-VALIDATION
    # =========================

    cv = GroupKFold(n_splits=CV_FOLDS)

    param_distributions = {
        "model__n_estimators": [200, 300, 500],
        "model__max_depth": [10, 20, 30, None],
        "model__min_samples_split": [2, 5, 10],
        "model__min_samples_leaf": [1, 2, 4, 8],
        "model__max_features": ["sqrt", "log2", 0.5, 0.7, 1.0]
    }

    search = RandomizedSearchCV(
        estimator=rf_pipeline,
        param_distributions=param_distributions,
        n_iter=N_ITER_SEARCH,
        scoring="neg_root_mean_squared_error",
        cv=cv,
        random_state=RANDOM_STATE,
        n_jobs=1,
        verbose=1,
        return_train_score=True
    )

    add_section(f"Random Forest Training - {scenario_name}")
    add_text("Model: Random Forest Regressor")
    add_text("Validation: GroupKFold")
    add_text(f"Group column: {GROUP_COL}")
    add_text(f"CV folds: {CV_FOLDS}")
    add_text(f"RandomizedSearchCV iterations: {N_ITER_SEARCH}")
    add_text("Scoring metric: negative root mean squared error")

    print("\n" + "=" * 70)
    print(f"Scenario          : {scenario_name}")
    print(f"Target            : {MODEL_TARGET_LABEL} ({MODEL_TARGET})")
    print(f"Unit              : {MODEL_TARGET_UNIT}")
    print(f"Use imputation    : {use_imputation}")
    print(f"Minimum PLZ obs   : {min_plz_obs}")
    print(f"Rows used         : {rows_final}")
    print(f"Groups used       : {n_groups}")
    print(f"Selected features : {len(selected_features)}")
    print("=" * 70)

    search.fit(X_train, y_train, groups=groups_train)

    best_model = search.best_estimator_
    best_params = search.best_params_
    best_cv_rmse = float(-search.best_score_)

    add_text("Training completed.")
    add_text(f"Best grouped CV RMSE: {best_cv_rmse:.4f}")

    add_text("Best parameters:")
    for k, v in best_params.items():
        add_text(f"{k}: {v}")

    # =========================
    # 11) FINAL PREDICTIONS
    # =========================

    y_pred_train = best_model.predict(X_train)
    y_pred_test = best_model.predict(X_test)

    train_metrics = regression_metrics(y_train, y_pred_train, "Train")
    test_metrics = regression_metrics(y_test, y_pred_test, "Grouped_Test")

    train_rmse = train_metrics["Train_RMSE"]
    train_mae = train_metrics["Train_MAE"]
    train_r2 = train_metrics["Train_R2"]

    test_rmse = test_metrics["Grouped_Test_RMSE"]
    test_mae = test_metrics["Grouped_Test_MAE"]
    test_r2 = test_metrics["Grouped_Test_R2"]

    rmse_gap = test_rmse - train_rmse
    r2_gap = train_r2 - test_r2

    # =========================
    # 11.1) OOB R2
    # =========================

    rf_model = best_model.named_steps["model"]

    try:
        oob_r2 = float(rf_model.oob_score_)
    except Exception:
        oob_r2 = np.nan

    add_section(f"Random Forest Metrics - {scenario_name}")

    add_subsection("Train Performance")
    add_text(f"RMSE: {train_rmse:.4f}")
    add_text(f"MAE: {train_mae:.4f}")
    add_text(f"R2: {train_r2:.4f}")

    add_subsection("Grouped Test Performance")
    add_text(f"RMSE: {test_rmse:.4f}")
    add_text(f"MAE: {test_mae:.4f}")
    add_text(f"R2: {test_r2:.4f}")

    add_subsection("Cross-Validation Performance")
    add_text(f"Best grouped CV RMSE: {best_cv_rmse:.4f}")

    add_subsection("Out-of-Bag Performance")
    add_text(f"OOB R2: {oob_r2:.4f}" if not pd.isna(oob_r2) else "OOB R2: NA")

    add_subsection("Overfitting Indicators")
    add_text(f"RMSE gap Test - Train: {rmse_gap:.4f}")
    add_text(f"R2 gap Train - Test: {r2_gap:.4f}")

    # =========================
    # 11.2) R2 COMPARISON PLOT
    # =========================

    add_subsection("Comparison of Train, Grouped Test and OOB R2")

    r2_comparison_df = pd.DataFrame({
        "Validation": ["Train", "Grouped test", "OOB"],
        "R2": [
            train_r2,
            test_r2,
            oob_r2
        ]
    })

    plt.figure(figsize=(7, 5))
    sns.barplot(data=r2_comparison_df, x="Validation", y="R2", color="steelblue")
    plt.ylim(0, 1)
    plt.title(f"Random Forest R2 Comparison - {scenario_name}")
    plt.ylabel("R2")
    plt.xlabel("Validation dataset")
    plt.tight_layout()
    add_current_figure_to_doc(width=6)
    plt.close()

    # =========================
    # 12) FEATURE NAMES AFTER PREPROCESSING
    # =========================

    fitted_preprocessor = best_model.named_steps["preprocessor"]

    try:
        feature_names_out = list(fitted_preprocessor.get_feature_names_out())
        feature_names_out = [
            n.replace("num__", "").replace("cat__", "")
            for n in feature_names_out
        ]
    except Exception:
        feature_names_out = [
            f"x_{i}" for i in range(len(rf_model.feature_importances_))
        ]

    add_rf_prediction_examples_and_tree(
        scenario_name=scenario_name,
        X_test=X_test,
        y_test=y_test,
        y_pred_test=y_pred_test,
        rf_model=rf_model,
        feature_names_out=feature_names_out,
        n_examples=EXAMPLE_N_CASES
    )

    # =========================
    # 13) INTERNAL FEATURE IMPORTANCE
    # =========================

    importance_df = pd.DataFrame({
        "feature": feature_names_out,
        "importance": rf_model.feature_importances_
    })

    importance_df = importance_df.sort_values("importance", ascending=False)

    add_section(f"Internal RF Feature Importance - {scenario_name}")
    add_text("Top 20 features by internal Random Forest feature importance:")

    for _, row in importance_df.head(20).iterrows():
        add_text(f"{row['feature']}: {row['importance']:.6f}")

    top_imp_plot = importance_df.head(20).sort_values("importance", ascending=True)

    fig, ax = plt.subplots(figsize=(8.2, 5.6))
    ax.barh(top_imp_plot["feature"], top_imp_plot["importance"], color="saddlebrown")
    ax.set_xlabel("Feature importance", fontsize=10)
    ax.set_ylabel("Feature", fontsize=10)
    ax.set_title(f"Top 20 RF Feature Importances - {scenario_name}", fontsize=10)
    ax.tick_params(axis="both", labelsize=10)
    plt.tight_layout()
    add_figure_to_doc(fig, width=6.2)
    plt.close(fig)

    # =========================
    # 14) PERMUTATION IMPORTANCE
    # =========================

    add_section(f"Permutation Importance - {scenario_name}")

    if len(X_test) > PERMUTATION_SAMPLE_SIZE:
        X_perm = X_test.sample(PERMUTATION_SAMPLE_SIZE, random_state=RANDOM_STATE)
        y_perm = y_test.loc[X_perm.index]
    else:
        X_perm = X_test.copy()
        y_perm = y_test.copy()

    add_text(f"Permutation importance sample size: {len(X_perm)}")
    add_text(f"N repeats: {PERMUTATION_N_REPEATS}")

    perm_result = permutation_importance(
        best_model,
        X_perm,
        y_perm,
        n_repeats=PERMUTATION_N_REPEATS,
        random_state=RANDOM_STATE,
        scoring="neg_root_mean_squared_error",
        n_jobs=1,
    )

    perm_importance_df = pd.DataFrame({
        "feature": X_perm.columns,
        "importance_mean": perm_result.importances_mean,
        "importance_std": perm_result.importances_std
    }).sort_values("importance_mean", ascending=False)

    add_text("Top 20 permutation importances by original feature:")

    for _, row in perm_importance_df.head(20).iterrows():
        add_text(
            f"{row['feature']}: mean={row['importance_mean']:.6f}, "
            f"std={row['importance_std']:.6f}"
        )

    top_perm_plot = perm_importance_df.head(20).sort_values("importance_mean", ascending=True)

    fig, ax = plt.subplots(figsize=(8.2, 5.6))
    ax.barh(top_perm_plot["feature"], top_perm_plot["importance_mean"], color="forestgreen")
    ax.set_xlabel("Permutation importance", fontsize=10)
    ax.set_ylabel("Original feature", fontsize=10)
    ax.set_title(f"Top 20 Permutation Importances - {scenario_name}", fontsize=10)
    ax.tick_params(axis="both", labelsize=10)
    plt.tight_layout()
    add_figure_to_doc(fig, width=6.2)
    plt.close(fig)

    # =========================
    # 15) PARTIAL DEPENDENCE PLOTS
    # =========================

    if GENERATE_PDP:
        add_section(f"Partial Dependence Plots - {scenario_name}")

        if MODEL_TARGET == TARGET:
            pdp_features = [
                "obj_livingSpace",
                "obj_picturecount",
                "obj_lift",
                "obj_balcony",
                "obj_hasKitchen"
            ]
        else:
            pdp_features = [
                "obj_buildingAge",
                "obj_picturecount",
                "obj_lift",
                "obj_balcony",
                "obj_hasKitchen"
            ]

        available_pdp_features = [f for f in pdp_features if f in X_test.columns]

        for feature in available_pdp_features:
            try:
                add_subsection(f"PDP: {feature}")

                fig, ax = plt.subplots(figsize=(7, 5))

                PartialDependenceDisplay.from_estimator(
                    best_model,
                    X_test,
                    features=[feature],
                    ax=ax,
                    grid_resolution=50,
                    line_kw={"color": "black", "linewidth": 1.5}
                )

                # Standardize PDP style for reporting readability.
                for line in ax.get_lines():
                    line.set_color("black")
                    line.set_linewidth(1.5)

                ax.tick_params(axis="both", labelsize=10)
                ax.xaxis.label.set_size(10)
                ax.yaxis.label.set_size(10)
                

                ax.set_title(
                    f"Partial Dependence Plot - {feature} - {scenario_name}",
                    fontsize=10
                )
                plt.tight_layout()
                add_figure_to_doc(fig, width=6)
                plt.close(fig)

            except Exception as exc:
                add_text(f"PDP for {feature} could not be generated: {exc}")

        if MODEL_TARGET == TARGET:
            try:
                if "obj_livingSpace" in X_test.columns and "obj_picturecount" in X_test.columns:
                    add_subsection("2D PDP: obj_livingSpace x obj_picturecount")

                    fig, ax = plt.subplots(figsize=(7, 5))

                    PartialDependenceDisplay.from_estimator(
                        best_model,
                        X_test,
                        features=[("obj_livingSpace", "obj_picturecount")],
                        ax=ax,
                        grid_resolution=30,
                        line_kw={"color": "black", "linewidth": 1.5}
                    )

                    ax.tick_params(axis="both", labelsize=10)
                    ax.xaxis.label.set_size(10)
                    ax.yaxis.label.set_size(10)

                    ax.set_title(
                        f"2D Partial Dependence Plot - living space x picture count - {scenario_name}",
                        fontsize=10
                    )
                    plt.tight_layout()
                    add_figure_to_doc(fig, width=6)
                    plt.close(fig)

            except Exception as exc:
                add_text(f"2D PDP could not be generated: {exc}")

    # =========================
    # 16) DIAGNOSTIC PLOTS
    # =========================

    add_section(f"Diagnostic Plots - {scenario_name}")
    add_text(
        "Status labels (OK/CHECK/NOT OK) in RF diagnostics are heuristic indicators "
        "for readability and are not formal statistical tests."
    )

    add_subsection("Actual vs Predicted")

    fig, ax = plt.subplots(figsize=(7, 7))
    ax.scatter(
        y_test,
        y_pred_test,
        alpha=0.25,
        color="black"
    )

    min_val = min(np.min(y_test), np.min(y_pred_test))
    max_val = max(np.max(y_test), np.max(y_pred_test))

    ax.plot(
        [min_val, max_val],
        [min_val, max_val],
        color="red",
        linestyle="--",
        label="Perfect prediction"
    )

    ax.set_xlabel(f"Actual {MODEL_TARGET_LABEL}")
    ax.set_ylabel(f"Predicted {MODEL_TARGET_LABEL}")
    ax.set_title(f"Random Forest: Actual vs Predicted - {scenario_name}")
    ax.legend()
    plt.tight_layout()
    add_figure_to_doc(fig, width=6)
    plt.close(fig)

    residuals = y_test - y_pred_test
    abs_errors = np.abs(residuals)
    eps = 1e-9
    y_test_arr = np.asarray(y_test.values, dtype=float)
    y_pred_arr = np.asarray(y_pred_test, dtype=float)
    residuals_arr = np.asarray(residuals.values, dtype=float)

    denom_actual = np.where(np.abs(y_test_arr) > eps, np.abs(y_test_arr), np.nan)
    pct_errors = (residuals_arr / denom_actual) * 100.0
    abs_pct_errors = np.abs(pct_errors)
    smape_pct_values = 200.0 * np.abs(residuals_arr) / (np.abs(y_test_arr) + np.abs(y_pred_arr) + eps)

    mape_pct = float(np.nanmean(abs_pct_errors)) if not pd.isna(np.nanmean(abs_pct_errors)) else np.nan
    mdape_pct = float(np.nanmedian(abs_pct_errors)) if not pd.isna(np.nanmedian(abs_pct_errors)) else np.nan
    smape_pct = float(np.nanmean(smape_pct_values)) if not pd.isna(np.nanmean(smape_pct_values)) else np.nan

    add_subsection("Residuals vs Predicted")

    fig, ax = plt.subplots(figsize=(8, 5))
    ax.scatter(
        y_pred_arr,
        residuals_arr,
        alpha=0.25,
        color="black"
    )
    ax.axhline(y=0, color="red", linestyle="--")

    ax.set_xlabel(f"Predicted {MODEL_TARGET_LABEL}")
    ax.set_ylabel("Residuals")
    ax.set_title(f"Random Forest: Residuals vs Predicted - {scenario_name}")
    plt.tight_layout()
    add_figure_to_doc(fig, width=6)
    plt.close(fig)

    add_subsection("Residual Distribution")

    fig, ax = plt.subplots(figsize=(8, 5))
    sns.histplot(residuals, kde=True, color="steelblue", ax=ax)
    ax.set_xlabel("Residuals")
    ax.set_title(f"Random Forest: Residual Distribution - {scenario_name}")
    format_axis_pretty(ax)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    add_figure_to_doc(fig, width=6)
    plt.close(fig)

    add_subsection("Absolute Error Distribution")

    fig, ax = plt.subplots(figsize=(8, 5))
    sns.histplot(abs_errors, kde=True, color="steelblue", ax=ax)
    ax.set_xlabel("Absolute error")
    ax.set_title(f"Random Forest: Absolute Error Distribution - {scenario_name}")
    format_axis_pretty(ax)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    add_figure_to_doc(fig, width=6)
    plt.close(fig)

    add_subsection("Absolute Percentage Error Distribution")

    fig, ax = plt.subplots(figsize=(8, 5))
    sns.histplot(pd.Series(abs_pct_errors).dropna(), kde=True, color="steelblue", ax=ax)

    mape_label, mape_color = pass_fail_label_from_threshold(
        mape_pct,
        ok_threshold=20.0,
        check_threshold=30.0,
        lower_is_better=True
    )
    add_status_box(
        ax,
        "MAPE heuristic",
        f"MAPE = {mape_pct:.2f}%" if not pd.isna(mape_pct) else "MAPE = NA",
        mape_label,
        mape_color
    )

    ax.set_xlabel("Absolute percentage error (%)")
    ax.set_title(f"Random Forest: Absolute Percentage Error Distribution - {scenario_name}")
    format_axis_pretty(ax)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    add_figure_to_doc(fig, width=6)
    plt.close(fig)

    # =========================
    # 17) ERROR SUMMARY
    # =========================

    error_summary = pd.Series(abs_errors).describe(
        percentiles=[0.25, 0.5, 0.75, 0.90, 0.95, 0.99]
    )

    add_section(f"Error Summary - {scenario_name}")
    add_text("Absolute error summary:")

    for idx, val in error_summary.items():
        add_text(f"{idx}: {val:.4f}")

    add_text(f"MAPE (%): {mape_pct:.4f}" if not pd.isna(mape_pct) else "MAPE (%): NA")
    add_text(f"MdAPE (%): {mdape_pct:.4f}" if not pd.isna(mdape_pct) else "MdAPE (%): NA")
    add_text(f"sMAPE (%): {smape_pct:.4f}" if not pd.isna(smape_pct) else "sMAPE (%): NA")

    error_summary_df = (
        error_summary
        .reset_index()
        .rename(columns={"index": "statistic", 0: "value"})
    )

    # =========================
    # 18) GROUP-LEVEL ERROR
    # =========================

    predictions_df = X_test.copy()
    predictions_df[GROUP_COL] = groups_test.values
    predictions_df["actual"] = y_test.values
    predictions_df["predicted"] = y_pred_test
    predictions_df["residual"] = residuals.values
    predictions_df["absolute_error"] = abs_errors.values
    predictions_df["pct_error"] = pct_errors
    predictions_df["abs_pct_error"] = abs_pct_errors
    predictions_df["smape_pct"] = smape_pct_values

    residuals_detail_df, residuals_by_zip_df = build_residual_zip_tables(
        df_source=df_model,
        index=y_test.index,
        y_true=y_test.values,
        y_pred=y_pred_test,
        zip_col=ZIP_COL_FOR_MAP,
        scenario_name=scenario_name
    )

    group_error_df = predictions_df.groupby(GROUP_COL).agg(
        n_obs=("absolute_error", "size"),
        mean_actual=("actual", "mean"),
        mean_predicted=("predicted", "mean"),
        mean_residual=("residual", "mean"),
        mean_absolute_error=("absolute_error", "mean"),
        median_absolute_error=("absolute_error", "median"),
        mean_pct_error=("pct_error", "mean"),
        mean_abs_pct_error=("abs_pct_error", "mean"),
        smape_pct=("smape_pct", "mean")
    ).reset_index()

    group_error_df = group_error_df.sort_values(
        "mean_absolute_error",
        ascending=False
    )

    group_error_reliable_df = group_error_df[
        group_error_df["n_obs"] >= MIN_GROUP_ERROR_N
    ].copy()

    group_error_reliable_df = group_error_reliable_df.sort_values(
        "mean_absolute_error",
        ascending=False
    )

    add_section(f"Group-Level Error Summary - {scenario_name}")

    add_text(f"Top 20 {GROUP_COL} groups with highest mean absolute error - all groups:")

    for _, row in group_error_df.head(20).iterrows():
        add_text(
            f"{row[GROUP_COL]} | "
            f"n={row['n_obs']} | "
            f"MAE={row['mean_absolute_error']:.4f} | "
            f"mean residual={row['mean_residual']:.4f} | "
            f"MAPE={row['mean_abs_pct_error']:.2f}%"
        )

    add_subsection(f"Group-Level Error Summary - n_obs >= {MIN_GROUP_ERROR_N}")

    add_text(
        f"Top 20 {GROUP_COL} groups with highest mean absolute error considering only "
        f"groups with n_obs >= {MIN_GROUP_ERROR_N}:"
    )

    for _, row in group_error_reliable_df.head(20).iterrows():
        add_text(
            f"{row[GROUP_COL]} | "
            f"n={row['n_obs']} | "
            f"MAE={row['mean_absolute_error']:.4f} | "
            f"mean residual={row['mean_residual']:.4f} | "
            f"MAPE={row['mean_abs_pct_error']:.2f}%"
        )

    # =========================
    # 19) CV RESULTS EXPORT
    # =========================

    cv_results_df = pd.DataFrame(search.cv_results_)

    cols_to_keep = [
        "mean_test_score",
        "std_test_score",
        "mean_train_score",
        "std_train_score",
        "rank_test_score",
        "params"
    ]

    cv_results_export = cv_results_df[cols_to_keep].copy()
    cv_results_export["mean_test_RMSE"] = -cv_results_export["mean_test_score"]
    cv_results_export["mean_train_RMSE"] = -cv_results_export["mean_train_score"]
    cv_results_export = cv_results_export.sort_values("rank_test_score")

    # =========================
    # 20) SCENARIO RESULT ROW
    # =========================

    result_row = {
        "scenario": scenario_name,
        "model_target": MODEL_TARGET,
        "model_target_label": MODEL_TARGET_LABEL,
        "model_target_unit": MODEL_TARGET_UNIT,
        "min_plz_obs": min_plz_obs,
        "use_imputation": use_imputation,

        "n_obs": rows_final,
        "n_train": len(y_train),
        "n_test": len(y_test),
        "n_total_groups": n_groups,
        "n_train_groups": groups_train.nunique(),
        "n_test_groups": groups_test.nunique(),
        "n_overlap_groups": len(overlap_groups),
        "n_selected_features_original": len(selected_features),

        "rows_initial": rows_initial,
        "rows_after_model_target_group_drop": rows_after_model_target_group_drop,
        "rows_after_missing_handling": rows_after_missing_handling,
        "rows_before_plz_filter": rows_before_plz_filter,
        "rows_after_plz_filter": rows_after_plz_filter,
        "groups_before_plz_filter": groups_before_plz_filter,
        "groups_after_plz_filter": groups_after_plz_filter,
        "removed_plz_groups": len(removed_plz_groups_df),

        "RF_Train_RMSE": train_rmse,
        "RF_Test_RMSE": test_rmse,
        "RF_RMSE_Gap": rmse_gap,

        "RF_Train_MAE": train_mae,
        "RF_Test_MAE": test_mae,
        "RF_Test_MAPE_pct": mape_pct,
        "RF_Test_MdAPE_pct": mdape_pct,
        "RF_Test_sMAPE_pct": smape_pct,

        "RF_Train_R2": train_r2,
        "RF_Test_R2": test_r2,
        "RF_R2_Gap": r2_gap,

        "RF_Best_Grouped_CV_RMSE": best_cv_rmse,
        "RF_OOB_R2": oob_r2,

        "RF_Best_Params": str(best_params)
    }

    metrics_df = pd.DataFrame({
        "metric": list(result_row.keys()),
        "value": list(result_row.values())
    })

    best_params_df = pd.DataFrame({
        "parameter": list(best_params.keys()),
        "value": [str(v) for v in best_params.values()]
    })

    selected_features_df = pd.DataFrame({
        "selected_feature": selected_features
    })

    # =========================
    # 21) EXPORT OBJECTS
    # =========================

    export_objects = {
        "result_row": result_row,
        "metrics_df": metrics_df,
        "best_params_df": best_params_df,
        "selected_features_df": selected_features_df,
        "pearson_target_corr_df": pearson_target_corr_df,
        "feature_missing_share_df": feature_missing_share_df,
        "removed_plz_groups_df": removed_plz_groups_df,
        "importance_df": importance_df,
        "perm_importance_df": perm_importance_df,
        "predictions_df": predictions_df,
        "group_error_df": group_error_df,
        "group_error_reliable_df": group_error_reliable_df,
        "cv_results_export": cv_results_export,
        "error_summary_df": error_summary_df,
        "residuals_detail_df": residuals_detail_df,
        "residuals_by_zip_df": residuals_by_zip_df
    }

    return export_objects


# ============================================================
# LOAD DATA
# ============================================================

READ_COLS = list(dict.fromkeys(
    BASE_FEATURES +
    RAW_ENGINEERING_FEATURES +
    [
        TARGET,
        "obj_livingSpace",
        ZIP_COL_FOR_MAP,
        GROUP_COL
    ]
))

try:
    df = pd.read_csv(
        FILE_PATH,
        usecols=READ_COLS,
        sep=CSV_SEP,
        low_memory=False
    )

    df["obj_totalRent_num"] = pd.to_numeric(df[TARGET], errors="coerce")
    df["obj_livingSpace_num"] = pd.to_numeric(df["obj_livingSpace"], errors="coerce")

    df[TARGET_RENT_SQM] = (
        df["obj_totalRent_num"] /
        df["obj_livingSpace_num"]
    )

    df.loc[
        (df["obj_livingSpace_num"] <= 0) |
        (~np.isfinite(df[TARGET_RENT_SQM])),
        TARGET_RENT_SQM
    ] = np.nan

except ValueError as e:
    print("\nERROR while loading columns.")
    print("Probably GROUP_COL or one selected feature does not exist in the CSV.")
    print(f"Current GROUP_COL: {GROUP_COL}")
    print("\nOriginal error:")
    print(e)

    print("\nAvailable columns in the CSV:")
    df_cols = pd.read_csv(FILE_PATH, sep=CSV_SEP, nrows=0).columns.tolist()
    for col in df_cols:
        print(col)

    raise


# ============================================================
# BUILD RANDOM FOREST SCENARIOS
# ============================================================

SCENARIOS = []

if DEBUG_MODE:

    selected_scenario = next(
        (
            x for x in PLZ_OBS_SCENARIOS
            if x["scenario_code"] == DEBUG_SCENARIO_CODE
        ),
        None
    )

    if selected_scenario is None:
        raise ValueError(
            f"DEBUG_SCENARIO_CODE='{DEBUG_SCENARIO_CODE}' is not valid. "
            "Use one of: N0, N3, N5, N10."
        )

    scenario_label = (
        "WITH IMPUTATION"
        if DEBUG_USE_IMPUTATION
        else "NO IMPUTATION"
    )

    SCENARIOS.append({
        "scenario_code": selected_scenario["scenario_code"],
        "min_plz_obs": selected_scenario["min_plz_obs"],
        "name": f"{selected_scenario['scenario_code']} - {scenario_label}",
        "use_imputation": DEBUG_USE_IMPUTATION
    })

else:

    for plz_scenario in PLZ_OBS_SCENARIOS:

        for use_imputation in IMPUTATION_SCENARIOS:

            scenario_label = (
                "WITH IMPUTATION"
                if use_imputation
                else "NO IMPUTATION"
            )

            SCENARIOS.append({
                "scenario_code": plz_scenario["scenario_code"],
                "min_plz_obs": plz_scenario["min_plz_obs"],
                "name": f"{plz_scenario['scenario_code']} - {scenario_label}",
                "use_imputation": use_imputation
            })

print(f"\nTotal Random Forest scenarios configured: {len(SCENARIOS)}")


# ============================================================
# RUN RANDOM FOREST SCENARIOS
# ============================================================

all_results = []
all_exports = {}

for scenario in SCENARIOS:

    scenario_name = scenario["name"]
    use_imputation = scenario["use_imputation"]
    min_plz_obs = scenario["min_plz_obs"]

    print(f"\nRunning Random Forest scenario: {scenario_name}")

    scenario_export = run_rf_scenario(
        df_original=df,
        scenario_name=scenario_name,
        use_imputation=use_imputation,
        min_plz_obs=min_plz_obs
    )

    if scenario_export is not None:
        all_exports[scenario_name] = scenario_export
        all_results.append(scenario_export["result_row"])


# ============================================================
# FINAL COMPARISON
# ============================================================

comparison_df = pd.DataFrame(all_results)

add_section(f"FINAL COMPARISON - RANDOM FOREST - {MODEL_TARGET_LABEL}")

if DEBUG_MODE:
    add_text("Debug mode was enabled. Only one scenario was executed.")
else:
    add_text("Full mode was enabled. All PLZ observation and imputation scenarios were executed.")

if comparison_df.empty:
    add_text("No Random Forest scenario returned valid results.")
else:
    add_text("Comparison table across Random Forest scenarios:")
    add_text(comparison_df.to_string(index=False))

    # RMSE / MAE / CV RMSE plot
    add_subsection("Comparison Plot - RMSE and MAE")

    error_metrics_df = comparison_df.set_index("scenario")[
        [
            "RF_Test_RMSE",
            "RF_Test_MAE",
            "RF_Best_Grouped_CV_RMSE"
        ]
    ]

    fig, ax = plt.subplots(figsize=(12, 6))
    error_metrics_df.plot(kind="bar", ax=ax)

    ax.set_title(f"Random Forest Error Metrics - {MODEL_TARGET_LABEL}")
    ax.set_ylabel(f"Metric value ({MODEL_TARGET_UNIT})")
    ax.set_xlabel("Scenario")
    ax.tick_params(axis="x", rotation=45)
    plt.tight_layout()

    add_figure_to_doc(fig, width=6.5)
    plt.close(fig)

    # R2 plot
    add_subsection("Comparison Plot - R2")

    r2_metrics_df = comparison_df.set_index("scenario")[
        [
            "RF_Train_R2",
            "RF_Test_R2",
            "RF_OOB_R2"
        ]
    ]

    fig, ax = plt.subplots(figsize=(12, 6))
    r2_metrics_df.plot(kind="bar", ax=ax)

    ax.set_title(f"Random Forest R2 - {MODEL_TARGET_LABEL}")
    ax.set_ylabel("R2")
    ax.set_xlabel("Scenario")
    ax.tick_params(axis="x", rotation=45)
    plt.tight_layout()

    add_figure_to_doc(fig, width=6.5)
    plt.close(fig)

    # Overfitting plot
    add_subsection("Comparison Plot - Overfitting Gaps")

    gap_metrics_df = comparison_df.set_index("scenario")[
        [
            "RF_RMSE_Gap",
            "RF_R2_Gap"
        ]
    ]

    fig, ax = plt.subplots(figsize=(12, 6))
    gap_metrics_df.plot(kind="bar", ax=ax)

    ax.set_title(f"Random Forest Overfitting Indicators - {MODEL_TARGET_LABEL}")
    ax.set_ylabel("Gap value")
    ax.set_xlabel("Scenario")
    ax.tick_params(axis="x", rotation=45)
    plt.tight_layout()

    add_figure_to_doc(fig, width=6.5)
    plt.close(fig)


# ============================================================
# SAVE WORD
# ============================================================

doc.save(DOCX_OUTPUT)


# ============================================================
# SAVE EXCEL
# ============================================================

with pd.ExcelWriter(OUTPUT_EXCEL, engine="openpyxl") as writer:

    if not comparison_df.empty:
        comparison_df.to_excel(
            writer,
            index=False,
            sheet_name="comparison"
        )

    for scenario_name, export_data in all_exports.items():

        clean_name = clean_sheet_name(scenario_name)

        export_data["metrics_df"].to_excel(
            writer,
            index=False,
            sheet_name=clean_sheet_name(f"{clean_name}_metrics")
        )

        export_data["best_params_df"].to_excel(
            writer,
            index=False,
            sheet_name=clean_sheet_name(f"{clean_name}_params")
        )

        export_data["selected_features_df"].to_excel(
            writer,
            index=False,
            sheet_name=clean_sheet_name(f"{clean_name}_selected")
        )

        export_data["pearson_target_corr_df"].to_excel(
            writer,
            index=False,
            sheet_name=clean_sheet_name(f"{clean_name}_pearson")
        )

        export_data["feature_missing_share_df"].to_excel(
            writer,
            index=False,
            sheet_name=clean_sheet_name(f"{clean_name}_missing")
        )

        export_data["removed_plz_groups_df"].to_excel(
            writer,
            index=False,
            sheet_name=clean_sheet_name(f"{clean_name}_removed_plz")
        )

        export_data["importance_df"].to_excel(
            writer,
            index=False,
            sheet_name=clean_sheet_name(f"{clean_name}_internal")
        )

        export_data["perm_importance_df"].to_excel(
            writer,
            index=False,
            sheet_name=clean_sheet_name(f"{clean_name}_perm")
        )

        export_data["group_error_df"].to_excel(
            writer,
            index=False,
            sheet_name=clean_sheet_name(f"{clean_name}_grp_all")
        )

        export_data["group_error_reliable_df"].to_excel(
            writer,
            index=False,
            sheet_name=clean_sheet_name(f"{clean_name}_grp_n3")
        )

        export_data["cv_results_export"].to_excel(
            writer,
            index=False,
            sheet_name=clean_sheet_name(f"{clean_name}_cv")
        )

        export_data["error_summary_df"].to_excel(
            writer,
            index=False,
            sheet_name=clean_sheet_name(f"{clean_name}_error")
        )

        export_data["residuals_detail_df"].to_excel(
            writer,
            index=False,
            sheet_name=clean_sheet_name(f"{clean_name}_residuals")
        )

        export_data["residuals_by_zip_df"].to_excel(
            writer,
            index=False,
            sheet_name=clean_sheet_name(f"{clean_name}_zip_resid")
        )

        # This sheet can be large.
        # If Excel becomes too heavy, comment this block.
        export_data["predictions_df"].to_excel(
            writer,
            index=False,
            sheet_name=clean_sheet_name(f"{clean_name}_pred")
        )


if ENABLE_GERMANY_RESIDUAL_MAP:
    add_section("GERMANY RESIDUAL MAPS")
    for scenario_name, export_data in all_exports.items():
        map_path = plot_germany_residual_map(
            zip_summary_df=export_data["residuals_by_zip_df"],
            scenario_name=scenario_name,
            output_prefix="rf_residual_map_de",
            min_listings_per_zip=RESIDUAL_MAP_MIN_LISTINGS_PER_ZIP,
            max_points=RESIDUAL_MAP_MAX_POINTS,
            color_column=RESIDUAL_MAP_COLOR_COLUMN,
            cmap=RESIDUAL_MAP_CMAP,
            color_label=RESIDUAL_MAP_COLOR_LABEL,
            symmetric_scale=RESIDUAL_MAP_SYMMETRIC_SCALE,
            clip_percentile_low=RESIDUAL_MAP_CLIP_PCT_LOW,
            clip_percentile_high=RESIDUAL_MAP_CLIP_PCT_HIGH,
        )

        if map_path is None:
            add_text(
                f"Map skipped for '{scenario_name}'. "
                "Possible reasons: no ZIP data after filtering, or missing optional package 'pgeocode'."
            )
            continue

        add_subsection(f"Residual map - {scenario_name}")
        add_text(
            f"Color metric: {RESIDUAL_MAP_COLOR_COLUMN}. "
            "Marker size scales with the number of listings in the ZIP area."
        )
        doc.add_picture(map_path, width=Inches(6.0))


# ============================================================
# PRINT FINAL RESULTS
# ============================================================

print(f"\nWord report saved: {DOCX_OUTPUT}")
print(f"Excel report saved: {OUTPUT_EXCEL}")

print("\nFinal Random Forest comparison:")

if comparison_df.empty:
    print("No valid Random Forest results were produced.")
else:
    print(comparison_df.to_string(index=False))

elapsed_seconds = perf_counter() - SCRIPT_START_TIME
elapsed_minutes = elapsed_seconds / 60.0
print(f"\nTotal runtime: {elapsed_seconds:.2f} seconds ({elapsed_minutes:.2f} minutes)")