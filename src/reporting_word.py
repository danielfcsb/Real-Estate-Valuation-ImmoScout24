from io import BytesIO
from docx import Document
from docx.shared import Inches


class WordReporter:
    def __init__(self, title: str):
        self.doc = Document()
        self.doc.add_heading(str(title), 0)

    def add_section(self, title):
        self.doc.add_heading(str(title), level=1)

    def add_subsection(self, title):
        self.doc.add_heading(str(title), level=2)

    def add_text(self, text):
        self.doc.add_paragraph(str(text))

    def add_list(self, items):
        if not items:
            self.doc.add_paragraph("(no items)")
            return
        for item in items:
            self.doc.add_paragraph(str(item), style="List Bullet")

    def add_figure(self, fig, width=6.5):
        buf = BytesIO()
        fig.savefig(buf, format="png", dpi=300, bbox_inches="tight")
        buf.seek(0)
        self.doc.add_picture(buf, width=Inches(width))
        buf.close()

    def save(self, output_path: str):
        self.doc.save(output_path)