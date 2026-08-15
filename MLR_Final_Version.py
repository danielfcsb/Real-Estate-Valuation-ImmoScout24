
# ============================================================
# MLR MODEL SCRIPT (OLS + RIDGE + STEPWISE)
# ------------------------------------------------------------
# Main outputs:
# 1) Performance metrics (RMSE, MAE, R2)
# 2) Assumption diagnostics (BP, JB, Shapiro, DW, RESET, Rainbow)
# 3) Feature interpretation tables (OLS, Ridge, VIF, Stepwise)
# 4) Optional Germany residual map by ZIP code
#
# Scenario logic:
# - PLZ observation filters (N0, N3, N4, N5, N10)
# - No-imputation vs imputation setup
#
# Report outputs:
# - Word (.docx)
# - Excel (.xlsx)
# ============================================================
# SCRIPT NAVIGATION (READ FIRST)
# 1) CONFIGURATION USER INPUTS
#    Edit scenario, target, grouping, and map controls.
# 2) FEATURE/ENGINEERING DEFINITIONS
#    Central place for explanatory variables and derived variables.
# 3) HELPER FUNCTIONS
#    Utilities for report writing, diagnostics formatting, preprocessing.
# 4) MODELING PIPELINE
#    OLS, Ridge, Stepwise training/evaluation per scenario.
# 5) RESIDUAL MAP + EXPORTS
#    ZIP residual aggregation, Germany map rendering, Excel/Word export.
# 6) FINAL COMPARISON OUTPUT
#    Consolidated metrics for scenario-level decision making.

import numpy as np
import pandas as pd
import warnings
from io import BytesIO
from datetime import datetime
from time import perf_counter
import json
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.ticker import FuncFormatter, MaxNLocator
import seaborn as sns
import scipy.stats as stats

import statsmodels.api as sm
from statsmodels.stats.outliers_influence import variance_inflation_factor
from statsmodels.stats.diagnostic import (
    het_breuschpagan,
    linear_reset,
    linear_rainbow
)
from statsmodels.stats.stattools import durbin_watson
from statsmodels.graphics.tsaplots import plot_acf
from statsmodels.nonparametric.smoothers_lowess import lowess

from sklearn.model_selection import train_test_split
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.impute import SimpleImputer
from sklearn.linear_model import RidgeCV, Ridge
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
from sklearn.model_selection import GroupShuffleSplit

from docx import Document
from docx.shared import Inches


# ============================================================
# CONFIGURATION USER INPUTS
# ============================================================

# ------------------------------
# USER GUIDE
# ------------------------------
# In most experiments, you only need to edit:
# 1) MODEL_TARGET              -> choose total rent or rent_per_sqm
# 2) GROUP_COL                 -> grouped split unit (ZIP recommended)
# 3) DEBUG_MODE settings       -> run one scenario or all scenarios
# 4) ENABLE_GERMANY_RESIDUAL_MAP and residual-map settings
#
# Safety note:
# - Keep GROUP_COL = "obj_zipCode" for spatially meaningful train/test separation.
# - Switching to "obj_regio1" increases aggregation and can reduce map coverage.

FILE_PATH = r"C:\Users\borgesd8828\OneDrive - ARCADIS\Dokumente\GitHub\Real-Estate-Valuation-ImmoScout24\apr20_rental_no_duplicates_for_python.csv"
CSV_SEP = ","

# ============================================================
# TARGET CONFIGURATION
# ============================================================

TARGET = "obj_totalRent"
TARGET_RENT_SQM = "rent_per_sqm"

# Choose the target used by all models and reports.
# Option A (recommended for macro valuation): total monthly rent in EUR.
# Option B (recommended for comparability across size): rent per sqm in EUR/sqm.
MODEL_TARGET = TARGET           # total monthly rent
# MODEL_TARGET = TARGET_RENT_SQM  # rent per sqm

# Grouping variable for GroupShuffleSplit.
# Recommended: obj_zipCode (finer spatial generalization check).
# Alternative: obj_regio1 (broader regions, fewer groups, less spatial detail).
GROUP_COL = "obj_zipCode"
# GROUP_COL = "obj_regio1"

TARGET_LABELS = {
    "obj_totalRent": "Total rent",
    "rent_per_sqm": "Rent per sqm"
}

TARGET_UNITS = {
    "obj_totalRent": "EUR/month",
    "rent_per_sqm": "EUR/sqm"
}

MODEL_TARGET_LABEL = TARGET_LABELS.get(MODEL_TARGET, MODEL_TARGET)
MODEL_TARGET_UNIT = TARGET_UNITS.get(MODEL_TARGET, "")

BASE_FEATURES_FULL = [
    # "obj_yearConstructed",
    "obj_heatingType",
    # "obj_serviceCharge",
    "obj_ExclusiveExpose",
    # "obj_newlyConst",
    "obj_regio1",
    "obj_balcony",
    "obj_cellar",
    "obj_hasKitchen",
    "obj_picturecount",
    "obj_lift",
    "obj_petsAllowed",
    # "obj_interiorQual",
    "obj_condition",
    "obj_livingSpace",
    # "obj_depositLink",
    # "ga_cd_via",
    "obj_typeOfFlat",
    "obj_garden",
    "obj_barrierFree",
    # "obj_zipCode"  # replaced by obj_zip3 (3-digit regional prefix)
]

if MODEL_TARGET == TARGET_RENT_SQM:
    BASE_FEATURES = [
        feature for feature in BASE_FEATURES_FULL
        if feature != "obj_livingSpace"
    ]
else:
    BASE_FEATURES = BASE_FEATURES_FULL.copy()


RAW_ENGINEERING_FEATURES = [
    "obj_energyType",
   # "obj_thermalChar",
    "obj_numberOfFloors",
    "obj_noParkSpaces",
    "obj_lastRefurbish",
    "obj_yearConstructed"
]

ENGINEERED_FEATURES = [
    "obj_energyType_cat",

  #  "obj_hasThermalCharInfo",
  #  "obj_thermalChar_num",

    "obj_hasNumberOfFloorsInfo",
    "obj_numberOfFloors_num",

    "obj_hasParkingInfo",
    "obj_noParkSpaces_num",

    "obj_hasLastRefurbishInfo",
    "obj_yearsSinceLastRefurbish",

   # "obj_hasYearConstructedInfo",
    "obj_buildingAge",

   # "obj_refurbishedAfterConstruction",
    "obj_yearsBetweenConstructionAndRefurbish",
  #  "obj_hasConstructionAndRefurbishInfo",
    "obj_zip3"
]

FEATURES = BASE_FEATURES + ENGINEERED_FEATURES

BINARY_FEATURES = [
    "obj_ExclusiveExpose",
    # "obj_newlyConst",
    "obj_balcony",
    "obj_cellar",
    "obj_hasKitchen",
    "obj_lift",
    "obj_garden",
    "obj_barrierFree",
    # "obj_depositLink",

   # "obj_hasThermalCharInfo",
    "obj_hasNumberOfFloorsInfo",
    "obj_hasParkingInfo",
    "obj_hasLastRefurbishInfo",
  #  "obj_hasYearConstructedInfo",
   # "obj_refurbishedAfterConstruction",
   # "obj_hasConstructionAndRefurbishInfo"
]

RANDOM_STATE = 42
# Test share used in grouped split. Higher values increase test size and may reduce train stability.
TEST_SIZE = 0.25
# VIF can be expensive in very large matrices; this caps rows sampled for VIF calculation.
VIF_SAMPLE_SIZE = 5000
VIF_REMOVE_THRESHOLD = 11
REFERENCE_YEAR = 2020

RIDGE_ALPHAS = np.logspace(-3, 3, 50)



STEPWISE_P_THRESHOLD = 0.05
STEPWISE_MAX_ITER = 100
SIGNIFICANCE_LEVEL = 0.05

# Global outlier clipping for model inputs.
# This is independent from map color clipping and affects model data directly.
ENABLE_GLOBAL_OUTLIER_CLIPPING = True
# Outlier handling mode:
# - "clip_values": keep all rows and cap extreme numeric values at percentile limits.
# - "remove_rows": remove rows containing outliers in monitored numeric variables.
OUTLIER_TREATMENT_MODE = "remove_rows"
# Percent to cut from each tail (example: 1.0 means P1 and P99).
OUTLIER_CLIP_TAIL_PCT = 0.25
# If True, clipping is also applied to the model target.
OUTLIER_CLIP_APPLY_TO_TARGET = True
# Optional feature names to exclude from clipping.
OUTLIER_CLIP_EXCLUDE_FEATURES = []
# Variables to explicitly show in outlier summary (only these are listed in report).
OUTLIER_REPORT_VARIABLES = [
    "obj_livingSpace",
    "obj_numberOfFloors",
    "obj_noParkSpaces"
]

# Worked-example configuration for report readability.
EXAMPLE_N_CASES = 5
EXAMPLE_CITY_PATTERNS = [
    ("Berlin", r"berlin"),
    ("Hamburg", r"hamburg"),
    ("Munich", r"muenchen|munich"),
    ("Cologne", r"koeln|cologne"),
    ("Frankfurt", r"frankfurt")
]

PLZ_OBS_SCENARIOS = [
    {"scenario_code": "N0", "min_plz_obs": 0},
    {"scenario_code": "N1", "min_plz_obs": 1},
    {"scenario_code": "N2", "min_plz_obs": 2},
    {"scenario_code": "N3", "min_plz_obs": 3},
    {"scenario_code": "N4", "min_plz_obs": 4},
    {"scenario_code": "N5", "min_plz_obs": 5},
    {"scenario_code": "N6", "min_plz_obs": 6},
    {"scenario_code": "N7", "min_plz_obs": 7},
    {"scenario_code": "N8", "min_plz_obs": 8},
    {"scenario_code": "N9", "min_plz_obs": 9},
    {"scenario_code": "N10", "min_plz_obs": 10}
]

# ============================================================
# EXECUTION MODE
# ============================================================
# DEBUG_MODE = True:
# - runs one scenario only (faster iteration)
# DEBUG_MODE = False:
# - runs full scenario grid (final comparison tables)

DEBUG_MODE = False

# Choose one PLZ scenario when DEBUG_MODE = True.
# Options: "N0", "N3", "N4", "N5", "N10"
DEBUG_SCENARIO_CODE = "N6"

# Choose imputation setting only when DEBUG_MODE = True
# False = NO IMPUTATION
# True  = WITH IMPUTATION
DEBUG_USE_IMPUTATION = True

IMPUTATION_SCENARIOS = [False, True]

timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

safe_target_name = MODEL_TARGET.replace("/", "_").replace(" ", "_")

OUTPUT_EXCEL = f"MLR_{safe_target_name}_PLZObs_no_imputation_vs_imputation_stepwise_{timestamp}.xlsx"
DOCX_OUTPUT = f"MLR_{safe_target_name}_PLZObs_no_imputation_vs_imputation_stepwise_{timestamp}.docx"
warnings.filterwarnings("ignore")

# Execution-time marker for full script runtime.
SCRIPT_START_TIME = perf_counter()

# Optional geospatial residual map for visualization.
ENABLE_GERMANY_RESIDUAL_MAP = True
ZIP_COL_FOR_MAP = "obj_zipCode"
# Minimum observations per ZIP in plotted map aggregation.
# 1 = maximum spatial coverage; 2/3 = more robust ZIP-level averages.
RESIDUAL_MAP_MIN_LISTINGS_PER_ZIP = 1
# Cap plotted ZIP points for readability/performance.
RESIDUAL_MAP_MAX_POINTS = 1200
# "test": map only out-of-sample errors (methodologically stricter).
# "all_rows": map all modeled rows (best visual territorial coverage).
RESIDUAL_MAP_SCOPE = "all_rows"
# Color-scale controls for map bubbles.
# Options for RESIDUAL_MAP_COLOR_COLUMN:
# - mean_residual       (signed error in target unit)
# - mean_abs_residual   (absolute error in target unit)
# - mean_pct_error      (signed percentage error)
# - mape_pct            (absolute percentage error)
# - smape_pct           (symmetric absolute percentage error)
RESIDUAL_MAP_COLOR_COLUMN = "mean_residual"
# Friendly palette name used in this script (user-friendly label).
# red_blue_diverging_reversed -> Matplotlib "RdBu_r" (negative = blue, positive = red)
# red_blue_diverging          -> Matplotlib "RdBu"
# yellow_orange_red           -> Matplotlib "YlOrRd"
# blue_green                  -> Matplotlib "YlGnBu"
RESIDUAL_MAP_PALETTE = "red_blue_diverging_reversed"

MAP_PALETTE_TO_CMAP = {
    "red_blue_diverging_reversed": "RdBu_r",
    "red_blue_diverging": "RdBu",
    "yellow_orange_red": "YlOrRd",
    "blue_green": "YlGnBu",
}

MAP_PALETTE_DESCRIPTION = {
    "red_blue_diverging_reversed": (
        "Blue indicates negative average residuals and red indicates positive average residuals."
    ),
    "red_blue_diverging": (
        "Red indicates negative average residuals and blue indicates positive average residuals."
    ),
    "yellow_orange_red": "Yellow to red indicates increasing average residual values.",
    "blue_green": "Light blue to green indicates increasing average residual values.",
}

RESIDUAL_MAP_CMAP = MAP_PALETTE_TO_CMAP.get(
    RESIDUAL_MAP_PALETTE,
    "RdBu_r"
)

# Symmetric scale around zero is recommended for signed errors.
RESIDUAL_MAP_SYMMETRIC_SCALE = True
# Percentile clipping reduces outlier domination in color scale.
RESIDUAL_MAP_CLIP_PCT_LOW = 2.0
RESIDUAL_MAP_CLIP_PCT_HIGH = 98.0

if RESIDUAL_MAP_COLOR_COLUMN == "mean_residual" and MODEL_TARGET_UNIT:
    RESIDUAL_MAP_COLOR_LABEL = f"Mean residual ({MODEL_TARGET_UNIT})"
elif RESIDUAL_MAP_COLOR_COLUMN in {"mean_pct_error", "mean_abs_pct_error", "mape_pct", "smape_pct"}:
    RESIDUAL_MAP_COLOR_LABEL = "Percentage error (%)"
else:
    RESIDUAL_MAP_COLOR_LABEL = None

RESIDUAL_MAP_PALETTE_TEXT = MAP_PALETTE_DESCRIPTION.get(
    RESIDUAL_MAP_PALETTE,
    "Color indicates average residual values."
)


# ============================================================
# GLOBAL PLOT STYLE
# ============================================================
plt.rcParams.update({
    "figure.facecolor": "white",
    "axes.facecolor": "white",
    "axes.edgecolor": "black",
    "axes.titlesize": 12,
    "axes.labelsize": 10,
    "xtick.labelsize": 9,
    "ytick.labelsize": 9,
    "legend.fontsize": 9,
    "font.size": 10
})

POINT_COLOR = "black"
SMOOTH_COLOR = "#d62728"
REFERENCE_COLOR = "gray"
OK_COLOR = "#228B22"
NOT_OK_COLOR = "#cc0000"
CHECK_COLOR = "darkorange"

DEFAULT_ALPHA = 0.65
DEFAULT_MARKER_SIZE = 18


# ============================================================
# WORD / DOCX HELPERS
# ============================================================
doc = Document()
doc.add_heading(
    f"MLR Debug Report - {MODEL_TARGET_LABEL} - No Imputation vs With Imputation",
    0
)

def add_section(title):
    """Add a level-1 heading to the Word report."""
    doc.add_heading(str(title), level=1)


def add_subsection(title):
    """Add a level-2 heading to the Word report."""
    doc.add_heading(str(title), level=2)


def add_text(text):
    """Add a paragraph to the Word report."""
    doc.add_paragraph(str(text))


def add_list(items):
    """Add a bullet list to the Word report."""
    if not items:
        doc.add_paragraph("(no items)")
    else:
        for item in items:
            doc.add_paragraph(str(item), style="List Bullet")


def add_figure_to_doc(fig, width=6.5):
    """
    Save a Matplotlib figure to an in-memory buffer and insert it into
    the Word document without writing any PNG file to disk.
    """
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=300, bbox_inches="tight")
    buf.seek(0)
    doc.add_picture(buf, width=Inches(width))
    buf.close()


# ============================================================
# GENERAL HELPERS
def add_model_target_info():
    """
    Add model target information to the Word report.
    """
    add_subsection("Model Target Variable")

    add_text(f"Model target: {MODEL_TARGET}")
    add_text(f"Target label: {MODEL_TARGET_LABEL}")

    if MODEL_TARGET_UNIT:
        add_text(f"Unit: {MODEL_TARGET_UNIT}")

    if MODEL_TARGET == TARGET_RENT_SQM:
        add_text(
            "The target variable represents rent per square meter and was calculated as "
            "obj_totalRent divided by obj_livingSpace."
        )
        add_text(
            "Because living space is part of the denominator of the dependent variable, "
            "obj_livingSpace was excluded from the explanatory variables to avoid a mechanical relationship."
        )

    elif MODEL_TARGET == TARGET:
        add_text(
            "The target variable represents the total monthly rent."
        )


def add_global_model_input_max_summary(df_source: pd.DataFrame):
    """
    Add global maximum values for selected model input variables.
    These values are computed from the loaded model dataset, not scenario subsets.
    """
    add_section("GLOBAL MODEL INPUT MAX VALUES")
    add_text(
        "These maxima are computed from the loaded model dataset before scenario filtering."
    )

    for col in ["obj_livingSpace", "obj_numberOfFloors", "obj_noParkSpaces"]:
        if col not in df_source.columns:
            add_text(f"{col}: not available in loaded columns.")
            continue

        series_num = pd.to_numeric(df_source[col], errors="coerce")
        if series_num.notna().any():
            max_val = float(series_num.max())
            add_text(f"Max {col}: {max_val:,.2f}")
        else:
            add_text(f"Max {col}: NA (no valid numeric values)")


def filter_by_min_plz_observations(X, y, groups, min_plz_obs):
    """
    Filter observations according to the minimum number of listings per PLZ.

    N0  -> no minimum PLZ-count filter
    N3  -> keep only PLZs with at least 3 observations
    N5  -> keep only PLZs with at least 5 observations
    N10 -> keep only PLZs with at least 10 observations
    """

    if min_plz_obs <= 0:
        removed_group_counts = pd.DataFrame(columns=[GROUP_COL, "n_obs"])
        return X.copy(), y.copy(), groups.copy(), removed_group_counts

    group_counts = groups.value_counts()

    valid_groups = group_counts[
        group_counts >= min_plz_obs
    ].index

    removed_group_counts = (
        group_counts[group_counts < min_plz_obs]
        .reset_index()
    )

    removed_group_counts.columns = [
        GROUP_COL,
        "n_obs"
    ]

    valid_mask = groups.isin(valid_groups)

    X_filtered = X.loc[valid_mask].copy()
    y_filtered = y.loc[valid_mask].copy()
    groups_filtered = groups.loc[valid_mask].copy()

    return X_filtered, y_filtered, groups_filtered, removed_group_counts


def create_feature_missing_report(X):
    """
    Create a descriptive missing-value report.
    This report is exported only for documentation.
    It is NOT used for feature selection in this corrected version.
    """

    missing_share = X.isna().mean().sort_values(ascending=False)

    missing_share_df = (
        missing_share
        .reset_index()
        .rename(columns={"index": "feature", 0: "missing_share"})
    )

    return missing_share_df

def save_individual_diagnostic_plots(
    fitted_vals, residuals,
    rainbow_pvalue, shapiro_pvalue,
    bp_pvalue, dw_stat,
    scenario_name
):
    fitted_vals = np.asarray(fitted_vals, dtype=float)
    residuals = np.asarray(residuals, dtype=float)

    if len(fitted_vals) == 0 or len(residuals) == 0:
        return

    max_points = 500
    rng = np.random.default_rng(RANDOM_STATE)
    if len(fitted_vals) > max_points:
        sample_idx = rng.choice(len(fitted_vals), size=max_points, replace=False)
        fitted_plot = fitted_vals[sample_idx]
        residuals_plot = residuals[sample_idx]
    else:
        fitted_plot = fitted_vals
        residuals_plot = residuals

    resid_std = float(np.nanstd(residuals_plot, ddof=1)) if len(residuals_plot) > 1 else 0.0
    if resid_std > 0 and np.isfinite(resid_std):
        std_resid_plot = residuals_plot / resid_std
    else:
        std_resid_plot = residuals_plot.copy()

    x_lim, _ = get_axis_limits(fitted_plot, residuals_plot)
    x_lim_bp = get_symmetric_centered_xlim(fitted_plot, low_pct=1, high_pct=99, pad_ratio=0.08)
    y_lim = get_symmetric_residual_ylim(residuals_plot, percentile=99, pad_ratio=0.12)
    sqrt_abs_resid = np.sqrt(np.abs(std_resid_plot))
    y_lim_dw = (y_lim[0] * 1.12, y_lim[1] * 1.12)

    # -------- 1 Residuals vs Fitted
    fig, ax = plt.subplots(figsize=(6, 4))
    ax.scatter(
        fitted_plot,
        residuals_plot,
        color=POINT_COLOR,
        alpha=0.30,
        s=10,
        edgecolors="none"
    )
    add_lowess_line(ax, fitted_plot, residuals_plot, frac=0.80)
    ax.axhline(0, linestyle='--')
    ax.set_xlim(x_lim)
    ax.set_ylim(y_lim)

    label, color = pass_fail_label_from_pvalue(rainbow_pvalue)
    add_status_box(ax, "Rainbow test", format_p_value(rainbow_pvalue), label, color)

    ax.set_title("Residuals vs Fitted")
    ax.set_xlabel("Fitted values")
    ax.set_ylabel("Residuals")
    format_axis_pretty(ax)
    add_figure_to_doc(fig)
    plt.close(fig)

    # -------- 2 QQ
    fig, ax = plt.subplots(figsize=(6, 4))
    osm_q, osr_q = stats.probplot(std_resid_plot, dist="norm", fit=False)
    ax.scatter(osm_q, osr_q, color=POINT_COLOR, alpha=0.30, s=12, edgecolors="none")
    qq_abs = max(np.max(np.abs(osm_q)), np.max(np.abs(osr_q)), 1e-9)
    qq_lim = qq_abs * 1.05
    ax.plot([-qq_lim, qq_lim], [-qq_lim, qq_lim], linestyle="--", color=REFERENCE_COLOR, linewidth=1)

    label, color = pass_fail_label_from_pvalue(shapiro_pvalue)
    add_status_box(ax, "Shapiro test", format_p_value(shapiro_pvalue), label, color)

    ax.set_title("QQ Plot")
    ax.set_xlabel("Theoretical Quantiles")
    ax.set_ylabel("Standardized Residuals")
    ax.set_xlim(-qq_lim, qq_lim)
    ax.set_ylim(-qq_lim, qq_lim)
    ax.set_aspect("equal", adjustable="box")
    format_axis_pretty(ax)
    add_figure_to_doc(fig)
    plt.close(fig)

    # -------- 3 Scale Location
    fig, ax = plt.subplots(figsize=(6, 4))
    ax.scatter(
        fitted_plot,
        sqrt_abs_resid,
        color=POINT_COLOR,
        alpha=0.30,
        s=10,
        edgecolors="none"
    )
    add_lowess_line(ax, fitted_plot, sqrt_abs_resid, frac=0.80)
    ax.set_xlim(x_lim_bp)
    y95 = float(np.nanpercentile(sqrt_abs_resid, 95)) if len(sqrt_abs_resid) else 1.0
    y99 = float(np.nanpercentile(sqrt_abs_resid, 99)) if len(sqrt_abs_resid) else y95
    y_upper = max(y95 * 1.10, y99)
    ax.set_ylim(0, y_upper)

    label, color = pass_fail_label_from_pvalue(bp_pvalue)
    add_status_box(ax, "Breusch-Pagan", format_p_value(bp_pvalue), label, color)

    ax.set_title("Scale Location")
    ax.set_xlabel("Fitted values")
    ax.set_ylabel("Sqrt(|Standardized residuals|)")
    format_axis_pretty(ax)
    add_figure_to_doc(fig)
    plt.close(fig)

    # -------- 4 Residual Order
    fig, ax = plt.subplots(figsize=(6, 4))
    ax.scatter(np.arange(len(residuals_plot)), residuals_plot, color=POINT_COLOR, alpha=0.30, s=10, edgecolors="none")
    ax.axhline(0, color=REFERENCE_COLOR, linestyle='--', linewidth=1)

    label, color = pass_fail_label_from_dw(dw_stat)
    add_status_box(ax, "Durbin Watson", format_stat_value("DW", dw_stat), label, color)

    ax.set_title("Residuals vs Order")
    ax.set_xlabel("Observation Order")
    ax.set_ylabel("Residuals")
    ax.set_ylim(y_lim_dw)
    format_axis_pretty(ax)
    add_figure_to_doc(fig)
    plt.close(fig)


def normalize_binary_series(series: pd.Series) -> pd.Series:
    """
    Normalize binary variables to 0.0 / 1.0.

    Accepted values include:
    - strings: Y/N, YES/NO, TRUE/FALSE, 1/0, 1.0/0.0
    - numbers: 1, 0, 1.0, 0.0
    """
    s = series.astype("object")
    s = s.where(pd.notna(s), np.nan)

    def _normalize(value):
        if pd.isna(value):
            return np.nan

        # Already numeric
        if isinstance(value, (int, float, np.integer, np.floating)):
            if np.isclose(value, 1):
                return 1.0
            if np.isclose(value, 0):
                return 0.0
            return np.nan

        # String value
        value = str(value).strip().upper()

        if value in {"Y", "YES", "TRUE", "T", "1", "1.0"}:
            return 1.0

        if value in {"N", "NO", "FALSE", "F", "0", "0.0"}:
            return 0.0

        return np.nan

    return s.map(_normalize).astype("float")


def apply_global_outlier_clipping(
    X: pd.DataFrame,
    y: pd.Series,
    tail_pct: float,
    apply_to_target: bool,
    mode: str = "clip_values",
    exclude_features=None
):
    """
        Apply global outlier treatment on numeric variables by percentile tails.

        mode="clip_values":
            - keep all rows and clip values to [P_low, P_high].

        mode="remove_rows":
            - remove rows where at least one monitored variable is outside [P_low, P_high].
    """
    if exclude_features is None:
        exclude_features = []

    X_out = X.copy()
    y_out = y.copy()

    try:
        tail_pct = float(tail_pct)
    except Exception:
        tail_pct = 0.0

    if tail_pct <= 0 or tail_pct >= 50:
        empty_thresholds = pd.DataFrame(columns=["variable", "low", "high", "n_rows_flagged"])
        empty_excluded = pd.DataFrame(columns=["row_index", "trigger_variables", "trigger_details", "n_trigger_variables"])
        return X_out, y_out, empty_thresholds, empty_excluded

    low_q = tail_pct
    high_q = 100.0 - tail_pct

    clipping_rows = []
    outlier_flags = pd.DataFrame(index=X_out.index)

    numeric_cols = X_out.select_dtypes(include=["number"]).columns.tolist()
    for col in numeric_cols:
        if col in exclude_features:
            continue

        non_na = X_out[col].dropna()
        if non_na.empty:
            continue

        unique_vals = set(non_na.unique())
        if unique_vals.issubset({0, 1}):
            continue

        low_val = np.nanpercentile(non_na.to_numpy(dtype=float), low_q)
        high_val = np.nanpercentile(non_na.to_numpy(dtype=float), high_q)

        if np.isfinite(low_val) and np.isfinite(high_val) and low_val < high_val:
            col_values = X_out[col].to_numpy(dtype=float)
            col_flags = (col_values < low_val) | (col_values > high_val)
            outlier_flags[col] = col_flags

            if mode == "clip_values":
                X_out[col] = X_out[col].clip(lower=low_val, upper=high_val)

            clipping_rows.append(
                {
                    "variable": col,
                    "low": low_val,
                    "high": high_val,
                    "n_rows_flagged": int(np.nansum(col_flags))
                }
            )

    if apply_to_target:
        y_non_na = y_out.dropna()
        if not y_non_na.empty:
            y_low = np.nanpercentile(y_non_na.to_numpy(dtype=float), low_q)
            y_high = np.nanpercentile(y_non_na.to_numpy(dtype=float), high_q)
            if np.isfinite(y_low) and np.isfinite(y_high) and y_low < y_high:
                y_values = y_out.to_numpy(dtype=float)
                y_flags = (y_values < y_low) | (y_values > y_high)
                outlier_flags[MODEL_TARGET] = y_flags

                if mode == "clip_values":
                    y_out = y_out.clip(lower=y_low, upper=y_high)

                clipping_rows.append(
                    {
                        "variable": MODEL_TARGET,
                        "low": y_low,
                        "high": y_high,
                        "n_rows_flagged": int(np.nansum(y_flags))
                    }
                )

    excluded_rows_df = pd.DataFrame(columns=["row_index", "trigger_variables", "trigger_details", "n_trigger_variables"])

    if mode == "remove_rows" and not outlier_flags.empty:
        row_outlier_mask = outlier_flags.any(axis=1)
        excluded_index = outlier_flags.index[row_outlier_mask]
        threshold_map = {
            row["variable"]: (row["low"], row["high"])
            for _, row in pd.DataFrame(clipping_rows).iterrows()
        }

        excluded_rows = []
        for idx in excluded_index:
            triggers = outlier_flags.columns[outlier_flags.loc[idx].to_numpy(dtype=bool)].tolist()
            details = []

            for var in triggers:
                low_val, high_val = threshold_map.get(var, (np.nan, np.nan))

                if var == MODEL_TARGET:
                    observed = y_out.loc[idx] if idx in y_out.index else np.nan
                else:
                    observed = X_out.loc[idx, var] if (idx in X_out.index and var in X_out.columns) else np.nan

                side = "inside"
                if pd.notna(observed) and np.isfinite(low_val) and observed < low_val:
                    side = "below_low"
                elif pd.notna(observed) and np.isfinite(high_val) and observed > high_val:
                    side = "above_high"

                if pd.notna(observed):
                    observed_txt = f"{float(observed):.4f}"
                else:
                    observed_txt = "NA"

                details.append(
                    f"{var}: observed={observed_txt}, ref=[{low_val:.4f}, {high_val:.4f}], side={side}"
                )

            excluded_rows.append(
                {
                    "row_index": int(idx) if isinstance(idx, (int, np.integer)) else str(idx),
                    "trigger_variables": ", ".join(triggers),
                    "trigger_details": " | ".join(details),
                    "n_trigger_variables": len(triggers)
                }
            )

        excluded_rows_df = pd.DataFrame(excluded_rows)

        keep_mask = ~row_outlier_mask
        X_out = X_out.loc[keep_mask].copy()
        y_out = y_out.loc[keep_mask].copy()

    clipping_df = pd.DataFrame(clipping_rows)
    return X_out, y_out, clipping_df, excluded_rows_df


def make_ohe():
    """
    Create a OneHotEncoder that works across sklearn versions
    (sparse_output=False in newer versions, sparse=False in older ones).
    """
    try:
        return OneHotEncoder(
            drop="first",
            handle_unknown="ignore",
            sparse_output=False
        )
    except TypeError:
        return OneHotEncoder(
            drop="first",
            handle_unknown="ignore",
            sparse=False
        )


def add_engineered_features(df_input: pd.DataFrame, reference_year=REFERENCE_YEAR) -> pd.DataFrame:

    """
    Create additional engineered features from variables with many missing values.

    Strategy:
    - obj_energyType: categorical with 'Unknown'
    - obj_thermalChar: numeric
    - obj_numberOfFloors: indicator + numeric value
    - obj_noParkSpaces: indicator + numeric value
    - obj_lastRefurbish: indicator + years since refurbishment
    - obj_yearConstructed: building age
    - interaction: years between construction and refurbishment
    """
    df = df_input.copy()
    current_year = reference_year

    # ----------------------------------------
    # 1) Energy type (categorical)
    # ----------------------------------------
    if "obj_energyType" in df.columns:
        df["obj_energyType_cat"] = (
            df["obj_energyType"]
            .astype("object")
            .where(pd.notna(df["obj_energyType"]), "Unknown")
            .astype(str)
        )
    else:
        df["obj_energyType_cat"] = "Unknown"

    # ----------------------------------------
    # 2) Thermal characteristic (numeric)
    # ----------------------------------------
    if "obj_thermalChar" in df.columns:
        thermal_num = pd.to_numeric(df["obj_thermalChar"], errors="coerce")
        df["obj_thermalChar_num"] = thermal_num.fillna(0)
    else:
        df["obj_thermalChar_num"] = 0.0

    # ----------------------------------------
    # 3) Number of floors
    # ----------------------------------------
    if "obj_numberOfFloors" in df.columns:
        number_of_floors_num = pd.to_numeric(df["obj_numberOfFloors"], errors="coerce")
        valid_floors = number_of_floors_num.notna()

        df["obj_hasNumberOfFloorsInfo"] = valid_floors.astype(float)
        df["obj_numberOfFloors_num"] = number_of_floors_num.fillna(0)
    else:
        df["obj_hasNumberOfFloorsInfo"] = 0.0
        df["obj_numberOfFloors_num"] = 0.0

    # ----------------------------------------
    # 4) Parking spaces
    # ----------------------------------------
    if "obj_noParkSpaces" in df.columns:
        parking_num = pd.to_numeric(df["obj_noParkSpaces"], errors="coerce")
        valid_parking = parking_num.notna()

        df["obj_hasParkingInfo"] = valid_parking.astype(float)
        df["obj_noParkSpaces_num"] = parking_num.fillna(0)
    else:
        df["obj_hasParkingInfo"] = 0.0
        df["obj_noParkSpaces_num"] = 0.0

    # ----------------------------------------
    # 5) Last refurbishment
    # ----------------------------------------
    if "obj_lastRefurbish" in df.columns:
        last_refurbish_year = pd.to_numeric(df["obj_lastRefurbish"], errors="coerce")
        valid_refurbish = (
            last_refurbish_year.notna()
            & (last_refurbish_year > 1800)
            & (last_refurbish_year <= current_year)
        )

        df["obj_hasLastRefurbishInfo"] = valid_refurbish.astype(float)
        df["obj_yearsSinceLastRefurbish"] = np.where(
            valid_refurbish,
            current_year - last_refurbish_year,
            0
        )
    else:
        last_refurbish_year = pd.Series(np.nan, index=df.index)
        valid_refurbish = pd.Series(False, index=df.index)

        df["obj_hasLastRefurbishInfo"] = 0.0
        df["obj_yearsSinceLastRefurbish"] = 0.0

    # ----------------------------------------
    # 6) Year constructed
    # ----------------------------------------
    if "obj_yearConstructed" in df.columns:
        constructed_year = pd.to_numeric(df["obj_yearConstructed"], errors="coerce")
        valid_constructed = (
            constructed_year.notna()
            & (constructed_year > 1800)
            & (constructed_year <= current_year)
        )

        df["obj_buildingAge"] = np.where(
            valid_constructed,
            current_year - constructed_year,
            0
        )
    else:
        constructed_year = pd.Series(np.nan, index=df.index)
        valid_constructed = pd.Series(False, index=df.index)

        df["obj_buildingAge"] = 0.0

    # ----------------------------------------
    # 7) Construction x refurbishment interaction
    # ----------------------------------------
    valid_both = valid_refurbish & valid_constructed

    df["obj_yearsBetweenConstructionAndRefurbish"] = np.where(
        valid_both & (last_refurbish_year >= constructed_year),
        last_refurbish_year - constructed_year,
        0
    )

    # ----------------------------------------
    # 8) ZIP 3-digit regional prefix
    # ----------------------------------------
    if "obj_zipCode" in df.columns:
        df["obj_zip3"] = (
            df["obj_zipCode"].astype(str)
            .str.replace(r"\.0$", "", regex=True)
            .str.zfill(5)
            .str[:3]
        )
    else:
        df["obj_zip3"] = "unknown"

    return df


def make_preprocessor(numeric_features, categorical_features, use_imputation=False):
    """
    Build the preprocessing object for both scenarios:
    - no imputation: passthrough for numeric features, one-hot for categorical
    - with imputation: median for numeric, most_frequent for categorical
    """
    if use_imputation:
        numeric_transformer = Pipeline([
            ("imputer", SimpleImputer(strategy="median"))
        ])

        categorical_transformer = Pipeline([
            ("imputer", SimpleImputer(strategy="most_frequent")),
            ("onehot", make_ohe())
        ])
    else:
        numeric_transformer = "passthrough"

        categorical_transformer = Pipeline([
            ("onehot", make_ohe())
        ])

    preprocessor = ColumnTransformer(
        transformers=[
            ("num", numeric_transformer, numeric_features),
            ("cat", categorical_transformer, categorical_features)
        ],
        remainder="drop",
        sparse_threshold=0
    )

    return preprocessor


def get_clean_feature_names(preprocessor, numeric_features, categorical_features):
    """
    Get clean feature names after the ColumnTransformer.
    """
    try:
        feature_names_out = list(preprocessor.get_feature_names_out())
        feature_names_out = [
            name.replace("num__", "").replace("cat__", "")
            for name in feature_names_out
        ]
        return feature_names_out

    except Exception:
        feature_names_out = []
        feature_names_out.extend(numeric_features)

        if len(categorical_features) > 0:
            cat_pipeline = preprocessor.named_transformers_["cat"]
            ohe = cat_pipeline.named_steps["onehot"]
            cat_names = ohe.get_feature_names_out(categorical_features)
            feature_names_out.extend(list(cat_names))

        return feature_names_out


# ============================================================
# DIAGNOSTIC PLOT HELPERS
# ============================================================
def format_p_value(p_value):
    """Format p-values for plot annotations."""
    if pd.isna(p_value):
        return "p = NA"
    if p_value < 0.0001:
        return "p < 0.0001"
    return f"p = {p_value:.4f}"


def format_stat_value(name, value):
    """Format test statistics for plot annotations."""
    if pd.isna(value):
        return f"{name} = NA"
    return f"{name} = {value:.4f}"


def pass_fail_label_from_pvalue(p_value, alpha=0.05):
    """
    For most assumption checks:
    - p > alpha  -> assumption not rejected -> OK
    - p <= alpha -> possible violation      -> NOT OK
    """
    if pd.isna(p_value):
        return "CHECK", CHECK_COLOR

    if p_value > alpha:
        return "OK", OK_COLOR
    return "NOT OK", NOT_OK_COLOR


def pass_fail_label_from_dw(dw_stat):
    """
    Soft heuristic label for residual independence based on Durbin-Watson.
    """
    if pd.isna(dw_stat):
        return "CHECK", CHECK_COLOR

    if 1.5 <= dw_stat <= 2.5:
        return "OK", OK_COLOR
    return "NOT OK", NOT_OK_COLOR


def get_axis_limits(fitted_vals, residuals):
    """
    Compute robust axis limits with small margins for cleaner thesis plots.
    """
    x_vals = np.asarray(fitted_vals, dtype=float)
    y_vals = np.asarray(residuals, dtype=float)

    x_vals = x_vals[np.isfinite(x_vals)]
    y_vals = y_vals[np.isfinite(y_vals)]

    if x_vals.size == 0 or y_vals.size == 0:
        return (-1.0, 1.0), (-1.0, 1.0)

    x_min = float(np.nanpercentile(x_vals, 1))
    x_max = float(np.nanpercentile(x_vals, 99))
    y_min = float(np.nanpercentile(y_vals, 1))
    y_max = float(np.nanpercentile(y_vals, 99))

    x_span = max(x_max - x_min, 1e-9)
    y_span = max(y_max - y_min, 1e-9)

    x_pad = 0.08 * x_span
    y_pad = 0.10 * y_span

    x_lim = (x_min + (-x_pad), x_max + x_pad)
    y_lim = (y_min + (-y_pad), y_max + y_pad)

    return x_lim, y_lim


def get_symmetric_residual_ylim(residuals, percentile=99, pad_ratio=0.12):
    """Build symmetric residual limits around zero for visually balanced plots."""
    values = np.asarray(residuals, dtype=float)
    values = values[np.isfinite(values)]

    if values.size == 0:
        return (-1.0, 1.0)

    robust_abs = float(np.nanpercentile(np.abs(values), percentile))
    robust_abs = max(robust_abs, 1e-9)
    lim = robust_abs * (1.0 + pad_ratio)
    return (-lim, lim)


def get_symmetric_centered_xlim(values, low_pct=1, high_pct=99, pad_ratio=0.08):
    """Build symmetric x-limits around the robust center of a variable."""
    arr = np.asarray(values, dtype=float)
    arr = arr[np.isfinite(arr)]

    if arr.size == 0:
        return (-1.0, 1.0)

    low = float(np.nanpercentile(arr, low_pct))
    high = float(np.nanpercentile(arr, high_pct))
    center = 0.5 * (low + high)
    half_span = max(high - center, center - low, 1e-9)
    half_span = half_span * (1.0 + pad_ratio)
    return (center - half_span, center + half_span)


def _pretty_tick(value, _pos):
    """Readable numeric ticks for thesis-ready plots."""
    def _format_eu_number(val, decimals):
        text = f"{val:,.{decimals}f}"
        return text.replace(",", "#").replace(".", ",").replace("#", ".")

    abs_val = abs(value)
    if abs_val >= 1000:
        return _format_eu_number(value, 0)
    if abs_val >= 100:
        return _format_eu_number(value, 1)
    if abs_val >= 1:
        return _format_eu_number(value, 2)
    return _format_eu_number(value, 3)


def format_axis_pretty(ax):
    """Apply consistent tick density and formatting to diagnostic axes."""
    ax.xaxis.set_major_locator(MaxNLocator(nbins=6))
    ax.yaxis.set_major_locator(MaxNLocator(nbins=6))
    ax.xaxis.set_major_formatter(FuncFormatter(_pretty_tick))
    ax.yaxis.set_major_formatter(FuncFormatter(_pretty_tick))


def annotate_top_residuals(ax, x_values, y_values, n_labels=3):
    """
    Annotate the observations with the largest absolute values.
    """
    x_values = np.asarray(x_values)
    y_values = np.asarray(y_values)

    if len(y_values) == 0:
        return

    n_labels = min(n_labels, len(y_values))
    top_idx = np.argsort(np.abs(y_values))[-n_labels:]

    for idx in top_idx:
        ax.annotate(
            str(idx),
            (x_values[idx], y_values[idx]),
            xytext=(4, 4),
            textcoords="offset points",
            fontsize=8,
            color="black"
        )


def add_debug_variable_inventory(df_model: pd.DataFrame, selected_features):
    """Add a debug-only inventory of columns available to the scenario report."""
    if not DEBUG_MODE:
        return

    available_columns = sorted(df_model.columns.tolist())
    selected_set = set(selected_features)
    non_model_columns = sorted([col for col in available_columns if col not in selected_set])

    add_subsection("DEBUG VARIABLE INVENTORY")
    add_text(f"Total columns available in scenario dataframe: {len(available_columns)}")
    add_text("All columns currently available to the scenario:")
    add_list(available_columns)
    add_text(f"Columns not used as direct model features: {len(non_model_columns)}")
    add_list(non_model_columns)


def add_status_box(ax, top_text, bottom_text, label, label_color):
    """
    Add a compact annotation box with a test summary inside a subplot.
    """
    annotation_text = f"{top_text}\n{bottom_text}"

    ax.text(
        0.03,
        0.97,
        annotation_text,
        transform=ax.transAxes,
        va="top",
        ha="left",
        fontsize=9,
        bbox=dict(
            boxstyle="round,pad=0.3",
            facecolor="white",
            edgecolor="lightgray",
            alpha=0.95
        )
    )

    ax.text(
        0.97,
        0.06,
        label,
        transform=ax.transAxes,
        va="bottom",
        ha="right",
        fontsize=16,
        fontweight="bold",
        color=label_color
    )


def add_lowess_line(ax, x, y, frac=0.65):
    """
    Add a LOWESS smoothing line to a scatter plot.
    """
    if len(x) < 3:
        return

    smooth = lowess(y, x, frac=frac, return_sorted=True)
    ax.plot(smooth[:, 0], smooth[:, 1], color=SMOOTH_COLOR, linewidth=1.5)


def map_outlier_report_variable_name(var_name: str) -> str:
    """Map report-friendly variable names to clipping variable names."""
    alias_map = {
        "obj_livingSpace": "obj_livingSpace",
        "obj_numberOfFloors": "obj_numberOfFloors_num",
        "obj_noParkSpaces": "obj_noParkSpaces_num"
    }
    return alias_map.get(var_name, var_name)


def build_selected_outlier_summary(clipping_df: pd.DataFrame, requested_variables) -> pd.DataFrame:
    """Return a concise outlier summary for selected variables only."""
    if clipping_df is None or clipping_df.empty:
        return pd.DataFrame(columns=["report_variable", "model_variable", "n_rows_flagged", "low", "high"])

    rows = []
    for report_var in requested_variables:
        model_var = map_outlier_report_variable_name(report_var)
        match = clipping_df.loc[clipping_df["variable"] == model_var]
        if match.empty:
            continue

        row = match.iloc[0]
        rows.append({
            "report_variable": report_var,
            "model_variable": model_var,
            "n_rows_flagged": int(row.get("n_rows_flagged", 0)),
            "low": row.get("low", np.nan),
            "high": row.get("high", np.nan)
        })

    return pd.DataFrame(rows)


def build_assumption_diagnostics_figure(
    fitted_vals,
    residuals,
    scenario_name,
    rainbow_pvalue=np.nan,
    shapiro_pvalue=np.nan,
    bp_pvalue=np.nan,
    dw_stat=np.nan
):
    """
    Build a publication-ready 2x2 OLS diagnostics panel with:
    1. Residuals vs Fitted
    2. Q-Q Residuals
    3. Scale-Location
    4. Residuals in Observation Order
    """

    # --------------------------------------------------------
    # Convert everything to NumPy arrays to avoid pandas index issues
    # --------------------------------------------------------
    fitted_vals = np.asarray(fitted_vals, dtype=float)
    residuals = np.asarray(residuals, dtype=float)
    residuals_raw = residuals.copy()

    # Optional clipping only for visualization
    residuals = np.clip(
        residuals,
        np.percentile(residuals, 1),
        np.percentile(residuals, 99)
    )

    residual_std = np.std(residuals, ddof=1)
    if residual_std > 0:
        std_resid = residuals / residual_std
    else:
        std_resid = residuals.copy()

    x_lim, _ = get_axis_limits(fitted_vals, residuals)
    x_lim_bp = get_symmetric_centered_xlim(fitted_vals, low_pct=1, high_pct=99, pad_ratio=0.08)
    y_lim = get_symmetric_residual_ylim(residuals, percentile=99, pad_ratio=0.12)
    sqrt_abs_resid = np.sqrt(np.abs(std_resid))

    order_idx = np.arange(1, len(residuals) + 1)

    fig, axes = plt.subplots(2, 2, figsize=(14, 9))
    axes = axes.flatten()

    # Reproducible subsampling for cleaner plots
    max_points = 150
    rng = np.random.default_rng(RANDOM_STATE)

    if len(fitted_vals) > max_points:
        idx = rng.choice(len(fitted_vals), size=max_points, replace=False)
    else:
        idx = np.arange(len(fitted_vals))

    # Subsampled versions for cleaner scatter plots
    fitted_plot = fitted_vals[idx]
    residuals_plot = residuals[idx]
    sqrt_abs_resid_plot = sqrt_abs_resid[idx]

    # --------------------------------------------------------
    # 1) Residuals vs Fitted (Linearity)
    # --------------------------------------------------------
    ax = axes[0]
    ax.scatter(
        fitted_plot,
        residuals_plot,
        color=POINT_COLOR,
        alpha=0.30,
        s=12,
        edgecolors="none"
    )
    ax.axhline(0, color=REFERENCE_COLOR, linestyle="--", linewidth=1)
    add_lowess_line(ax, fitted_plot, residuals_plot, frac=0.85)
    annotate_top_residuals(ax, fitted_plot, residuals_plot, n_labels=3)

    linearity_label, linearity_color = pass_fail_label_from_pvalue(
        rainbow_pvalue,
        alpha=SIGNIFICANCE_LEVEL
    )

    add_status_box(
        ax,
        "Rainbow linearity test",
        format_p_value(rainbow_pvalue),
        linearity_label,
        linearity_color
    )

    ax.set_title("Residuals vs Fitted")
    ax.set_xlabel("Fitted values")
    ax.set_ylabel("Residuals")
    ax.set_xlim(x_lim)
    ax.set_ylim(y_lim)
    format_axis_pretty(ax)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    # --------------------------------------------------------
    # 2) Q-Q Plot (Normality)
    # --------------------------------------------------------
    ax = axes[1]
    osm, osr = stats.probplot(std_resid, dist="norm", fit=False)

    if len(osm) > max_points:
        qq_idx = rng.choice(len(osm), size=max_points, replace=False)
        osm_plot = np.asarray(osm)[qq_idx]
        osr_plot = np.asarray(osr)[qq_idx]
    else:
        osm_plot = np.asarray(osm)
        osr_plot = np.asarray(osr)

    ax.scatter(
        osm_plot,
        osr_plot,
        color=POINT_COLOR,
        alpha=0.30,
        s=12,
        edgecolors="none"
    )

    qq_min = min(np.min(osm), np.min(osr))
    qq_max = max(np.max(osm), np.max(osr))
    qq_abs = max(abs(qq_min), abs(qq_max), 1e-9)
    qq_lim = qq_abs * 1.05
    ax.plot(
        [-qq_lim, qq_lim],
        [-qq_lim, qq_lim],
        linestyle="--",
        color=REFERENCE_COLOR,
        linewidth=1
    )
    ax.set_xlim(-qq_lim, qq_lim)
    ax.set_ylim(-qq_lim, qq_lim)
    ax.set_aspect("equal", adjustable="box")

    normality_label, normality_color = pass_fail_label_from_pvalue(
        shapiro_pvalue,
        alpha=SIGNIFICANCE_LEVEL
    )

    add_status_box(
        ax,
        "Shapiro-Wilk normality test",
        format_p_value(shapiro_pvalue),
        normality_label,
        normality_color
    )

    ax.set_title("Q-Q Residuals")
    ax.set_xlabel("Theoretical Quantiles")
    ax.set_ylabel("Standardized Residuals")
    format_axis_pretty(ax)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    # --------------------------------------------------------
    # 3) Scale-Location (Homoscedasticity)
    # --------------------------------------------------------
    ax = axes[2]
    ax.scatter(
        fitted_plot,
        sqrt_abs_resid_plot,
        color=POINT_COLOR,
        alpha=0.30,
        s=12,
        edgecolors="none"
    )
    add_lowess_line(ax, fitted_plot, sqrt_abs_resid_plot, frac=0.85)
    annotate_top_residuals(ax, fitted_plot, sqrt_abs_resid_plot, n_labels=3)

    homosced_label, homosced_color = pass_fail_label_from_pvalue(
        bp_pvalue,
        alpha=SIGNIFICANCE_LEVEL
    )

    add_status_box(
        ax,
        "Breusch-Pagan test",
        format_p_value(bp_pvalue),
        homosced_label,
        homosced_color
    )

    ax.set_title("Scale-Location")
    ax.set_xlabel("Fitted values")
    ax.set_ylabel("Sqrt(|Standardized residuals|)")
    ax.set_xlim(x_lim_bp)
    y95 = float(np.nanpercentile(sqrt_abs_resid, 95)) if len(sqrt_abs_resid) else 1.0
    y99 = float(np.nanpercentile(sqrt_abs_resid, 99)) if len(sqrt_abs_resid) else y95
    y_upper = max(y95 * 1.10, y99)
    ax.set_ylim(0, y_upper)
    format_axis_pretty(ax)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    # --------------------------------------------------------
    # 4) Residuals in Observation Order (Independence)
    # --------------------------------------------------------
    ax = axes[3]
    ax.plot(
        order_idx,
        residuals,
        color="black",
        linewidth=0.8,
        marker="o",
        markersize=2
    )
    ax.axhline(0, color=REFERENCE_COLOR, linestyle="--", linewidth=1)

    independence_label, independence_color = pass_fail_label_from_dw(dw_stat)

    add_status_box(
        ax,
        "Durbin-Watson statistic",
        format_stat_value("DW", dw_stat),
        independence_label,
        independence_color
    )

    ax.set_title("Residuals in Observation Order")
    ax.set_xlabel("Observation Order")
    ax.set_ylabel("Residuals")
    y_lim_dw = (y_lim[0] * 1.12, y_lim[1] * 1.12)
    ax.set_ylim(y_lim_dw)
    format_axis_pretty(ax)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    fig.suptitle(
        f"OLS Assumption Diagnostics - {scenario_name}",
        fontsize=14,
        fontweight="bold"
    )
    plt.tight_layout(rect=[0, 0.02, 1, 0.96])

    return fig

# ============================================================
# PEARSON CORRELATION
# ============================================================
def add_pearson_correlation_plot(X_input, y_input, target, scenario_name, use_imputation=False):
    """
    Create Pearson correlation visualizations using numeric variables only.
    Pure categorical variables are excluded before one-hot encoding.
    """
    add_section(f"PEARSON CORRELATION - {scenario_name}")

    df_corr = X_input.copy()
    df_corr[target] = y_input

    numeric_corr_df = df_corr.select_dtypes(include=["number"]).copy()

    if target not in numeric_corr_df.columns:
        add_text("The target is not available as a numeric variable for Pearson correlation.")
        return

    numeric_corr_df = numeric_corr_df.dropna(axis=1, how="all")

    if numeric_corr_df.shape[1] < 2:
        add_text("Too few numeric variables are available for Pearson correlation analysis.")
        return

    if use_imputation:
        for col in numeric_corr_df.columns:
            if numeric_corr_df[col].isna().any():
                numeric_corr_df[col] = numeric_corr_df[col].fillna(numeric_corr_df[col].median())
        add_text("Pearson correlation was computed after median imputation for numeric variables.")
    else:
        add_text("Pearson correlation was computed using available data. Pandas uses pairwise valid observations.")

    corr_matrix = numeric_corr_df.corr(method="pearson")

    # ------------------------------
    # Full heatmap
    # ------------------------------
    fig, ax = plt.subplots(figsize=(14, 10))
    sns.heatmap(
        corr_matrix,
        annot=True,
        fmt=".2f",
        cmap="coolwarm",
        center=0,
        square=False,
        linewidths=0.3,
        annot_kws={"size": 10},
        cbar_kws={"shrink": 0.8},
        ax=ax
    )
    ax.set_title(f"Pearson Correlation Matrix - {scenario_name}", fontsize=16)
    plt.xticks(rotation=45, ha="right", fontsize=10)
    plt.yticks(rotation=0, fontsize=10)
    plt.tight_layout()

    add_figure_to_doc(fig, width=7.2)
    plt.close(fig)

    # ------------------------------
    # Correlation with target
    # ------------------------------
    add_subsection(f"Pearson Correlation with {target}")

    target_corr = (
        corr_matrix[target]
        .drop(target)
        .sort_values(key=lambda x: x.abs(), ascending=False)
    )

    add_text("Top Pearson correlations with the target:")
    for feature, corr_value in target_corr.head(25).items():
        add_text(f"{feature}: {corr_value:.4f}")

    target_corr_plot = target_corr.head(25).sort_values()

    fig, ax = plt.subplots(figsize=(10, 8))
    ax.barh(target_corr_plot.index, target_corr_plot.values, color= "steelblue")
    ax.axvline(x=0, color="black", linewidth=0.8)
    ax.tick_params(axis="y", labelsize=10)
    ax.tick_params(axis="x", labelsize=10)
    ax.set_xlabel("Pearson correlation with target")
    ax.set_ylabel("Feature")
    ax.set_title(f"Top Pearson Correlations with {target} - {scenario_name}")
    plt.tight_layout()

    add_figure_to_doc(fig, width=6.8)
    plt.close(fig)

    add_text(
        "Note: Pearson correlation measures linear relationships among numeric variables. "
        "Pure categorical variables were excluded before one-hot encoding."
    )


def add_mlr_prediction_examples(
    scenario_name,
    ols_model,
    X_test_proc,
    X_test_raw,
    y_test,
    y_pred_test,
    feature_names_out,
    n_examples=5,
    top_contrib=8
):
    """
    Add worked OLS examples to the Word report, showing how betas compose prediction values.
    """
    if len(y_test) == 0:
        return

    add_section(f"WORKED OLS EXAMPLES - {scenario_name}")
    add_text(
        "Each example below shows the OLS prediction decomposition: "
        "predicted = intercept + sum(beta_i * x_i)."
    )

    x_index = X_test_raw.index
    selected_idx = []

    if "obj_regio1" in X_test_raw.columns:
        regio = X_test_raw["obj_regio1"].astype(str)
        for _, pattern in EXAMPLE_CITY_PATTERNS:
            city_idx = x_index[regio.str.contains(pattern, case=False, na=False, regex=True)]
            if len(city_idx) > 0 and city_idx[0] not in selected_idx:
                selected_idx.append(city_idx[0])
            if len(selected_idx) >= n_examples:
                break

    if len(selected_idx) < n_examples:
        remaining = [idx for idx in x_index if idx not in selected_idx]
        for idx in remaining:
            selected_idx.append(idx)
            if len(selected_idx) >= n_examples:
                break

    selected_idx = selected_idx[:n_examples]

    y_test_series = y_test.copy()
    y_pred_series = pd.Series(np.asarray(y_pred_test, dtype=float), index=y_test.index)

    coef_arr = np.asarray(ols_model.params, dtype=float)
    intercept = float(coef_arr[0])
    beta = coef_arr[1:]

    # Map row index to transformed matrix position.
    test_pos_map = {idx: pos for pos, idx in enumerate(x_index)}

    for k, idx in enumerate(selected_idx, start=1):
        pos = test_pos_map[idx]
        x_row = np.asarray(X_test_proc[pos], dtype=float)
        contrib = beta * x_row

        contrib_df = pd.DataFrame({
            "feature": feature_names_out,
            "x_value": x_row,
            "beta": beta,
            "contribution": contrib
        })
        contrib_df["abs_contribution"] = contrib_df["contribution"].abs()
        contrib_df = contrib_df.sort_values("abs_contribution", ascending=False)

        top_df = contrib_df.head(top_contrib).copy()
        others_sum = float(contrib.sum() - top_df["contribution"].sum())
        pred_from_equation = float(intercept + contrib.sum())

        actual_val = float(y_test_series.loc[idx]) if idx in y_test_series.index else np.nan
        predicted_val = float(y_pred_series.loc[idx]) if idx in y_pred_series.index else np.nan

        add_subsection(f"Example {k}")
        add_text(f"Row index: {idx}")
        if "obj_regio1" in X_test_raw.columns and idx in X_test_raw.index:
            add_text(f"Region: {X_test_raw.loc[idx, 'obj_regio1']}")
        if "obj_zipCode" in X_test_raw.columns and idx in X_test_raw.index:
            add_text(f"ZIP: {X_test_raw.loc[idx, 'obj_zipCode']}")

        add_text(f"Actual rent: {actual_val:.4f}")
        add_text(f"Predicted rent (model output): {predicted_val:.4f}")
        add_text(f"Predicted rent (equation check): {pred_from_equation:.4f}")

        add_text(f"Intercept (const): {intercept:.4f}")
        add_text("Top contributions beta_i * x_i:")
        for _, row in top_df.iterrows():
            add_text(
                f"{row['feature']}: x={row['x_value']:.4f}, "
                f"beta={row['beta']:.4f}, contrib={row['contribution']:.4f}"
            )

        add_text(f"Sum of remaining contributions: {others_sum:.4f}")
        add_text(
            "Equation summary: predicted = intercept + top contributions + remaining contributions."
        )


# ============================================================
# BACKWARD STEPWISE REGRESSION
# ============================================================
def run_backward_stepwise(
    X_train_proc,
    X_test_proc,
    y_train,
    y_test,
    feature_names_out,
    scenario_name,
    p_threshold=0.05,
    max_iter=100
):
    """
    Perform backward stepwise regression based on p-values.

    The algorithm starts with all processed features, removes the feature
    with the highest p-value above the threshold, and repeats until all
    remaining p-values are <= threshold.

    This function also generates:
    1) p-values of variables removed by iteration
    2) number of remaining features by iteration
    3) top final coefficients
    4) actual vs predicted plot
    5) residuals vs predicted plot
    """
    add_section(f"STEPWISE REGRESSION - {scenario_name}")

    X_train_sw = pd.DataFrame(
        X_train_proc,
        columns=feature_names_out,
        index=y_train.index
    )

    X_test_sw = pd.DataFrame(
        X_test_proc,
        columns=feature_names_out,
        index=y_test.index
    )

    selected_features = list(X_train_sw.columns)
    removed_features = []

    add_text(f"Backward stepwise regression using p-value threshold = {p_threshold}")

    for iteration in range(max_iter):
        if len(selected_features) == 0:
            add_text("No features remain in the stepwise model.")
            break

        X_train_sm_step = sm.add_constant(
            X_train_sw[selected_features],
            has_constant="add"
        )

        try:
            step_model = sm.OLS(y_train, X_train_sm_step).fit()
        except Exception as exc:
            add_text(f"Stepwise regression stopped due to an OLS error: {exc}")
            break

        pvalues = step_model.pvalues.drop("const", errors="ignore")
        pvalues = pvalues.replace([np.inf, -np.inf], np.nan).dropna()

        if pvalues.empty:
            add_text("No valid p-values were found during stepwise regression.")
            break

        worst_pvalue = float(pvalues.max())
        worst_feature = pvalues.idxmax()

        if worst_pvalue > p_threshold:
            selected_features.remove(worst_feature)

            removed_features.append({
                "iteration": iteration + 1,
                "removed_feature": worst_feature,
                "p_value": worst_pvalue,
                "n_features_remaining_after_removal": len(selected_features)
            })
        else:
            break

    # Final stepwise model
    X_train_sm_final = sm.add_constant(
        X_train_sw[selected_features],
        has_constant="add"
    )
    X_test_sm_final = sm.add_constant(
        X_test_sw[selected_features],
        has_constant="add"
    )

    stepwise_model = sm.OLS(y_train, X_train_sm_final).fit()
    y_pred_stepwise = stepwise_model.predict(X_test_sm_final)

    stepwise_rmse = float(np.sqrt(mean_squared_error(y_test, y_pred_stepwise)))
    stepwise_mae = float(mean_absolute_error(y_test, y_pred_stepwise))
    stepwise_r2 = float(r2_score(y_test, y_pred_stepwise))

    removed_df = pd.DataFrame(removed_features)
    selected_df = pd.DataFrame({"selected_feature": selected_features})

    stepwise_coef_df = pd.DataFrame({
        "feature": stepwise_model.params.index,
        "coef": stepwise_model.params.values,
        "p_value": stepwise_model.pvalues.values
    })

    add_text(f"Initial number of processed features: {len(feature_names_out)}")
    add_text(f"Final number of selected features: {len(selected_features)}")
    add_text(f"Number of removed features: {len(removed_features)}")

    add_subsection("Stepwise Metrics")
    add_text(f"RMSE: {stepwise_rmse:.4f}")
    add_text(f"MAE: {stepwise_mae:.4f}")
    add_text(f"R²: {stepwise_r2:.4f}")

    add_subsection("Removed Features")
    if removed_df.empty:
        add_text("No features were removed.")
    else:
        for _, row in removed_df.head(80).iterrows():
            add_text(
                f"Iteration {int(row['iteration'])}: removed {row['removed_feature']} "
                f"(p = {row['p_value']:.5f})"
            )

        if len(removed_df) > 80:
            add_text("The list was truncated in the Word report. The full list is available in Excel.")


    add_subsection("Stepwise OLS Summary")
    doc.add_paragraph(stepwise_model.summary().as_text())

    # --------------------------------------------------------
    # Plot 1: Removed variables by p-value
    # --------------------------------------------------------
    add_subsection("Stepwise Selection Path - Removed Variables")

    if not removed_df.empty:
        plot_removed_df = removed_df.copy()
        plot_removed_df["removed_feature_short"] = plot_removed_df["removed_feature"].astype(str)

        max_label_len = 45
        plot_removed_df["removed_feature_short"] = plot_removed_df["removed_feature_short"].apply(
            lambda x: x if len(x) <= max_label_len else x[:max_label_len] + "..."
        )

        plot_df = plot_removed_df.head(30).sort_values("iteration", ascending=True)

        fig, ax = plt.subplots(figsize=(12, max(6, 0.35 * len(plot_df))))
        ax.barh(plot_df["removed_feature_short"], plot_df["p_value"], color="steelblue")
        ax.axvline(
            x=p_threshold,
            color="red",
            linestyle="--",
            label=f"Threshold p = {p_threshold}"
        )
        ax.set_xlabel("p-value at removal")
        ax.set_ylabel("Removed feature")
        ax.set_title(f"Stepwise Removed Variables by p-value - {scenario_name}")
        ax.legend()
        plt.tight_layout()

        add_text(
            "This plot shows the variables removed by backward stepwise regression. "
            "The red vertical line shows the p-value threshold used to keep variables in the model."
        )

        add_figure_to_doc(fig, width=6.0)
        plt.close(fig)

    else:
        add_text("No variables were removed, so the removal plot was not generated.")

    # --------------------------------------------------------
    # Plot 2: Number of remaining features by iteration
    # --------------------------------------------------------
    add_subsection("Stepwise Model Size by Iteration")

    if not removed_df.empty:
        fig, ax = plt.subplots(figsize=(9, 5))
        ax.plot(
            removed_df["iteration"],
            removed_df["n_features_remaining_after_removal"],
            marker="o",
            color="black"
        )
        ax.set_xlabel("Iteration")
        ax.set_ylabel("Number of remaining features")
        ax.set_title(f"Number of Features During Stepwise Selection - {scenario_name}")
        plt.tight_layout()

        add_text(
            "This plot shows how the model size decreases across stepwise iterations."
        )

        add_figure_to_doc(fig, width=6.0)
        plt.close(fig)

    else:
        add_text("Because no variables were removed, the model size remained constant.")

    # --------------------------------------------------------
    # Plot 3: Top final coefficients
    # --------------------------------------------------------
    add_subsection("Top Stepwise Coefficients")

    stepwise_plot_df = stepwise_coef_df.copy()
    stepwise_plot_df = stepwise_plot_df[stepwise_plot_df["feature"] != "const"]
    stepwise_plot_df["abs_coef"] = stepwise_plot_df["coef"].abs()
    stepwise_plot_df = stepwise_plot_df.sort_values("abs_coef", ascending=False).head(20)
    stepwise_plot_df = stepwise_plot_df.sort_values("abs_coef", ascending=True)

    if not stepwise_plot_df.empty:
        fig, ax = plt.subplots(figsize=(10, 8))
        ax.barh(stepwise_plot_df["feature"], stepwise_plot_df["abs_coef"], color="slateblue")
        ax.set_xlabel("|Coefficient|")
        ax.set_ylabel("Feature")
        ax.set_title(f"Top 20 Stepwise Coefficients - {scenario_name}")
        plt.tight_layout()

        add_text(
            "This plot shows the largest absolute coefficients in the final stepwise model."
        )

        add_figure_to_doc(fig, width=6.0)
        plt.close(fig)

    # --------------------------------------------------------
    # Plot 4: Actual vs Predicted
    # --------------------------------------------------------
    add_subsection("Stepwise Actual vs Predicted")

    fig, ax = plt.subplots(figsize=(7, 7))
    ax.scatter(y_test, y_pred_stepwise, alpha=0.25, color="black")

    min_val = min(np.min(y_test), np.min(y_pred_stepwise))
    max_val = max(np.max(y_test), np.max(y_pred_stepwise))

    ax.plot(
        [min_val, max_val],
        [min_val, max_val],
        color="red",
        linestyle="--",
        label="Perfect prediction"
    )

    ax.set_xlabel(f"Actual {MODEL_TARGET_LABEL}")
    ax.set_ylabel(f"Predicted {MODEL_TARGET_LABEL}")
    ax.set_title(f"Stepwise Actual vs Predicted - {scenario_name}")
    ax.legend()
    plt.tight_layout()

    add_text(
        "This plot compares the observed rent values to the values predicted by the stepwise model. "
        "Points closer to the red diagonal line indicate better predictions."
    )

    add_figure_to_doc(fig, width=6.0)
    plt.close(fig)

    # --------------------------------------------------------
    # Plot 5: Residuals vs Predicted
    # --------------------------------------------------------
    add_subsection("Stepwise Residuals vs Predicted")

    stepwise_residuals = y_test - y_pred_stepwise

    fig, ax = plt.subplots(figsize=(8, 5))
    ax.scatter(y_pred_stepwise, stepwise_residuals, alpha=0.25, color="black")
    ax.axhline(y=0, color="red", linestyle="--")
    ax.set_xlabel(f"Predicted {MODEL_TARGET_LABEL}")
    ax.set_xlabel(f"Predicted {MODEL_TARGET_LABEL} ({MODEL_TARGET_UNIT})")
    ax.set_ylabel(f"Actual {MODEL_TARGET_LABEL} ({MODEL_TARGET_UNIT})")
    ax.set_title(f"Stepwise Residuals vs Predicted - {scenario_name}")
    plt.tight_layout()

    add_text(
        "This plot shows the residual behavior of the stepwise model. "
        "Ideally, the points should be randomly scattered around zero."
    )

    add_figure_to_doc(fig, width=6.0)
    plt.close(fig)

    return {
        "stepwise_model": stepwise_model,
        "stepwise_rmse": stepwise_rmse,
        "stepwise_mae": stepwise_mae,
        "stepwise_r2": stepwise_r2,
        "stepwise_selected_df": selected_df,
        "stepwise_removed_df": removed_df,
        "stepwise_coef_df": stepwise_coef_df
    }


def sanitize_german_zip_series(zip_series: pd.Series) -> pd.Series:
    """
    Convert zip-like values to canonical 5-digit German postal code strings.
    """
    cleaned = (
        zip_series.astype(str)
        .str.strip()
        .str.replace(r"\.0$", "", regex=True)
        .str.replace(r"\D", "", regex=True)
    )
    cleaned = cleaned.where(cleaned.str.len().between(4, 5))
    cleaned = cleaned.str.zfill(5)
    return cleaned


def build_residual_zip_table(df_model, y_test, y_pred_test_ols, scenario_name):
    """
    Build test-level and aggregated residual tables by German ZIP code.
    """
    if ZIP_COL_FOR_MAP not in df_model.columns:
        return pd.DataFrame(), pd.DataFrame()

    base = pd.DataFrame({
        "row_index": y_test.index,
        "scenario": scenario_name,
        "actual": y_test.values,
        "predicted": np.asarray(y_pred_test_ols, dtype=float)
    })
    base["residual"] = base["actual"] - base["predicted"]
    base["abs_residual"] = base["residual"].abs()
    eps = 1e-9
    denom_actual = np.where(np.abs(base["actual"].to_numpy(dtype=float)) > eps, np.abs(base["actual"].to_numpy(dtype=float)), np.nan)
    base["pct_error"] = (base["residual"].to_numpy(dtype=float) / denom_actual) * 100.0
    base["abs_pct_error"] = np.abs(base["pct_error"].to_numpy(dtype=float))
    smape_denom = np.abs(base["actual"].to_numpy(dtype=float)) + np.abs(base["predicted"].to_numpy(dtype=float)) + eps
    base["smape_pct"] = 200.0 * np.abs(base["residual"].to_numpy(dtype=float)) / smape_denom
    base["zip_code"] = sanitize_german_zip_series(
        df_model.loc[y_test.index, ZIP_COL_FOR_MAP]
    ).to_numpy()
    base = base.dropna(subset=["zip_code"]).copy()

    if base.empty:
        return base, pd.DataFrame()

    zip_summary = (
        base.groupby("zip_code", as_index=False)
        .agg(
            n_listings=("residual", "size"),
            mean_residual=("residual", "mean"),
            median_residual=("residual", "median"),
            mean_abs_residual=("abs_residual", "mean"),
            mean_pct_error=("pct_error", "mean"),
            median_pct_error=("pct_error", "median"),
            mean_abs_pct_error=("abs_pct_error", "mean"),
            mape_pct=("abs_pct_error", "mean"),
            smape_pct=("smape_pct", "mean")
        )
        .sort_values("mean_abs_residual", ascending=False)
    )

    return base, zip_summary


def _draw_germany_outline_from_geojson(ax):
    """
    Draw Germany polygon from local GeoJSON file.
    Returns True when the outline was drawn.
    """
    geojson_path = Path(__file__).resolve().parent / "src" / "data" / "germany.geojson"
    if not geojson_path.exists():
        return False

    try:
        with geojson_path.open("r", encoding="utf-8") as f:
            geo = json.load(f)
    except Exception:
        return False

    features = geo.get("features", [])
    if not features:
        return False

    geometry = features[0].get("geometry", {})
    gtype = geometry.get("type")
    coords = geometry.get("coordinates", [])
    if not coords:
        return False

    polygons = []
    if gtype == "Polygon":
        polygons = [coords]
    elif gtype == "MultiPolygon":
        polygons = coords
    else:
        return False

    for poly in polygons:
        if not poly:
            continue
        outer_ring = poly[0]
        if len(outer_ring) < 3:
            continue
        xs = [pt[0] for pt in outer_ring]
        ys = [pt[1] for pt in outer_ring]
        ax.fill(xs, ys, color="#f5f5f5", zorder=1)
        ax.plot(xs, ys, color="#666666", linewidth=0.9, zorder=2)

    return True


def _draw_bundesland_boundaries(ax):
    """
    Draw Bundesland internal boundaries from local GeoJSON.
    """
    geojson_path = Path(__file__).resolve().parent / "src" / "data" / "bundeslaender.geojson"
    if not geojson_path.exists():
        return False

    try:
        with geojson_path.open("r", encoding="utf-8") as f:
            geo = json.load(f)
    except Exception:
        return False

    features = geo.get("features", [])
    if not features:
        return False

    for feat in features:
        geometry = feat.get("geometry", {})
        gtype = geometry.get("type")
        coords = geometry.get("coordinates", [])
        if not coords:
            continue

        polygons = []
        if gtype == "Polygon":
            polygons = [coords]
        elif gtype == "MultiPolygon":
            polygons = coords
        else:
            continue

        for poly in polygons:
            if not poly:
                continue
            outer_ring = poly[0]
            if len(outer_ring) < 3:
                continue
            xs = [pt[0] for pt in outer_ring]
            ys = [pt[1] for pt in outer_ring]
            ax.plot(xs, ys, color="#8a8a8a", linewidth=0.45, alpha=0.95, zorder=2)

    return True


def plot_germany_residual_map(
    zip_summary_df,
    scenario_name,
    output_prefix,
    min_listings_per_zip=3,
    max_points=1200,
    color_column="mean_residual",
    cmap="RdBu_r",
    color_label=None,
    symmetric_scale=True,
    clip_percentile_low=2.0,
    clip_percentile_high=98.0,
    vmin=None,
    vmax=None
):
    """
    Plot a Germany map with ZIP centroids colored by mean residual.
    Returns path to the generated PNG or None.
    """
    if zip_summary_df.empty:
        return None

    plot_df = zip_summary_df.copy()
    if min_listings_per_zip > 1:
        plot_df = plot_df.loc[plot_df["n_listings"] >= min_listings_per_zip].copy()

    if plot_df.empty:
        return None

    try:
        import pgeocode  # type: ignore[import-not-found]
    except Exception:
        return None

    nomi = pgeocode.Nominatim("de")
    zip_geo = nomi.query_postal_code(plot_df["zip_code"].tolist())

    geo_df = pd.DataFrame({
        "zip_code": plot_df["zip_code"].values,
        "latitude": zip_geo["latitude"].values,
        "longitude": zip_geo["longitude"].values
    })

    map_df = plot_df.merge(geo_df, on="zip_code", how="left")
    map_df = map_df.dropna(subset=["latitude", "longitude"]).copy()
    if map_df.empty:
        return None

    if len(map_df) > max_points:
        map_df = map_df.nlargest(max_points, "n_listings").copy()

    if color_column not in map_df.columns:
        return None

    color_values = map_df[color_column].to_numpy(dtype=float)
    valid_colors = color_values[np.isfinite(color_values)]
    if valid_colors.size == 0:
        return None

    if clip_percentile_low is not None and clip_percentile_high is not None:
        lo = np.nanpercentile(valid_colors, clip_percentile_low)
        hi = np.nanpercentile(valid_colors, clip_percentile_high)
        if np.isfinite(lo) and np.isfinite(hi) and lo < hi:
            color_values = np.clip(color_values, lo, hi)

    if vmin is None or vmax is None:
        cmin = float(np.nanmin(color_values))
        cmax = float(np.nanmax(color_values))
        if symmetric_scale:
            cabs = max(abs(cmin), abs(cmax))
            vmin_eff, vmax_eff = -cabs, cabs
        else:
            vmin_eff, vmax_eff = cmin, cmax
    else:
        vmin_eff, vmax_eff = float(vmin), float(vmax)

    fig, ax = plt.subplots(figsize=(7.5, 9.2))

    germany_outline_drawn = _draw_germany_outline_from_geojson(ax)
    _draw_bundesland_boundaries(ax)

    size_scale = np.clip(map_df["n_listings"].to_numpy(dtype=float), 1.0, None)
    marker_sizes = 20.0 + 4.0 * np.sqrt(size_scale)

    sc = ax.scatter(
        map_df["longitude"],
        map_df["latitude"],
        c=color_values,
        s=marker_sizes,
        cmap=cmap,
        vmin=vmin_eff,
        vmax=vmax_eff,
        alpha=0.88,
        edgecolors="black",
        linewidths=0.2,
        zorder=3
    )

    cbar = plt.colorbar(sc, ax=ax, pad=0.02, fraction=0.05)
    if color_label is None:
        default_labels = {
            "mean_residual": "Mean residual",
            "mean_abs_residual": "Mean absolute residual",
            "mean_pct_error": "Mean percentage error (%)",
            "mean_abs_pct_error": "Mean absolute percentage error (%)",
            "mape_pct": "MAPE (%)",
            "smape_pct": "sMAPE (%)"
        }
        color_label = default_labels.get(color_column, color_column)
    cbar.set_label(color_label)

    ax.set_title(f"Germany residual map by ZIP - {scenario_name}")
    ax.set_xlabel("Longitude")
    ax.set_ylabel("Latitude")
    ax.grid(alpha=0.25, linewidth=0.5)

    if not germany_outline_drawn:
        ax.set_xlim(5.5, 15.6)
        ax.set_ylim(47.0, 55.2)

    plt.tight_layout()

    safe_name = scenario_name.replace(" ", "_").replace("-", "_")
    output_path = f"{output_prefix}_{safe_name}.png"
    fig.savefig(output_path, dpi=250, bbox_inches="tight")
    plt.close(fig)
    return output_path


# ============================================================
# SCENARIO RUNNER
# ============================================================
def run_model_scenario(df_original, scenario_name, use_imputation, min_plz_obs):
    """
    Run the full MLR / Ridge / Stepwise pipeline for one scenario.

    Scenario structure:
    - N0, N3, N5, N10 based on minimum number of observations per PLZ
    - NO IMPUTATION or WITH IMPUTATION
    """

    add_section(f"SCENARIO: {scenario_name}")
    add_model_target_info()


    # --------------------------------------------------------
    # 1) Prepare data
    # --------------------------------------------------------
    df_model = df_original.copy()

    df_model[TARGET] = pd.to_numeric(df_model[TARGET], errors="coerce")
    df_model[MODEL_TARGET] = pd.to_numeric(df_model[MODEL_TARGET], errors="coerce")

    # Create engineered features before selecting the final feature list
    df_model = add_engineered_features(df_model)

    rows_initial = len(df_model)

    # The model target is never imputed
    df_model = df_model.dropna(subset=[MODEL_TARGET])

    rows_after_target_group_drop = len(df_model)

    X = df_model[FEATURES].copy()
    y = df_model[MODEL_TARGET].copy()
    groups = df_model[GROUP_COL].copy()

    selected_features = FEATURES.copy()
    feature_missing_share_df = create_feature_missing_report(X)

    add_subsection("PLZ Observation Scenario")
    add_text(f"Scenario name: {scenario_name}")
    add_text(f"Minimum number of observations per PLZ: {min_plz_obs}")
    add_text(f"Number of model features before preprocessing: {len(selected_features)}")

    add_text("Model features:")
    add_list(selected_features)
    add_debug_variable_inventory(df_model, selected_features)

    # --------------------------------------------------------
    # 2) Normalize binary variables
    # --------------------------------------------------------
    for col in BINARY_FEATURES:
        if col in X.columns:
            X[col] = normalize_binary_series(X[col])

    # If a binary variable became entirely NA, revert it to categorical
    binary_all_nan = [
        c for c in BINARY_FEATURES
        if c in X.columns and X[c].notna().sum() == 0
    ]

    if binary_all_nan:
        add_subsection("BINARY VARIABLE WARNING")
        add_text(
            "The following binary variables did not contain recognizable values and "
            "were therefore treated as categorical variables:"
        )
        add_list(binary_all_nan)

        for c in binary_all_nan:
            X[c] = df_model[c].astype("object")

    # --------------------------------------------------------
    # 2.1) Global outlier clipping (user-configured)
    # --------------------------------------------------------
    if ENABLE_GLOBAL_OUTLIER_CLIPPING:
        rows_before_outlier_treatment = len(X)

        X, y, clipping_df, excluded_rows_df = apply_global_outlier_clipping(
            X=X,
            y=y,
            tail_pct=OUTLIER_CLIP_TAIL_PCT,
            apply_to_target=OUTLIER_CLIP_APPLY_TO_TARGET,
            mode=OUTLIER_TREATMENT_MODE,
            exclude_features=OUTLIER_CLIP_EXCLUDE_FEATURES
        )

        groups = groups.loc[X.index].copy()

        rows_after_outlier_treatment = len(X)
        rows_removed_outliers = rows_before_outlier_treatment - rows_after_outlier_treatment
        variables_flagged = int((clipping_df["n_rows_flagged"] > 0).sum()) if not clipping_df.empty else 0

        add_subsection("GLOBAL OUTLIER CLIPPING")
        add_text(f"Enabled: {ENABLE_GLOBAL_OUTLIER_CLIPPING}")
        add_text(f"Mode: {OUTLIER_TREATMENT_MODE}")
        add_text(f"Tail percentage per side: {OUTLIER_CLIP_TAIL_PCT}%")
        add_text(f"Target clipping enabled: {OUTLIER_CLIP_APPLY_TO_TARGET}")
        add_text(f"Excluded features: {OUTLIER_CLIP_EXCLUDE_FEATURES}")
        add_text(f"Variables monitored: {len(clipping_df)}")
        add_text(f"Variables that triggered outliers: {variables_flagged}")
        add_text(f"Rows removed by outlier rule: {rows_removed_outliers}")

        selected_outlier_df = build_selected_outlier_summary(
            clipping_df,
            OUTLIER_REPORT_VARIABLES
        )

        add_subsection("SELECTED CLIPPED VARIABLES")
        if selected_outlier_df.empty:
            add_text("No selected variables were found in clipping results.")
        else:
            for _, row in selected_outlier_df.iterrows():
                add_text(
                    f"{row['report_variable']} -> flagged rows: {int(row['n_rows_flagged'])}; "
                    f"limits [{row['low']:.4f}, {row['high']:.4f}]"
                )

        add_text(
            "Detailed outlier listings are hidden to keep the report concise. "
            "Use this summary to compare scenarios."
        )
    else:
        add_subsection("GLOBAL OUTLIER CLIPPING")
        add_text("Enabled: False")

    # --------------------------------------------------------
    # 3) Missing value handling
    # --------------------------------------------------------
    if not use_imputation:
        # No imputation:
        # observations with missing values in any selected model feature are removed.
        valid_mask = X.notna().all(axis=1) & y.notna() & groups.notna()

        X = X.loc[valid_mask].copy()
        y = y.loc[valid_mask].copy()
        groups = groups.loc[valid_mask].copy()

    else:
        # With imputation:
        # only target and group must be non-missing.
        # feature missing values are handled inside the preprocessing pipeline.
        valid_mask = y.notna() & groups.notna()

        X = X.loc[valid_mask].copy()
        y = y.loc[valid_mask].copy()
        groups = groups.loc[valid_mask].copy()

    rows_after_missing_handling = len(X)

    # --------------------------------------------------------
    # 4) Filter by minimum PLZ observations
    # --------------------------------------------------------
    rows_before_plz_filter = len(X)
    groups_before_plz_filter = groups.nunique()

    X, y, groups, removed_plz_groups_df = filter_by_min_plz_observations(
        X=X,
        y=y,
        groups=groups,
        min_plz_obs=min_plz_obs
    )

    rows_after_plz_filter = len(X)
    groups_after_plz_filter = groups.nunique()

    add_subsection("PLZ Group Filtering")
    add_text(f"Rows before PLZ filtering: {rows_before_plz_filter}")
    add_text(f"PLZ groups before filtering: {groups_before_plz_filter}")
    add_text(f"Rows after PLZ filtering: {rows_after_plz_filter}")
    add_text(f"PLZ groups after filtering: {groups_after_plz_filter}")
    add_text(f"Removed PLZ groups: {len(removed_plz_groups_df)}")

    if not removed_plz_groups_df.empty:
        add_text("Examples of removed PLZ groups:")
        for _, row in removed_plz_groups_df.head(20).iterrows():
            add_text(f"{row[GROUP_COL]}: n={row['n_obs']}")

    rows_final = len(X)

    add_subsection("Rows Used")
    add_text(f"Initial rows: {rows_initial}")
    add_text(f"Rows after dropping missing model target/group: {rows_after_target_group_drop}")
    add_text(f"Model target used in this scenario: {MODEL_TARGET}")
    add_text(f"Rows after missing value handling: {rows_after_missing_handling}")
    add_text(f"Rows used after PLZ filtering: {rows_final}")
    add_text(f"Use imputation: {use_imputation}")

    if ZIP_COL_FOR_MAP not in df_model.columns:
        add_text(
            f"Optional residual map warning: '{ZIP_COL_FOR_MAP}' is not available. "
            "The map export will be skipped for this run."
        )

    if rows_final == 0:
        add_text("No rows are available for this scenario. The scenario was skipped.")
        return None

    # --------------------------------------------------------
    # 5) Detect feature types
    # --------------------------------------------------------
    numeric_features = X.select_dtypes(include=["number"]).columns.tolist()
    categorical_features = [c for c in X.columns if c not in numeric_features]

    binary_detected = []
    numeric_non_binary = []

    for col in numeric_features:
        unique_vals = set(X[col].dropna().unique())
        if unique_vals.issubset({0, 1}):
            binary_detected.append(col)
        else:
            numeric_non_binary.append(col)

    add_subsection("VARIABLE TYPE SUMMARY")

    add_text(f"Number of binary variables: {len(binary_detected)}")
    add_list(binary_detected)

    add_text(f"Number of numeric variables (non-binary): {len(numeric_non_binary)}")
    add_list(numeric_non_binary)

    add_text(f"Number of categorical variables: {len(categorical_features)}")
    add_list(categorical_features)

    # --------------------------------------------------------
    # 6) Categorical variables + levels
    # --------------------------------------------------------
    add_subsection("Categorical Variables and Their Levels")

    MAX_LEVELS = 40

    for col in categorical_features:
        add_text(f"--- {col} ---")

        unique_vals = X[col].dropna().unique()

        try:
            unique_vals = sorted(unique_vals)
        except Exception:
            unique_vals = list(unique_vals)

        add_text(f"Number of categories: {len(unique_vals)}")

        if len(unique_vals) <= MAX_LEVELS:
            add_list(unique_vals)
        else:
            add_list(unique_vals[:MAX_LEVELS])
            add_text("... truncated")

    # --------------------------------------------------------
    # 7) Pearson correlation
    # --------------------------------------------------------
    add_pearson_correlation_plot(
        X_input=X,
        y_input=y,
        target=MODEL_TARGET,
        scenario_name=scenario_name,
        use_imputation=use_imputation
    )

    # --------------------------------------------------------
    # 8) Grouped train/test split
    # --------------------------------------------------------
    groups = groups.loc[X.index]

    gss = GroupShuffleSplit(
        n_splits=1,
        test_size=TEST_SIZE,
        random_state=RANDOM_STATE
    )

    train_idx, test_idx = next(gss.split(X, y, groups=groups))

    X_train = X.iloc[train_idx].copy()
    X_test = X.iloc[test_idx].copy()

    y_train = y.iloc[train_idx].copy()
    y_test = y.iloc[test_idx].copy()

    groups_train = groups.iloc[train_idx].copy()
    groups_test = groups.iloc[test_idx].copy()

    add_subsection("Grouped Train-Test Split")
    add_text(f"Grouping variable: {GROUP_COL}")
    add_text(f"Test size: {TEST_SIZE}")
    add_text(f"Training observations: {len(X_train)}")
    add_text(f"Test observations: {len(X_test)}")
    add_text(f"Training ZIP-code groups: {groups_train.nunique()}")
    add_text(f"Test ZIP-code groups: {groups_test.nunique()}")
    add_text(f"Overlapping ZIP-code groups: {len(set(groups_train) & set(groups_test))}")

    # --------------------------------------------------------
    # 9) Preprocessing
    # --------------------------------------------------------
    preprocessor = make_preprocessor(
        numeric_features=numeric_features,
        categorical_features=categorical_features,
        use_imputation=use_imputation
    )

    X_train_proc = preprocessor.fit_transform(X_train)
    X_test_proc = preprocessor.transform(X_test)

    X_train_proc = np.array(X_train_proc, dtype=float)
    X_test_proc = np.array(X_test_proc, dtype=float)

    feature_names_out = get_clean_feature_names(
        preprocessor,
        numeric_features,
        categorical_features
    )

    if len(feature_names_out) != X_train_proc.shape[1]:
        feature_names_out = [f"x_{i}" for i in range(X_train_proc.shape[1])]

    # --------------------------------------------------------
    # 10) OLS
    # --------------------------------------------------------
    X_train_sm = sm.add_constant(X_train_proc, has_constant="add")
    X_test_sm = sm.add_constant(X_test_proc, has_constant="add")

    ols_model = sm.OLS(y_train, X_train_sm).fit()

    # --------------------------------------------------------
    # 11) Ridge
    # --------------------------------------------------------
    ridge_scaler = StandardScaler()
    X_train_ridge = ridge_scaler.fit_transform(X_train_proc)
    X_test_ridge = ridge_scaler.transform(X_test_proc)

    ridge_model = RidgeCV(alphas=RIDGE_ALPHAS, cv=5)
    ridge_model.fit(X_train_ridge, y_train)

    y_pred_ridge = ridge_model.predict(X_test_ridge)
    ridge_best_alpha = float(ridge_model.alpha_)

    ridge_coef_df = pd.DataFrame({
        "feature": feature_names_out,
        "coef_scaled_space": ridge_model.coef_
    })

    ridge_coef_df["abs_coef"] = ridge_coef_df["coef_scaled_space"].abs()
    ridge_coef_df = ridge_coef_df.sort_values("abs_coef", ascending=False)

    add_section(f"RIDGE SUMMARY - {scenario_name}")
    add_text(f"Best alpha (RidgeCV): {ridge_best_alpha:.6f}")
    add_text("Note: Ridge does not provide p-values like OLS.")
    add_text("The coefficients shown below are in standardized feature space.")

    add_subsection("Top Ridge Coefficients (Table)")
    doc.add_paragraph(ridge_coef_df.head(20).to_string(index=False))

    # --------------------------------------------------------
    # 12) Ridge coefficient plot
    # --------------------------------------------------------
    add_section(f"TOP RIDGE COEFFICIENTS - {scenario_name}")

    top_ridge_plot = ridge_coef_df.head(20).sort_values("abs_coef", ascending=True)

    fig, ax = plt.subplots(figsize=(10, 8))
    ax.barh(top_ridge_plot["feature"], top_ridge_plot["abs_coef"], color="steelblue")
    ax.set_xlabel("|Standardized coefficient|")
    ax.set_ylabel("Feature")
    ax.set_title(f"Top 20 Absolute Ridge Coefficients - {scenario_name}")
    plt.tight_layout()

    add_figure_to_doc(fig, width=6.0)
    plt.close(fig)

    # --------------------------------------------------------
    # 13) Ridge regularization path
    # --------------------------------------------------------
    add_section(f"RIDGE REGULARIZATION PATH - {scenario_name}")

    alphas_path = np.logspace(-3, 3, 30)
    coefs_path = []

    for alpha in alphas_path:
        ridge_tmp = Ridge(alpha=alpha)
        ridge_tmp.fit(X_train_ridge, y_train)
        coefs_path.append(ridge_tmp.coef_)

    coefs_path = np.array(coefs_path)

    top_features = ridge_coef_df.head(10)["feature"].tolist()
    top_idx = [feature_names_out.index(f) for f in top_features if f in feature_names_out]

    fig, ax = plt.subplots(figsize=(9, 6))
    for idx in top_idx:
        ax.plot(alphas_path, coefs_path[:, idx], label=feature_names_out[idx])

    ax.set_xscale("log")
    ax.set_xlabel("Alpha (log scale)")
    ax.set_ylabel("Standardized coefficient")
    ax.set_title(f"Ridge Regularization Path - {scenario_name}")
    ax.legend(fontsize=8)
    plt.tight_layout()

    add_figure_to_doc(fig, width=6.0)
    plt.close(fig)

    # --------------------------------------------------------
    # 14) RESET
    # --------------------------------------------------------
    add_section(f"MODEL SPECIFICATION / EXOGENEITY PROXY - {scenario_name}")

    try:
        reset_res = linear_reset(ols_model, power=2, use_f=True)

        reset_f = float(reset_res.fvalue)
        reset_p = float(reset_res.pvalue)

        add_text("Ramsey RESET test:")
        add_text(f"F-statistic: {reset_f:.6f}")
        add_text(f"p-value: {reset_p:.6f}")
        add_text(
            "Note: RESET does not prove exogeneity or endogeneity; "
            "it only signals possible model misspecification."
        )

    except Exception as exc:
        add_text(f"Ramsey RESET could not be computed: {exc}")
        reset_f = np.nan
        reset_p = np.nan

    # --------------------------------------------------------
    # 15) OLS coefficients
    # --------------------------------------------------------
    add_section(f"OLS COEFFICIENTS - {scenario_name}")

    coef_df = pd.DataFrame({
        "feature": ["const"] + feature_names_out,
        "coef": ols_model.params,
        "p_value": ols_model.pvalues
    })

    add_subsection("OLS Coefficients Table")
    doc.add_paragraph(coef_df.to_string(index=False))

    # --------------------------------------------------------
    # 16) Model performance
    # --------------------------------------------------------
    y_pred_train_ols = ols_model.predict(X_train_sm)
    y_pred_test_ols = ols_model.predict(X_test_sm)

    add_mlr_prediction_examples(
        scenario_name=scenario_name,
        ols_model=ols_model,
        X_test_proc=X_test_proc,
        X_test_raw=X_test,
        y_test=y_test,
        y_pred_test=y_pred_test_ols,
        feature_names_out=feature_names_out,
        n_examples=EXAMPLE_N_CASES,
        top_contrib=8
    )

    ols_train_rmse = float(np.sqrt(mean_squared_error(y_train, y_pred_train_ols)))
    ols_train_mae = float(mean_absolute_error(y_train, y_pred_train_ols))
    ols_train_r2 = float(r2_score(y_train, y_pred_train_ols))

    ols_test_rmse = float(np.sqrt(mean_squared_error(y_test, y_pred_test_ols)))
    ols_test_mae = float(mean_absolute_error(y_test, y_pred_test_ols))
    ols_test_r2 = float(r2_score(y_test, y_pred_test_ols))

    ols_rmse_gap = ols_test_rmse - ols_train_rmse
    ols_r2_gap = ols_train_r2 - ols_test_r2

    y_pred_train_ridge = ridge_model.predict(X_train_ridge)
    y_pred_test_ridge = ridge_model.predict(X_test_ridge)

    ridge_train_rmse = float(np.sqrt(mean_squared_error(y_train, y_pred_train_ridge)))
    ridge_train_mae = float(mean_absolute_error(y_train, y_pred_train_ridge))
    ridge_train_r2 = float(r2_score(y_train, y_pred_train_ridge))

    ridge_test_rmse = float(np.sqrt(mean_squared_error(y_test, y_pred_test_ridge)))
    ridge_test_mae = float(mean_absolute_error(y_test, y_pred_test_ridge))
    ridge_test_r2 = float(r2_score(y_test, y_pred_test_ridge))

    ridge_rmse_gap = ridge_test_rmse - ridge_train_rmse
    ridge_r2_gap = ridge_train_r2 - ridge_test_r2

    # --------------------------------------------------------
    # 17) Stepwise regression
    # --------------------------------------------------------
    try:
        stepwise_output = run_backward_stepwise(
            X_train_proc=X_train_proc,
            X_test_proc=X_test_proc,
            y_train=y_train,
            y_test=y_test,
            feature_names_out=feature_names_out,
            scenario_name=scenario_name,
            p_threshold=STEPWISE_P_THRESHOLD,
            max_iter=STEPWISE_MAX_ITER
        )

        stepwise_rmse = stepwise_output["stepwise_rmse"]
        stepwise_mae = stepwise_output["stepwise_mae"]
        stepwise_r2 = stepwise_output["stepwise_r2"]

        stepwise_selected_df = stepwise_output["stepwise_selected_df"]
        stepwise_removed_df = stepwise_output["stepwise_removed_df"]
        stepwise_coef_df = stepwise_output["stepwise_coef_df"]

    except Exception as exc:
        add_section(f"STEPWISE REGRESSION - {scenario_name}")
        add_text(f"Stepwise regression failed: {exc}")

        stepwise_rmse = np.nan
        stepwise_mae = np.nan
        stepwise_r2 = np.nan

        stepwise_selected_df = pd.DataFrame()
        stepwise_removed_df = pd.DataFrame()
        stepwise_coef_df = pd.DataFrame()

    add_section(f"METRICS AND OVERFITTING CHECK - {scenario_name}")

    add_text("OLS:")
    add_text(f"Train RMSE: {ols_train_rmse:.4f}")
    add_text(f"Test RMSE: {ols_test_rmse:.4f}")
    add_text(f"RMSE gap: {ols_rmse_gap:.4f}")
    add_text(f"Train R²: {ols_train_r2:.4f}")
    add_text(f"Test R²: {ols_test_r2:.4f}")
    add_text(f"R² gap: {ols_r2_gap:.4f}")

    add_text("Ridge:")
    add_text(f"Best alpha: {ridge_best_alpha:.6f}")
    add_text(f"Train RMSE: {ridge_train_rmse:.4f}")
    add_text(f"Test RMSE: {ridge_test_rmse:.4f}")
    add_text(f"RMSE gap: {ridge_rmse_gap:.4f}")
    add_text(f"Train R²: {ridge_train_r2:.4f}")
    add_text(f"Test R²: {ridge_test_r2:.4f}")
    add_text(f"R² gap: {ridge_r2_gap:.4f}")

    # --------------------------------------------------------
    # 18) OLS vs Ridge vs Stepwise metrics plot
    # --------------------------------------------------------
    add_section(f"OLS VS RIDGE VS STEPWISE - METRICS - {scenario_name}")

    metrics_plot_df = pd.DataFrame({
        "Metric": ["RMSE", "MAE", "R2"],
        "OLS": [ols_test_rmse, ols_test_mae, ols_test_r2],
        "Ridge": [ridge_test_rmse, ridge_test_mae, ridge_test_r2],
        "Stepwise": [stepwise_rmse, stepwise_mae, stepwise_r2]
    })

    x = np.arange(len(metrics_plot_df))
    width = 0.25

    fig, ax = plt.subplots(figsize=(9, 5))
    ax.bar(x - width, metrics_plot_df["OLS"], width, label="OLS")
    ax.bar(x, metrics_plot_df["Ridge"], width, label="Ridge")
    ax.bar(x + width, metrics_plot_df["Stepwise"], width, label="Stepwise")

    ax.set_xticks(x)
    ax.set_xticklabels(metrics_plot_df["Metric"])
    ax.set_ylabel("Value")
    ax.set_title(f"Metric Comparison: OLS vs Ridge vs Stepwise - {scenario_name}")
    ax.legend()
    plt.tight_layout()

    add_figure_to_doc(fig, width=6.0)
    plt.close(fig)

    # --------------------------------------------------------
    # 19) VIF
    # --------------------------------------------------------
    add_section(f"VIF - {scenario_name}")

    X_vif = X_train_proc.copy()

    if VIF_SAMPLE_SIZE and X_vif.shape[0] > VIF_SAMPLE_SIZE:
        rng = np.random.default_rng(RANDOM_STATE)
        idx = rng.choice(X_vif.shape[0], size=VIF_SAMPLE_SIZE, replace=False)
        X_vif = X_vif[idx, :]

    vif_values = []

    for i in range(X_vif.shape[1]):
        try:
            vif_val = variance_inflation_factor(X_vif, i)
        except Exception:
            vif_val = np.inf

        vif_values.append(vif_val)

    vif_df = pd.DataFrame({
        "feature": feature_names_out,
        "VIF": vif_values
    }).sort_values("VIF", ascending=False)

    add_text("Top 20 highest VIF values:")
    for _, row in vif_df.head(20).iterrows():
        add_text(f"{row['feature']}: {row['VIF']:.4f}")

    top_vif_plot = vif_df.head(20).sort_values("VIF", ascending=True)

    fig, ax = plt.subplots(figsize=(10, 8))
    bars = ax.barh(top_vif_plot["feature"], top_vif_plot["VIF"], color="slateblue")
    ax.axvline(5, color="orange", linestyle="--", linewidth=1, label="VIF = 5")
    ax.axvline(10, color="red", linestyle="--", linewidth=1, label="VIF = 10")

    ax.set_xlabel("Variance Inflation Factor (VIF)")
    ax.set_ylabel("Feature")
    ax.set_title(f"Top 20 VIF Values - {scenario_name}")
    ax.legend()
    ax.tick_params(axis="x", labelsize=10)
    ax.tick_params(axis="y", labelsize=10)
    
    for bar in bars:
        width_val = bar.get_width()
        ax.text(
            width_val + 0.05,
            bar.get_y() + bar.get_height() / 2,
            f"{width_val:.2f}",
            va="center",
            fontsize=10
        )

    plt.tight_layout()

    add_figure_to_doc(fig, width=6.2)
    plt.close(fig)

    # --------------------------------------------------------
    # 20) Assumption checks
    # --------------------------------------------------------
    add_section(f"ASSUMPTION CHECKS - {scenario_name}")

    residuals = y_test - y_pred_test_ols
    fitted_vals = y_pred_test_ols

    if RESIDUAL_MAP_SCOPE == "all_rows":
        X_all_proc = np.array(preprocessor.transform(X), dtype=float)
        X_all_sm = sm.add_constant(X_all_proc, has_constant="add")
        y_for_map = y.copy()
        y_pred_for_map = ols_model.predict(X_all_sm)
    else:
        y_for_map = y_test.copy()
        y_pred_for_map = y_pred_test_ols

    residuals_detail_df, residuals_by_zip_df = build_residual_zip_table(
        df_model=df_model,
        y_test=y_for_map,
        y_pred_test_ols=y_pred_for_map,
        scenario_name=scenario_name
    )

    # Linearity
    try:
        rainbow_stat, rainbow_pvalue = linear_rainbow(ols_model)
        rainbow_stat = float(rainbow_stat)
        rainbow_pvalue = float(rainbow_pvalue)
    except Exception as exc:
        add_text(f"Rainbow test could not be computed: {exc}")
        rainbow_stat = np.nan
        rainbow_pvalue = np.nan

    # Homoscedasticity
    try:
        bp_test = het_breuschpagan(residuals, X_test_sm)
        bp_labels = ["LM Statistic", "LM-Test p-value", "F-Statistic", "F-Test p-value"]
        bp_results = dict(zip(bp_labels, bp_test))

        bp_lm_stat = float(bp_results["LM Statistic"])
        bp_pvalue = float(bp_results["LM-Test p-value"])
        bp_f_stat = float(bp_results["F-Statistic"])
        bp_f_pvalue = float(bp_results["F-Test p-value"])

    except Exception as exc:
        add_text(f"Breusch-Pagan test could not be computed: {exc}")
        bp_lm_stat = np.nan
        bp_pvalue = np.nan
        bp_f_stat = np.nan
        bp_f_pvalue = np.nan

    # Normality
    try:
        jb_stat, jb_pvalue = stats.jarque_bera(residuals)
        jb_stat = float(jb_stat)
        jb_pvalue = float(jb_pvalue)
    except Exception as exc:
        add_text(f"Jarque-Bera test could not be computed: {exc}")
        jb_stat = np.nan
        jb_pvalue = np.nan

    try:
        sample_size = min(5000, len(residuals))
        resid_sample = pd.Series(residuals).sample(sample_size, random_state=RANDOM_STATE)
        shapiro_stat, shapiro_p = stats.shapiro(resid_sample)

        shapiro_stat = float(shapiro_stat)
        shapiro_p = float(shapiro_p)
    except Exception as exc:
        add_text(f"Shapiro-Wilk test could not be computed: {exc}")
        shapiro_stat = np.nan
        shapiro_p = np.nan
        sample_size = np.nan

    # Independence
    try:
        dw_stat = float(durbin_watson(residuals))
    except Exception as exc:
        add_text(f"Durbin-Watson statistic could not be computed: {exc}")
        dw_stat = np.nan

    add_subsection("Diagnostic Test Summary")

    add_text("Linearity")
    add_text(f"Rainbow statistic: {rainbow_stat:.6f}" if not pd.isna(rainbow_stat) else "Rainbow statistic: NA")
    add_text(f"Rainbow p-value: {rainbow_pvalue:.6f}" if not pd.isna(rainbow_pvalue) else "Rainbow p-value: NA")

    add_text("Homoscedasticity")
    add_text(f"Breusch-Pagan LM statistic: {bp_lm_stat:.6f}" if not pd.isna(bp_lm_stat) else "Breusch-Pagan LM statistic: NA")
    add_text(f"Breusch-Pagan LM p-value: {bp_pvalue:.6f}" if not pd.isna(bp_pvalue) else "Breusch-Pagan LM p-value: NA")
    add_text(f"Breusch-Pagan F statistic: {bp_f_stat:.6f}" if not pd.isna(bp_f_stat) else "Breusch-Pagan F statistic: NA")
    add_text(f"Breusch-Pagan F p-value: {bp_f_pvalue:.6f}" if not pd.isna(bp_f_pvalue) else "Breusch-Pagan F p-value: NA")

    add_text("Normality")
    add_text(f"Jarque-Bera statistic: {jb_stat:.6f}" if not pd.isna(jb_stat) else "Jarque-Bera statistic: NA")
    add_text(f"Jarque-Bera p-value: {jb_pvalue:.6f}" if not pd.isna(jb_pvalue) else "Jarque-Bera p-value: NA")
    add_text(
        f"Shapiro-Wilk statistic (sample = {sample_size}): {shapiro_stat:.6f}"
        if not pd.isna(shapiro_stat) else
        f"Shapiro-Wilk statistic (sample = {sample_size}): NA"
    )
    add_text(
        f"Shapiro-Wilk p-value (sample = {sample_size}): {shapiro_p:.6f}"
        if not pd.isna(shapiro_p) else
        f"Shapiro-Wilk p-value (sample = {sample_size}): NA"
    )

    add_text("Independence")
    add_text(f"Durbin-Watson statistic: {dw_stat:.6f}" if not pd.isna(dw_stat) else "Durbin-Watson statistic: NA")
    add_text("Rule of thumb: values close to 2 suggest residual independence.")

    # Diagnostic plots
    add_subsection("Unified Diagnostic Panel")

    save_individual_diagnostic_plots(
        fitted_vals=fitted_vals,
        residuals=residuals,
        rainbow_pvalue=rainbow_pvalue,
        shapiro_pvalue=shapiro_p,
        bp_pvalue=bp_pvalue,
        dw_stat=dw_stat,
        scenario_name=scenario_name
    )

    add_text(
        "This unified panel summarizes the main OLS assumption checks. "
        "Each subplot includes the corresponding p-value or statistic and an automated visual status label."
    )
    add_text(
        "How to read the panel for thesis interpretation: "
        "Residuals vs Fitted should look like a horizontal cloud around zero; "
        "Q-Q points should stay close to the diagonal; "
        "Scale-Location should be relatively flat; "
        "Residuals vs Order should oscillate around zero without clear trend. "
        "When these patterns are not observed, the assumption may be violated."
    )

    # Residual distribution
    add_subsection("Residual Distribution")

    fig, ax = plt.subplots(figsize=(8, 5))
    sns.histplot(residuals, kde=True, color="steelblue", ax=ax)
    ax.set_xlabel("Residuals")
    ax.set_ylabel("Frequency")
    ax.set_title(f"Residual Distribution - {scenario_name}")
    plt.tight_layout()

    add_figure_to_doc(fig, width=6.2)
    plt.close(fig)

    # Residual ACF
    add_subsection("Residual Autocorrelation Function")

    try:
        fig, ax = plt.subplots(figsize=(8, 5))
        plot_acf(pd.Series(residuals).dropna(), lags=40, ax=ax)
        ax.set_title(f"ACF of Residuals - {scenario_name}")
        plt.tight_layout()

        add_text(
            "The residual ACF can help identify autocorrelation patterns. "
            "For cross-sectional data, this plot should be interpreted with caution."
        )

        add_figure_to_doc(fig, width=6.2)
        plt.close(fig)

    except Exception as exc:
        add_text(f"The residual ACF could not be computed: {exc}")

    # --------------------------------------------------------
    # 21) Export objects
    # --------------------------------------------------------

    metrics_df = pd.DataFrame({
        "metric": [
            "scenario",
            "min_plz_obs",
            "use_imputation",
            "n_obs",
            "n_train",
            "n_test",
            "n_train_zip_groups",
            "n_test_zip_groups",

            "OLS_Train_RMSE",
            "OLS_Test_RMSE",
            "OLS_Train_MAE",
            "OLS_Test_MAE",
            "OLS_Train_R2",
            "OLS_Test_R2",

            "Ridge_Best_Alpha",
            "Ridge_Train_RMSE",
            "Ridge_Test_RMSE",
            "Ridge_Train_MAE",
            "Ridge_Test_MAE",
            "Ridge_Train_R2",
            "Ridge_Test_R2",

            "Stepwise_Test_RMSE",
            "Stepwise_Test_MAE",
            "Stepwise_Test_R2",

            "Jarque_Bera_p",
            "Shapiro_p",
            "Durbin_Watson",
            "BP_pvalue",
            "RESET_F",
            "RESET_pvalue",
            "Rainbow_pvalue"
        ],
        "value": [
            scenario_name,
            min_plz_obs,
            use_imputation,
            rows_final,
            len(y_train),
            len(y_test),
            groups_train.nunique(),
            groups_test.nunique(),

            ols_train_rmse,
            ols_test_rmse,
            ols_train_mae,
            ols_test_mae,
            ols_train_r2,
            ols_test_r2,

            ridge_best_alpha,
            ridge_train_rmse,
            ridge_test_rmse,
            ridge_train_mae,
            ridge_test_mae,
            ridge_train_r2,
            ridge_test_r2,

            stepwise_rmse,
            stepwise_mae,
            stepwise_r2,

            jb_pvalue,
            shapiro_p,
            dw_stat,
            bp_pvalue,
            reset_f,
            reset_p,
            rainbow_pvalue
        ]
    })

    ridge_coef_export = ridge_coef_df.drop(columns=["abs_coef"]).copy()

    result_row = {
        "scenario": scenario_name,
        "model_target": MODEL_TARGET,
        "model_target_label": MODEL_TARGET_LABEL,
        "model_target_unit": MODEL_TARGET_UNIT,
        
        "min_plz_obs": min_plz_obs,
        "use_imputation": use_imputation,

        "n_obs": rows_final,
        "n_train": len(y_train),
        "n_test": len(y_test),
        "n_train_zip_groups": groups_train.nunique(),
        "n_test_zip_groups": groups_test.nunique(),
        "n_overlap_zip_groups": len(set(groups_train) & set(groups_test)),

        "rows_initial": rows_initial,
        "rows_after_target_group_drop": rows_after_target_group_drop,
        "rows_after_missing_handling": rows_after_missing_handling,
        "rows_before_plz_filter": rows_before_plz_filter,
        "rows_after_plz_filter": rows_after_plz_filter,
        "groups_before_plz_filter": groups_before_plz_filter,
        "groups_after_plz_filter": groups_after_plz_filter,
        "removed_plz_groups": len(removed_plz_groups_df),

        "OLS_Train_RMSE": ols_train_rmse,
        "OLS_Test_RMSE": ols_test_rmse,
        "OLS_RMSE_Gap": ols_rmse_gap,
        "OLS_Train_MAE": ols_train_mae,
        "OLS_Test_MAE": ols_test_mae,
        "OLS_Train_R2": ols_train_r2,
        "OLS_Test_R2": ols_test_r2,
        "OLS_R2_Gap": ols_r2_gap,

        "Ridge_Best_Alpha": ridge_best_alpha,
        "Ridge_Train_RMSE": ridge_train_rmse,
        "Ridge_Test_RMSE": ridge_test_rmse,
        "Ridge_RMSE_Gap": ridge_rmse_gap,
        "Ridge_Train_MAE": ridge_train_mae,
        "Ridge_Test_MAE": ridge_test_mae,
        "Ridge_Train_R2": ridge_train_r2,
        "Ridge_Test_R2": ridge_test_r2,
        "Ridge_R2_Gap": ridge_r2_gap,

        "Stepwise_Test_RMSE": stepwise_rmse,
        "Stepwise_Test_MAE": stepwise_mae,
        "Stepwise_Test_R2": stepwise_r2,

        "Jarque_Bera_p": jb_pvalue,
        "Shapiro_p": shapiro_p,
        "Durbin_Watson": dw_stat,
        "BP_pvalue": bp_pvalue,
        "RESET_pvalue": reset_p,
        "Rainbow_pvalue": rainbow_pvalue
    }

    export_objects = {
        "metrics_df": metrics_df,
        "coef_df": coef_df,
        "ridge_coef_df": ridge_coef_export,
        "vif_df": vif_df,
        "result_row": result_row,
        "stepwise_selected_df": stepwise_selected_df,
        "stepwise_removed_df": stepwise_removed_df,
        "stepwise_coef_df": stepwise_coef_df,
        "feature_missing_share_df": feature_missing_share_df,
        "removed_plz_groups_df": removed_plz_groups_df,
        "residuals_detail_df": residuals_detail_df,
        "residuals_by_zip_df": residuals_by_zip_df,
        "min_plz_obs": min_plz_obs
    }

    return export_objects

# ============================================================
# LOAD DATA
# ============================================================

# #metodo para target = total rent
# READ_COLS = list(dict.fromkeys(BASE_FEATURES + RAW_ENGINEERING_FEATURES + [TARGET, GROUP_COL]))

# ============================================================
# LOAD DATA
# ============================================================

READ_COLS = list(dict.fromkeys(
    BASE_FEATURES +
    RAW_ENGINEERING_FEATURES +
    [
        TARGET,
        "obj_livingSpace",
        ZIP_COL_FOR_MAP,
        GROUP_COL
    ]
))

df = pd.read_csv(
    FILE_PATH,
    usecols=READ_COLS,
    sep=CSV_SEP,
    low_memory=False
)

df["obj_totalRent_num"] = pd.to_numeric(df[TARGET], errors="coerce")
df["obj_livingSpace_num"] = pd.to_numeric(df["obj_livingSpace"], errors="coerce")

df[TARGET_RENT_SQM] = (
    df["obj_totalRent_num"] /
    df["obj_livingSpace_num"]
)

df.loc[
    (df["obj_livingSpace_num"] <= 0) |
    (~np.isfinite(df[TARGET_RENT_SQM])),
    TARGET_RENT_SQM
] = np.nan

add_global_model_input_max_summary(df)


# # ============================================================
# # RUN BOTH SCENARIOS
# # ============================================================
# SCENARIOS = []

# for threshold_scenario in MISSING_THRESHOLD_SCENARIOS:
#     for use_imputation in [False, True]:
#         SCENARIOS.append({
#             "scenario_code": threshold_scenario["scenario_code"],
#             "max_missing_pct": threshold_scenario["max_missing_pct"],
#             "name": f"{threshold_scenario['scenario_code']} - {'WITH IMPUTATION' if use_imputation else 'NO IMPUTATION'}",
#             "use_imputation": use_imputation
#         })

# all_results = []
# all_exports = {}

# for scenario in SCENARIOS:
#     scenario_name = scenario["name"]
#     use_imputation = scenario["use_imputation"]

#     print(f"\nRunning scenario: {scenario_name}")

#     scenario_export = run_model_scenario(
#         df_original=df,
#         scenario_name=scenario_name,
#         use_imputation=use_imputation
#     )

#     if scenario_export is not None:
#         all_exports[scenario_name] = scenario_export
#         all_results.append(scenario_export["result_row"])

# ============================================================
# BUILD SCENARIOS
# ============================================================

# ============================================================
# BUILD SCENARIOS
# ============================================================

SCENARIOS = []

if DEBUG_MODE:
    selected_plz_scenario = None

    for plz_scenario in PLZ_OBS_SCENARIOS:
        if plz_scenario["scenario_code"] == DEBUG_SCENARIO_CODE:
            selected_plz_scenario = plz_scenario
            break

    if selected_plz_scenario is None:
        raise ValueError(
            f"DEBUG_SCENARIO_CODE='{DEBUG_SCENARIO_CODE}' is not valid. "
            "Use one of: N0, N3, N5, N10."
        )

    scenario_label = "WITH IMPUTATION" if DEBUG_USE_IMPUTATION else "NO IMPUTATION"

    SCENARIOS.append({
        "scenario_code": selected_plz_scenario["scenario_code"],
        "min_plz_obs": selected_plz_scenario["min_plz_obs"],
        "name": f"{selected_plz_scenario['scenario_code']} - {scenario_label}",
        "use_imputation": DEBUG_USE_IMPUTATION
    })

else:
    for plz_scenario in PLZ_OBS_SCENARIOS:
        for use_imputation in IMPUTATION_SCENARIOS:

            scenario_label = "WITH IMPUTATION" if use_imputation else "NO IMPUTATION"

            SCENARIOS.append({
                "scenario_code": plz_scenario["scenario_code"],
                "min_plz_obs": plz_scenario["min_plz_obs"],
                "name": f"{plz_scenario['scenario_code']} - {scenario_label}",
                "use_imputation": use_imputation
            })





all_results = []
all_exports = {}

for scenario in SCENARIOS:
    scenario_name = scenario["name"]
    use_imputation = scenario["use_imputation"]
    min_plz_obs = scenario["min_plz_obs"]

    print(f"\nRunning scenario: {scenario_name}")

    scenario_export = run_model_scenario(
        df_original=df,
        scenario_name=scenario_name,
        use_imputation=use_imputation,
        min_plz_obs=min_plz_obs
    )

    if scenario_export is not None:
        all_exports[scenario_name] = scenario_export
        all_results.append(scenario_export["result_row"])


# ============================================================
# FINAL COMPARISON
# ============================================================
comparison_df = pd.DataFrame(all_results)

add_section(f"FINAL COMPARISON - {MODEL_TARGET_LABEL}")

if DEBUG_MODE:
    add_text("Debug mode was enabled. Only one scenario was executed.")
else:
    add_text("Full mode was enabled. All PLZ observation and imputation scenarios were executed.")

if comparison_df.empty:
    add_text("No scenario returned valid results.")
else:
    add_text("Comparison table between the two scenarios:")
    add_text(comparison_df.to_string(index=False))

    # --------------------------------------------------------
    # Comparison plot: RMSE and MAE
    # --------------------------------------------------------
    add_subsection("Comparison Plot - RMSE and MAE")

    error_metrics_df = comparison_df.set_index("scenario")[
    [
        "OLS_Test_RMSE",
        "Ridge_Test_RMSE",
        "Stepwise_Test_RMSE",
        "OLS_Test_MAE",
        "Ridge_Test_MAE",
        "Stepwise_Test_MAE"
    ]
]

    fig, ax = plt.subplots(figsize=(12, 6))
    error_metrics_df.plot(kind="bar", ax=ax)
    ax.set_title("Error Metric Comparison Across Scenarios")
    ax.set_ylabel("Metric value")
    ax.set_xlabel("Scenario")
    ax.tick_params(axis="x", rotation=45)
    plt.tight_layout()

    add_figure_to_doc(fig, width=6.0)
    plt.close(fig)

    # --------------------------------------------------------
    # Comparison plot: R²
    # --------------------------------------------------------
    add_subsection("Comparison Plot - R²")

    r2_metrics_df = comparison_df.set_index("scenario")[
     [
        "OLS_Test_R2",
        "Ridge_Test_R2",
        "Stepwise_Test_R2"
     ]   
    ]

    fig, ax = plt.subplots(figsize=(10, 5))
    r2_metrics_df.plot(kind="bar", ax=ax)
    ax.set_title("R² Comparison Across Scenarios")
    ax.set_ylabel("R²")
    ax.set_xlabel("Scenario")
    ax.tick_params(axis="x", rotation=0)
    plt.tight_layout()

    add_figure_to_doc(fig, width=6.0)
    plt.close(fig)


# ============================================================
# EXPORT WORD
# ============================================================
doc.save(DOCX_OUTPUT)


# ============================================================
# EXPORT EXCEL
# ============================================================
with pd.ExcelWriter(OUTPUT_EXCEL, engine="openpyxl") as writer:

    if not comparison_df.empty:
        comparison_df.to_excel(writer, index=False, sheet_name="comparison")

    for scenario_name, export_data in all_exports.items():
        clean_name = scenario_name.replace(" ", "_").replace("-", "_")

        export_data["metrics_df"].to_excel(
            writer,
            index=False,
            sheet_name=f"{clean_name}_metrics"[:31]
        )

        export_data["coef_df"].to_excel(
            writer,
            index=False,
            sheet_name=f"{clean_name}_ols_coef"[:31]
        )

        export_data["ridge_coef_df"].to_excel(
            writer,
            index=False,
            sheet_name=f"{clean_name}_ridge_coef"[:31]
        )

        export_data["vif_df"].to_excel(
            writer,
            index=False,
            sheet_name=f"{clean_name}_VIF"[:31]
        )

        export_data["stepwise_coef_df"].to_excel(
            writer,
            index=False,
            sheet_name=f"{clean_name}_step_coef"[:31]
        )

        export_data["stepwise_selected_df"].to_excel(
            writer,
            index=False,
            sheet_name=f"{clean_name}_step_selected"[:31]
        )

        export_data["stepwise_removed_df"].to_excel(
            writer,
            index=False,
            sheet_name=f"{clean_name}_step_removed"[:31]
        )
        export_data["feature_missing_share_df"].to_excel(
            writer,
            index=False,
            sheet_name=f"{clean_name}_missing"[:31]
        )

        export_data["removed_plz_groups_df"].to_excel(
            writer,
            index=False,
            sheet_name=f"{clean_name}_removed_plz"[:31]
            )

        export_data["residuals_detail_df"].to_excel(
            writer,
            index=False,
            sheet_name=f"{clean_name}_residuals"[:31]
        )

        export_data["residuals_by_zip_df"].to_excel(
            writer,
            index=False,
            sheet_name=f"{clean_name}_zip_resid"[:31]
        )


if ENABLE_GERMANY_RESIDUAL_MAP:
    add_section("GERMANY RESIDUAL MAPS")
    for scenario_name, export_data in all_exports.items():
        map_path = plot_germany_residual_map(
            zip_summary_df=export_data["residuals_by_zip_df"],
            scenario_name=scenario_name,
            output_prefix="residual_map_de",
            min_listings_per_zip=RESIDUAL_MAP_MIN_LISTINGS_PER_ZIP,
            max_points=RESIDUAL_MAP_MAX_POINTS,
            color_column=RESIDUAL_MAP_COLOR_COLUMN,
            cmap=RESIDUAL_MAP_CMAP,
            color_label=RESIDUAL_MAP_COLOR_LABEL,
            symmetric_scale=RESIDUAL_MAP_SYMMETRIC_SCALE,
            clip_percentile_low=RESIDUAL_MAP_CLIP_PCT_LOW,
            clip_percentile_high=RESIDUAL_MAP_CLIP_PCT_HIGH,
        )

        if map_path is None:
            add_text(
                f"Map skipped for '{scenario_name}'. "
                "Possible reasons: no ZIP data after filtering, or missing optional package 'pgeocode'."
            )
            continue

        add_subsection(f"Residual map - {scenario_name}")
        add_text(
            f"Color metric: {RESIDUAL_MAP_COLOR_COLUMN}. "
            f"{RESIDUAL_MAP_PALETTE_TEXT} "
            "Marker size scales with the number of listings in the ZIP area."
        )
        doc.add_picture(map_path, width=Inches(6.0))


# ============================================================
# PRINT RESULTS
# ============================================================
print(f"\nWord report saved: {DOCX_OUTPUT}")
print(f"Excel report saved: {OUTPUT_EXCEL}")

print("\nFinal comparison:")

if comparison_df.empty:
    print("No valid results were produced.")
else:
    print(comparison_df.to_string(index=False))

elapsed_seconds = perf_counter() - SCRIPT_START_TIME
elapsed_minutes = elapsed_seconds / 60.0
print(f"\nTotal runtime: {elapsed_seconds:.2f} seconds ({elapsed_minutes:.2f} minutes)")
